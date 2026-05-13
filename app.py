"""
Personal RAG Chatbot — 2026 stack.

Features
--------
- OpenAI-compatible chat completions (HF Router / DashScope / vLLM / local)
- Streaming responses with reasoning_content split (Gemma-4 / R1 / Qwen3)
- BGE-M3 (default) or MiniLM multilingual embeddings — sidebar selectable
- Hybrid retrieval: BM25 + dense + Reciprocal Rank Fusion
- Cross-encoder reranker (BAAI/bge-reranker-v2-m3) — optional
- Persistent vector store on disk (`./.data/{embedder}/{doc_hash}/`)
- Inline [N] citation parsing + per-chunk popover annotations
- HF Hub cache list/delete
- Configurable chunking (size, overlap)
"""

import base64
import datetime
import hashlib
import json
import os
import re
import shutil
import uuid
import warnings
from pathlib import Path

# Silence transformers deprecation logs before any heavy import.
os.environ.setdefault('TRANSFORMERS_VERBOSITY', 'error')
os.environ.setdefault('TRANSFORMERS_NO_ADVISORY_WARNINGS', '1')
warnings.filterwarnings('ignore', category=FutureWarning)
warnings.filterwarnings('ignore', message=r'.*__path__.*')

import numpy as np
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv(Path(__file__).parent / '.env')
load_dotenv(Path(__file__).parent.parent / '.env')

# Bridge: when running on Streamlit Cloud, secrets live in st.secrets, not in
# a .env file. Promote them to os.environ so the existing os.getenv code path
# works unchanged. Local .env values (loaded above) take priority on the dev
# machine because we only set keys that aren't already in os.environ.
try:
    if hasattr(st, 'secrets'):
        for _k, _v in dict(st.secrets).items():
            if isinstance(_v, (str, int, float, bool)) and _k not in os.environ:
                os.environ[_k] = str(_v)
except Exception:
    pass

LOGO_PATH = Path(__file__).parent / 'logo' / 'real_logo.png'           # full lockup
FAVICON_PATH = Path(__file__).parent / 'logo' / 'real_1_엠블럼.png'    # square emblem for browser tab


def _b64(path: Path) -> str:
    if not path.exists():
        return ''
    try:
        return base64.b64encode(path.read_bytes()).decode('ascii')
    except Exception:
        return ''


_LOGO_B64 = _b64(LOGO_PATH)
_LOGO_URI = f'data:image/png;base64,{_LOGO_B64}' if _LOGO_B64 else ''


# Favicon needs a square-ish source so the tab icon is recognizable.
_FAVICON = str(FAVICON_PATH) if FAVICON_PATH.exists() else (
    str(LOGO_PATH) if LOGO_PATH.exists() else None
)


st.set_page_config(
    page_title='Personal RAG',
    page_icon=_FAVICON,
    layout='wide',
    initial_sidebar_state='expanded',
)

# CSS — minimal custom styling. We deliberately do NOT force display/width on
# Streamlit's own layout elements (stSidebar, stHeader, etc.) so the framework
# can manage its own responsive behavior, the sidebar can be collapsed and
# reopened normally, and future Streamlit DOM changes don't break our layout.
st.markdown(
    """
    <style>
      /* App brand inside sidebar. */
      .sb-brand { font-size: 18px; font-weight: 700; letter-spacing: -0.01em; margin-bottom: 2px; }
      .sb-tagline { color: rgba(128,128,128,0.95); font-size: 12px; margin-bottom: 4px; }
      .sb-section { color: rgba(128,128,128,0.95); font-size: 11px; text-transform: uppercase; letter-spacing: 0.06em; margin-top: 16px; margin-bottom: 6px; }

      /* Section labels inside main views. */
      .section-title { font-size: 18px; font-weight: 600; margin: 6px 0 6px 0; }
      .section-sub   { color: rgba(128,128,128,0.95); font-size: 13px; margin-bottom: 10px; line-height: 1.5; }

      /* Empty-state cards and chat hero. */
      .empty-state { padding: 28px; border: 1px dashed rgba(128,128,128,0.28); border-radius: 12px; text-align: center; color: rgba(128,128,128,0.95); line-height: 1.6; }
      .empty-hero { text-align: center; padding: 64px 16px 24px 16px; }
      .empty-hero h2 { font-size: 26px; font-weight: 600; margin: 0 0 8px 0; letter-spacing: -0.01em; }
      .empty-hero p { color: rgba(128,128,128,0.95); font-size: 14px; margin: 0 0 24px 0; }

      /* Status chips used inside sidebar bottom status. */
      .chip { display: inline-block; padding: 3px 9px; background: rgba(120,120,120,0.12); border: 1px solid rgba(120,120,120,0.22); border-radius: 999px; font-size: 11px; line-height: 1.4; margin-right: 4px; margin-bottom: 4px; }
      .chip.active { background: rgba(46,160,67,0.12); border-color: rgba(46,160,67,0.30); }
      .chip.muted  { color: rgba(128,128,128,0.95); }

      /* Make hr/divider lines softer. */
      hr { border-color: rgba(128,128,128,0.16) !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

DATA_DIR = Path(__file__).parent / '.data'
DATA_DIR.mkdir(parents=True, exist_ok=True)
LOGS_DIR = Path(__file__).parent / 'logs'
LOGS_DIR.mkdir(parents=True, exist_ok=True)


# ---------- Multi-user isolation ----------
# Each user gets their own subtree under DATA_DIR and LOGS_DIR so uploaded
# documents, indexes, conversation history, and agent runs do not leak between
# users sharing this deployment. If no users are configured in secrets, the
# app falls back to a single '_local' user (suitable for local dev).

USERS_FROM_SECRETS = {}
try:
    USERS_FROM_SECRETS = dict(st.secrets.get('users', {}) or {})
except Exception:
    USERS_FROM_SECRETS = {}


def _safe_uid(uid: str) -> str:
    return re.sub(r'[^A-Za-z0-9._-]+', '_', uid or '_local') or '_local'


def _user_data_dir() -> Path:
    d = DATA_DIR / _safe_uid(st.session_state.get('user_id', '_local'))
    d.mkdir(parents=True, exist_ok=True)
    return d


def _user_sessions_dir() -> Path:
    d = _user_data_dir() / 'sessions'
    d.mkdir(parents=True, exist_ok=True)
    return d


def _user_logs_dir() -> Path:
    d = LOGS_DIR / _safe_uid(st.session_state.get('user_id', '_local'))
    d.mkdir(parents=True, exist_ok=True)
    return d


def _agent_log_path() -> Path:
    return _user_logs_dir() / 'agents.jsonl'


def _events_log_path() -> Path:
    return _user_logs_dir() / 'events.jsonl'


# ---------- Persistent Supabase logging ----------
# Streamlit Community Cloud's filesystem is ephemeral — every redeploy or
# container restart wipes /mount/src/.../logs. To preserve user activity for
# real analysis and audit, every log function ALSO inserts into Supabase
# Postgres when SUPABASE_URL + SUPABASE_KEY are configured (via .env or
# Streamlit secrets). If neither is set, only the local JSONL is written and
# the function is a no-op. Tables expected (one-time setup via SQL editor):
#   chat_turns, agent_runs, events  — schema in db_schema.sql.

@st.cache_resource(show_spinner=False)
def _supabase_client():
    """Return a Supabase client if configured, else None. Cached so we don't
    re-import / re-connect on every turn."""
    url = os.getenv('SUPABASE_URL', '').strip()
    key = os.getenv('SUPABASE_KEY', '').strip() or os.getenv('SUPABASE_SERVICE_KEY', '').strip()
    if not url or not key:
        return None
    try:
        from supabase import create_client
        return create_client(url, key)
    except Exception:
        return None


def _scrub_for_postgres(obj):
    """Recursively strip control characters Postgres can't store.

    PDF text extraction routinely embeds NULL bytes (\\x00) into chunk text.
    Postgres rejects them in TEXT and JSONB columns with code 22P05. We
    also drop other ASCII control chars (except tab/newline/cr) since they
    contribute nothing for log analysis and trip JSON validators in some
    Postgres setups."""
    if isinstance(obj, str):
        if '\x00' not in obj and not any(
            ord(c) < 32 and c not in '\t\n\r' for c in obj[:1024]
        ):
            return obj  # fast path — clean string
        return ''.join(
            c for c in obj if ord(c) >= 32 or c in '\t\n\r'
        ).replace('\x00', '')
    if isinstance(obj, dict):
        return {k: _scrub_for_postgres(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_scrub_for_postgres(x) for x in obj]
    return obj


# ---------- Supabase-backed user auth (signup/login) ----------
# The actual passwords never reach Python. We call two SECURITY DEFINER
# RPCs in Postgres that handle bcrypt hashing/verification, and only see
# (success, message, user_id) back. Configure via db_schema_users.sql.

def _supabase_signup(username: str, password: str) -> tuple:
    """Returns (success: bool, message: str). Requires Supabase configured
    and db_schema_users.sql applied."""
    client = _supabase_client()
    if client is None:
        return False, 'Supabase 가 설정되지 않아 회원가입을 사용할 수 없습니다.'
    try:
        resp = client.rpc(
            'signup_user',
            {'p_username': username, 'p_password': password},
        ).execute()
        rows = resp.data or []
        if not rows:
            return False, '회원가입 처리 중 알 수 없는 오류가 발생했습니다.'
        row = rows[0]
        return bool(row.get('success')), str(row.get('message') or '')
    except Exception as e:
        msg = str(e)
        if 'function' in msg.lower() and 'does not exist' in msg.lower():
            return False, (
                '회원가입 시스템이 아직 설정되지 않았습니다. '
                'Supabase SQL Editor 에서 db_schema_users.sql 을 실행하세요.'
            )
        return False, f'회원가입 실패: {msg[:200]}'


def _supabase_login(username: str, password: str) -> tuple:
    """Returns (success: bool, message: str, user_id_in_db: int|None)."""
    client = _supabase_client()
    if client is None:
        return False, 'Supabase 미설정.', None
    try:
        resp = client.rpc(
            'login_user',
            {'p_username': username, 'p_password': password},
        ).execute()
        rows = resp.data or []
        if not rows:
            return False, '로그인 처리 중 알 수 없는 오류.', None
        row = rows[0]
        return (
            bool(row.get('success')),
            str(row.get('message') or ''),
            row.get('user_id'),
        )
    except Exception as e:
        msg = str(e)
        if 'function' in msg.lower() and 'does not exist' in msg.lower():
            return False, (
                '로그인 시스템이 아직 설정되지 않았습니다. '
                'Supabase SQL Editor 에서 db_schema_users.sql 을 실행하세요.'
            ), None
        return False, f'로그인 실패: {msg[:200]}', None


def _supabase_users_enabled() -> bool:
    """True if Supabase is configured. We assume db_schema_users.sql has
    been applied; the RPC helpers handle the "function missing" case with
    a friendly message instead of crashing."""
    return _supabase_client() is not None


def _should_sync_prefs_to_supabase() -> bool:
    """Sync prefs to Supabase only when a user has a stable identity. Anon
    UUIDs change each visit so syncing them wastes DB writes for no
    benefit; logged-in usernames and the local-dev `_local` slot are fine."""
    client = _supabase_client()
    if client is None:
        return False
    uid = st.session_state.get('user_id', '')
    if not uid or uid.startswith('_anon_'):
        return False
    return True


def _supabase_load_prefs() -> dict:
    """Return this user's prefs blob from Supabase, or {} if none / error.
    Designed to extend (not replace) whatever the disk file already gave us:
    callers should merge with disk dict, Supabase wins on conflicts."""
    if not _should_sync_prefs_to_supabase():
        return {}
    client = _supabase_client()
    uid = st.session_state['user_id']
    try:
        resp = client.rpc('get_prefs', {'p_user_id': uid}).execute()
        # rpc on a scalar-returning function gives us the value directly
        data = resp.data
        if isinstance(data, dict):
            return data
        if isinstance(data, str) and data:
            try:
                return json.loads(data)
            except Exception:
                return {}
        return {}
    except Exception:
        return {}


def _supabase_save_prefs(prefs: dict) -> None:
    """Upsert prefs blob. Best-effort; failures don't break the local save."""
    if not _should_sync_prefs_to_supabase():
        return
    client = _supabase_client()
    uid = st.session_state['user_id']
    try:
        client.rpc(
            'set_prefs',
            {'p_user_id': uid, 'p_prefs': _scrub_for_postgres(prefs)},
        ).execute()
    except Exception as e:
        # Don't loop noise — record once per session.
        if not st.session_state.get('_prefs_sync_warned'):
            st.session_state['_prefs_sync_warned'] = True
            st.session_state['_prefs_sync_last_err'] = (
                f'{type(e).__name__}: {str(e)[:300]}'
            )


def _supabase_insert(table: str, record: dict) -> None:
    """Best-effort INSERT. Never raises — local JSONL remains the source of
    truth for the live container; Supabase is the durable copy. Failures are
    tracked in session_state so the cache tab can surface them instead of
    leaving the user wondering why rows aren't appearing."""
    client = _supabase_client()
    if client is None:
        return
    st.session_state['_sb_attempts'] = st.session_state.get('_sb_attempts', 0) + 1
    try:
        client.table(table).insert(_scrub_for_postgres(record)).execute()
        st.session_state['_sb_successes'] = st.session_state.get('_sb_successes', 0) + 1
    except Exception as e:
        st.session_state['_sb_failures'] = st.session_state.get('_sb_failures', 0) + 1
        st.session_state['_sb_last_err'] = f'{table}: {type(e).__name__}: {str(e)[:600]}'


def _log_event(event_type: str, payload: dict = None) -> None:
    """Append one event to ./logs/{user_id}/events.jsonl.

    Captures everything outside the chat/agent JSONL: login, logout, document
    ingest, document delete, session delete, LLM call failures. One line per
    event so it merges cleanly with the other JSONL files for analytics.
    Best-effort — failures swallowed so logging never breaks the user flow.
    """
    record = {
        'kind': 'event',
        'event_type': event_type,
        'timestamp': datetime.datetime.now().isoformat(timespec='microseconds'),
        'user_id': st.session_state.get('user_id', '_local'),
        'payload': payload or {},
    }
    try:
        with _events_log_path().open('a', encoding='utf-8') as f:
            f.write(json.dumps(record, ensure_ascii=False, default=str) + '\n')
    except Exception:
        pass
    _supabase_insert('events', record)

EMBEDDER_CHOICES = [
    'BAAI/bge-m3',
    'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2',
]
RERANKER_MODEL = 'BAAI/bge-reranker-v2-m3'

# Provider presets for OpenAI-compatible endpoints.
PROVIDERS = {
    'Hugging Face Router': {
        'base_url': 'https://router.huggingface.co/v1',
        'env_key': 'HF_TOKEN',
        'default_model': 'google/gemma-4-31B-it',
    },
    'OpenAI': {
        'base_url': 'https://api.openai.com/v1',
        'env_key': 'OPENAI_API_KEY',
        'default_model': 'gpt-5-mini',
    },
    'Anthropic (Claude)': {
        # Anthropic's OpenAI-SDK compatibility endpoint. Most chat-completions
        # features work; tool calling and some sampling params have caveats.
        'base_url': 'https://api.anthropic.com/v1/',
        'env_key': 'ANTHROPIC_API_KEY',
        'default_model': 'claude-sonnet-4-6',
    },
    'Fireworks AI': {
        # OpenAI-compatible. Model ids follow accounts/fireworks/models/<name>.
        'base_url': 'https://api.fireworks.ai/inference/v1',
        'env_key': 'FIREWORKS_API_KEY',
        'default_model': 'accounts/fireworks/models/llama-v3p3-70b-instruct',
    },
    'DashScope (Qwen)': {
        'base_url': 'https://dashscope-intl.aliyuncs.com/compatible-mode/v1',
        'env_key': 'DASHSCOPE_API_KEY',
        'default_model': 'qwen3.6-27b',
    },
    'vLLM / local': {
        'base_url': 'http://localhost:8000/v1',
        'env_key': None,
        'default_model': '',
    },
    'Custom': {
        'base_url': '',
        'env_key': None,
        'default_model': '',
    },
}
PROVIDER_NAMES = list(PROVIDERS.keys())

# Popular model IDs per provider. "직접 입력" sentinel lets the user type any
# other model ID; the actual model string sent to the API is always the raw ID.
_CUSTOM = '__custom__'
PROVIDER_MODELS = {
    'Hugging Face Router': [
        # NOTE: Korean-native models from LG (EXAONE) and Naver (HyperCLOVAX)
        # exist on HF Hub but Inference Providers (Together/Cerebras/Hyperbolic)
        # don't typically deploy them. Calling them through the Router returns
        # "model_not_supported". Self-host via vLLM or Ollama instead — see the
        # 'vLLM / local' provider preset below for suggested model IDs.
        # DeepSeek — flagship general purpose
        'deepseek-ai/DeepSeek-V4-Pro',
        'deepseek-ai/DeepSeek-V4-Flash',
        'deepseek-ai/DeepSeek-V3.2',
        'deepseek-ai/DeepSeek-R1',
        # Qwen3 — strong multilingual incl. Korean
        'Qwen/Qwen3.6-35B-A3B',
        'Qwen/Qwen3-Next-80B-A3B-Instruct',
        'Qwen/Qwen3-235B-A22B-Instruct-2507',
        'Qwen/Qwen3-235B-A22B-Thinking-2507',
        'Qwen/Qwen3-VL-30B-A3B-Instruct',
        'Qwen/Qwen2.5-7B-Instruct',
        # Google Gemma 4 (multimodal)
        'google/gemma-4-31B-it',
        'google/gemma-4-26B-A4B-it',
        # Meta Llama
        'meta-llama/Llama-3.3-70B-Instruct',
        'meta-llama/Llama-3.1-8B-Instruct',
        # Moonshot Kimi K2 (long-context, agentic)
        'moonshotai/Kimi-K2.6',
        'moonshotai/Kimi-K2-Thinking',
        # GLM (Z.ai)
        'zai-org/GLM-5.1',
        'zai-org/GLM-4.7',
        # OpenAI open-source
        'openai/gpt-oss-120b',
        'openai/gpt-oss-20b',
        # MiniMax
        'MiniMaxAI/MiniMax-M2.7',
        # Cohere
        'CohereLabs/c4ai-command-a-03-2025',
    ],
    'OpenAI': [
        # GPT-5 tier — Chat Completions compatible only.
        # NOTE: *-pro variants (gpt-5-pro, gpt-5.5-pro, gpt-5.4-pro, o1-pro)
        # are OpenAI's Responses-API-only models and would 404 here, so they
        # are intentionally omitted. Use gpt-5 or o3 for top-tier reasoning.
        'gpt-5.5',
        'gpt-5.4',
        'gpt-5.4-mini',
        'gpt-5.4-nano',
        'gpt-5.3-chat-latest',
        'gpt-5.2',
        'gpt-5.1',
        'gpt-5',
        'gpt-5-mini',
        'gpt-5-nano',
        'gpt-5-chat-latest',
        # GPT-4.1 series
        'gpt-4.1',
        'gpt-4.1-mini',
        'gpt-4.1-nano',
        # GPT-4o family
        'gpt-4o',
        'gpt-4o-mini',
        # Reasoning models
        'o4-mini',
        'o3',
        'o3-mini',
        'o1',
    ],
    'Anthropic (Claude)': [
        # Latest Claude 4.x tier
        'claude-opus-4-7',
        'claude-sonnet-4-6',
        'claude-haiku-4-5',
        # 4.x earlier
        'claude-opus-4-5',
        'claude-sonnet-4-5',
        'claude-haiku-4-1',
        # 3.7 / 3.5 (legacy but stable)
        'claude-3-7-sonnet-latest',
        'claude-3-5-sonnet-latest',
        'claude-3-5-haiku-latest',
    ],
    'Fireworks AI': [
        # NOTE: model IDs change occasionally on Fireworks; some require
        # account tier upgrades (e.g. 405B). If you get 404 NOT_FOUND check
        # the current list at https://fireworks.ai/models.
        # Llama
        'accounts/fireworks/models/llama-v3p3-70b-instruct',
        'accounts/fireworks/models/llama-v3p1-70b-instruct',
        'accounts/fireworks/models/llama-v3p1-8b-instruct',
        # Qwen
        'accounts/fireworks/models/qwen2p5-72b-instruct',
        'accounts/fireworks/models/qwen2p5-coder-32b-instruct',
        # DeepSeek
        'accounts/fireworks/models/deepseek-v3',
        'accounts/fireworks/models/deepseek-r1',
        # Mixtral
        'accounts/fireworks/models/mixtral-8x22b-instruct',
    ],
    'DashScope (Qwen)': [
        'qwen3.6-27b',
        'qwen3-72b',
        'qwen-max',
        'qwen-plus',
        'qwen-turbo',
        'qwen-vl-max',
    ],
    'vLLM / local': [
        # Self-hosted suggestions (you bring the GPU + run `vllm serve <id>`).
        # Strong open multilingual that handle Korean well:
        'Qwen/Qwen3-Next-80B-A3B-Instruct',
        'Qwen/Qwen2.5-7B-Instruct',
        'meta-llama/Llama-3.3-70B-Instruct',
        'meta-llama/Llama-3.1-8B-Instruct',
    ],
    'Custom': [],
}


# =============================================================================
# Session state
# =============================================================================

def _init_state():
    defaults = {
        # chat
        'user_inputs': [],
        'generated_responses': [],
        'thinking_traces': [],
        'retrieved_per_turn': [],

        # model / endpoint
        'provider': 'Hugging Face Router',
        'model': 'google/gemma-4-31B-it',
        'base_url': os.getenv('OPENAI_BASE_URL', 'https://router.huggingface.co/v1'),

        # Per-provider API keys. Stored separately so switching providers
        # does not overwrite a manually entered key.
        'hf_api_key': os.getenv('HF_TOKEN', ''),
        'openai_api_key': os.getenv('OPENAI_API_KEY', ''),
        'anthropic_api_key': os.getenv('ANTHROPIC_API_KEY', ''),
        'fireworks_api_key': os.getenv('FIREWORKS_API_KEY', ''),
        'dashscope_api_key': os.getenv('DASHSCOPE_API_KEY', ''),
        'custom_api_key': '',

        # generation
        'max_tokens': 8192,
        'temperature': 1.0,
        'top_p': 0.95,
        'sampling_top_k': 64,
        'presence_penalty': 0.0,
        'enable_thinking': True,
        'stream': True,

        # rag config — on Streamlit Cloud free tier RAM is ~1 GB, which is
        # below BGE-M3 (~2.2 GB) + reranker (~580 MB). Default to the lighter
        # MiniLM + reranker off there so the app doesn't OOM on first PDF.
        'embedder_model': (
            'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2'
            if Path('/mount/src').exists() else EMBEDDER_CHOICES[0]
        ),
        'chunk_size': 600,
        'chunk_overlap': 80,
        'retrieval_mode': 'hybrid',   # 'dense' | 'bm25' | 'hybrid'
        'retrieve_top_n': 20,
        'final_top_k': 5,
        'use_reranker': not Path('/mount/src').exists(),
        # Route dense retrieval through Supabase pgvector when both this flag
        # is ON and a Supabase client is configured. Default OFF — user
        # opts in after running db_schema_pgvector.sql and seeing the cache
        # tab confirm that chunks are being upserted.
        'use_pgvector_search': False,
        # Agentic loop: after the initial retrieval, the LLM may issue
        # follow-up searches via a tool call. Default OFF; toggle in
        # settings tab. Only active when the model/provider supports
        # OpenAI-style tool calling (gpt-5, gpt-4.1, Qwen3, Llama-3.x, ...).
        'use_agentic_search': False,
        'agentic_max_iters': 3,
        # When True, skip retrieval entirely and let the LLM answer from its
        # own knowledge — useful when the user wants a quick general chat
        # without uploading documents or when documents are loaded but
        # the question is off-topic.
        'general_chat_mode': False,

        # query expansion config
        'use_hyde': False,
        'use_multi_query': False,
        'n_paraphrases': 3,
        'use_contextual_rewrite': True,

        # cross-document retrieval
        'per_doc_balance': True,
        'per_doc_reserve': 1,
        'comparison_autodetect': True,

        # multimodal config
        'include_page_images': False,
        'max_page_images': 3,

        # web search config
        'web_enabled': False,
        'web_provider': 'duckduckgo',  # 'duckduckgo' | 'tavily' | 'brave'
        'web_top_n': 5,
        'tavily_key': os.getenv('TAVILY_API_KEY', ''),
        'brave_key': os.getenv('BRAVE_API_KEY', ''),

        # per-turn parallel arrays
        'query_variants_per_turn': [],

        # navigation
        'active_view': 'chat',
        '_pending_input': None,

        # chat-time document filter (None = use all loaded docs)
        'chat_doc_filter': None,

        # session (conversation) tracking
        'current_session_id': None,
        'current_session_title': '',
        'current_session_created_at': '',

        # in-memory doc store: list of dicts {id, name, raw_text, chunks, embeddings, chunk_size, chunk_overlap}
        'docs': [],
        '_loaded_for_embedder': None,   # tracks which embedder's persisted docs are loaded
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)


_init_state()


# ---------- Login gate ----------

def _render_login_screen():
    """Render brand + login form. Sets user_id on success and reruns; otherwise
    st.stop()s so the rest of the app is gated off.

    Also renders a minimal sidebar (brand only) so the page does not look like
    the sidebar disappeared — it just has no nav until login completes.
    """
    with st.sidebar:
        if _LOGO_URI:
            st.markdown(
                f'<div style="text-align:center; padding:4px 0 4px 0;">'
                f'<img src="{_LOGO_URI}" '
                f'style="width:100%; max-width:220px; height:auto; display:block; margin:0 auto;" />'
                f'</div>',
                unsafe_allow_html=True,
            )
        st.markdown(
            '<div class="sb-brand" style="text-align:center; font-size:14px;">Personal RAG</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            '<div class="sb-tagline" style="text-align:center;">로그인 후 이용</div>',
            unsafe_allow_html=True,
        )

    if _LOGO_URI:
        st.markdown(
            f'<div style="text-align:center; margin-top:80px; margin-bottom:8px;">'
            f'<img src="{_LOGO_URI}" '
            f'style="max-width:240px; height:auto; opacity:0.95;" /></div>',
            unsafe_allow_html=True,
        )

    # Two distinct login backends:
    #  1) Supabase users table (signup + login, bcrypt). Active when Supabase
    #     is configured. New users can register themselves.
    #  2) Legacy secrets [users] block. Admin-managed; no signup.
    use_supabase_auth = _supabase_users_enabled()

    if use_supabase_auth:
        st.markdown(
            '<div style="text-align:center; font-size:14px; color:rgba(128,128,128,0.9); '
            'margin-bottom:16px;">계정으로 로그인하거나 새로 가입하세요.</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="text-align:center; font-size:14px; color:rgba(128,128,128,0.9); '
            'margin-bottom:24px;">로그인이 필요합니다.</div>',
            unsafe_allow_html=True,
        )

    _, form_col, _ = st.columns([1, 2, 1])
    with form_col:
        if use_supabase_auth:
            tab_login, tab_signup = st.tabs(['로그인', '회원가입'])
            with tab_login:
                with st.form('_login_form', clear_on_submit=False):
                    username = st.text_input(
                        '아이디', key='login_username',
                        placeholder='가입한 아이디',
                    )
                    password = st.text_input(
                        '비밀번호', type='password', key='login_password',
                        placeholder='비밀번호 입력',
                    )
                    submitted = st.form_submit_button(
                        '로그인', use_container_width=True, type='primary',
                    )
                if submitted:
                    ok, msg, _user_id = _supabase_login(username, password)
                    if ok:
                        st.session_state['user_id'] = (username or '').strip()
                        st.session_state['_loaded_for_embedder'] = None
                        _log_event('login', {
                            'method': 'supabase',
                            'username': st.session_state['user_id'],
                        })
                        st.rerun()
                    else:
                        _log_event('login_failed', {
                            'method': 'supabase',
                            'attempted_username': (username or '')[:64],
                            'reason': msg[:200],
                        })
                        st.error(msg)
            with tab_signup:
                with st.form('_signup_form', clear_on_submit=False):
                    new_username = st.text_input(
                        '새 아이디',
                        key='signup_username',
                        placeholder='2자 이상 64자 이하',
                    )
                    new_password = st.text_input(
                        '새 비밀번호', type='password',
                        key='signup_password',
                        placeholder='6자 이상 입력',
                    )
                    new_password2 = st.text_input(
                        '비밀번호 확인', type='password',
                        key='signup_password2',
                        placeholder='위와 동일하게 한 번 더',
                    )
                    signup_submitted = st.form_submit_button(
                        '회원가입', use_container_width=True, type='primary',
                    )
                if signup_submitted:
                    if new_password != new_password2:
                        st.error('비밀번호 확인이 일치하지 않습니다.')
                    else:
                        ok, msg = _supabase_signup(new_username, new_password)
                        if ok:
                            # Auto-login after signup so user doesn't have to
                            # retype credentials.
                            ok2, _msg2, _uid = _supabase_login(
                                new_username, new_password,
                            )
                            if ok2:
                                st.session_state['user_id'] = (
                                    (new_username or '').strip()
                                )
                                st.session_state['_loaded_for_embedder'] = None
                                _log_event('signup', {
                                    'username': st.session_state['user_id'],
                                })
                                _log_event('login', {
                                    'method': 'supabase',
                                    'username': st.session_state['user_id'],
                                    'first_login': True,
                                })
                                st.success('회원가입 완료 — 자동 로그인했습니다.')
                                st.rerun()
                            else:
                                st.success('회원가입 성공. 로그인 탭에서 다시 로그인해 주세요.')
                        else:
                            st.error(msg)
        else:
            with st.form('_login_form', clear_on_submit=False):
                username = st.text_input('사용자 ID')
                password = st.text_input('비밀번호', type='password')
                submitted = st.form_submit_button(
                    '로그인', use_container_width=True, type='primary',
                )
            if submitted:
                if (username in USERS_FROM_SECRETS
                        and str(USERS_FROM_SECRETS[username]) == str(password)):
                    st.session_state['user_id'] = username
                    st.session_state['_loaded_for_embedder'] = None
                    _log_event('login', {'method': 'password', 'username': username})
                    st.rerun()
                else:
                    _log_event('login_failed', {
                        'method': 'password',
                        'attempted_username': (username or '')[:64],
                    })
                    st.error('사용자 ID 또는 비밀번호가 올바르지 않습니다.')
    st.stop()


def _is_streamlit_cloud() -> bool:
    """Detect Streamlit Community Cloud — apps are mounted under /mount/src/."""
    return Path('/mount/src').exists()


def _auth_gate():
    """Resolve the active user_id. Priority:
      1) Supabase users (signup + login, bcrypt) — when Supabase is wired up.
         This is the recommended setup for any shared deployment.
      2) Legacy secrets [users] block — admin-managed, no signup.
      3) Streamlit Cloud, no auth backend at all → anonymous per-browser UUID.
      4) Local dev, no auth backend → single-tenant '_local'.
    """
    current_uid = st.session_state.get('user_id', '')

    # Migration: if an auth backend was just enabled but the browser still has
    # a stale anonymous user_id from before the backend existed, drop it so
    # the login screen actually appears. Otherwise the gate keeps short-
    # circuiting on the old anon id and the user never sees a login form.
    if current_uid.startswith('_anon_') and (
            _supabase_users_enabled() or USERS_FROM_SECRETS):
        st.session_state['user_id'] = ''
        current_uid = ''

    if current_uid:
        return  # already logged in (real account or anon/local fallback)

    if _supabase_users_enabled():
        _render_login_screen()
        return

    if USERS_FROM_SECRETS:
        _render_login_screen()
        return

    if _is_streamlit_cloud():
        # New visitor on a shared deployment without auth → give them their
        # own isolated workspace for this browser session. Closing the tab
        # ends the session; previous anonymous data remains on disk under
        # its UUID until an admin cleans it up.
        st.session_state['user_id'] = '_anon_' + uuid.uuid4().hex[:10]
        _log_event('login', {'method': 'anonymous'})
    else:
        st.session_state['user_id'] = '_local'
        if not st.session_state.get('_local_login_logged'):
            _log_event('login', {'method': 'local'})
            st.session_state['_local_login_logged'] = True


_auth_gate()


# ---------- Persisted per-user preferences (survives idle reset) ----------
# Streamlit Cloud disconnects idle WebSocket sessions, and reconnection
# resets st.session_state. We persist the slowly-changing config (API keys,
# provider, model, retrieval/sampling settings, active view) to disk so the
# user does not lose them when they come back after stepping away. Conversation
# history is already persisted via the session JSONL files.

_PERSIST_KEYS = (
    'provider', 'model', 'base_url',
    'hf_api_key', 'openai_api_key', 'anthropic_api_key', 'fireworks_api_key',
    'dashscope_api_key', 'custom_api_key',
    'tavily_key', 'brave_key',
    'embedder_model', 'chunk_size', 'chunk_overlap',
    'retrieval_mode', 'retrieve_top_n', 'final_top_k', 'use_reranker',
    'use_pgvector_search', 'use_agentic_search', 'agentic_max_iters',
    'general_chat_mode',
    'use_contextual_rewrite', 'use_multi_query', 'use_hyde', 'n_paraphrases',
    'per_doc_balance', 'per_doc_reserve', 'comparison_autodetect',
    'include_page_images', 'max_page_images',
    'web_enabled', 'web_provider', 'web_top_n',
    'max_tokens', 'temperature', 'top_p', 'sampling_top_k',
    'presence_penalty', 'stream', 'enable_thinking',
    'chat_doc_filter', 'active_view',
    # Active conversation pointer — combined with auto-restore in view_chat
    # this lets the user pick up exactly where they were after a Cloud
    # idle reset, not just "new chat on the chat tab".
    'current_session_id',
)


def _user_prefs_path() -> Path:
    return _user_data_dir() / 'preferences.json'


def _load_user_prefs():
    """Override session_state defaults with whatever was last saved for this
    user. Tries disk first (fast, available locally), then merges in
    Supabase prefs (survives Cloud container restarts and follows the user
    across devices). Supabase values win on conflict — they are the
    authoritative source for logged-in users.

    Crucial: this must run only ONCE per Streamlit session. Otherwise every
    rerun (e.g. when a sidebar button sets active_view='settings') would
    immediately overwrite the just-clicked value with the previously-saved
    disk value, making the UI feel unresponsive.

    After an idle reconnect, session_state is wiped → the _prefs_loaded flag
    is gone → this function runs again and restores the user's last state.
    """
    if st.session_state.get('_prefs_loaded'):
        return

    data = {}

    # Disk first — present on warm Cloud container or local dev.
    p = _user_prefs_path()
    if p.exists():
        try:
            data = json.loads(p.read_text())
        except Exception:
            data = {}

    # Supabase next — wins on conflict so cross-device sync works.
    sb_prefs = _supabase_load_prefs()
    if isinstance(sb_prefs, dict) and sb_prefs:
        data.update(sb_prefs)

    for k in _PERSIST_KEYS:
        if k not in data:
            continue
        v = data[k]
        if v is None:
            continue
        if isinstance(v, str) and not v:
            # Don't clobber an env-loaded key with an empty saved string.
            continue
        st.session_state[k] = v
    # Fallback: deprecated models (Responses-only / Korean-native not on
    # HF Inference Providers) should be swapped before any API call. We
    # only have the swap table after PROVIDERS is defined further down the
    # file, so just check by name here — the table itself is referenced at
    # call time by _resolve_deprecated_model(). This pass handles "user's
    # picker still shows the bad name" on session start.
    _DEPRECATED_NOW = {
        'gpt-5-pro', 'o1-pro', 'gpt-5.5-pro', 'gpt-5.4-pro',
        'LGAI-EXAONE/EXAONE-4.5-33B',
        'naver-hyperclovax/HyperCLOVAX-SEED-Think-32B',
    }
    if st.session_state.get('model') in _DEPRECATED_NOW:
        # If we know the provider, swap to its default; otherwise leave a
        # neutral safe choice. The call-time guard handles the rest.
        prov = st.session_state.get('provider', '')
        if prov == 'OpenAI':
            st.session_state['model'] = 'gpt-5-mini'
        elif prov == 'Hugging Face Router':
            st.session_state['model'] = 'Qwen/Qwen3-Next-80B-A3B-Instruct'
        else:
            st.session_state['model'] = 'Qwen/Qwen3-Next-80B-A3B-Instruct'
    st.session_state['_prefs_loaded'] = True


def _save_user_prefs():
    """Persist current configurable session_state. Writes to local disk
    (fast, works offline) AND to Supabase user_preferences (survives
    container restarts, follows the user across devices). Each side is
    best-effort; if one fails the other still happened."""
    data = {k: st.session_state.get(k) for k in _PERSIST_KEYS}
    try:
        encoded = json.dumps(data, ensure_ascii=False, default=str, sort_keys=True)
    except Exception:
        return
    if st.session_state.get('_prefs_snapshot') == encoded:
        return
    try:
        _user_prefs_path().write_text(encoded)
    except Exception:
        pass
    _supabase_save_prefs(data)
    st.session_state['_prefs_snapshot'] = encoded


_load_user_prefs()


# =============================================================================
# Persistent vector store
# =============================================================================

def _safe_name(s: str) -> str:
    return re.sub(r'[^A-Za-z0-9._-]+', '_', s)


def _embedder_dir(embedder_id: str) -> Path:
    return _user_data_dir() / _safe_name(embedder_id)


def compute_doc_id(name: str, raw_text: str, chunk_size: int, chunk_overlap: int) -> str:
    h = hashlib.sha256()
    h.update(name.encode('utf-8'))
    h.update(b'\0')
    h.update(raw_text.encode('utf-8', errors='ignore'))
    h.update(f'\0{chunk_size}\0{chunk_overlap}'.encode())
    return h.hexdigest()[:16]


def save_doc(embedder_id: str, doc: dict) -> None:
    d = _embedder_dir(embedder_id) / doc['id']
    d.mkdir(parents=True, exist_ok=True)
    meta = {
        'id': doc['id'],
        'name': doc['name'],
        'chunks': doc['chunks'],
        'chunk_pages': doc.get('chunk_pages', []),
        'page_count': doc.get('page_count', 0),
        'has_page_images': bool(doc.get('has_page_images', False)),
        'is_pdf': bool(doc.get('is_pdf', False)),
        'raw_text': doc['raw_text'],
        'chunk_size': doc['chunk_size'],
        'chunk_overlap': doc['chunk_overlap'],
    }
    (d / 'meta.json').write_text(json.dumps(meta, ensure_ascii=False))
    np.save(d / 'embeddings.npy', doc['embeddings'])


# ---------- pgvector dual-write (optional) ----------
# Maps the embedder model id we use locally to the short name we store in
# doc_chunks.embedder and the vector column that holds the actual values.
_EMBEDDER_TABLE_MAP = {
    'BAAI/bge-m3': ('bge-m3', 'embedding_bgem3', 1024),
    'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2':
        ('paraphrase-multilingual-MiniLM-L12-v2', 'embedding_minilm', 384),
}


def _pgvector_upsert_doc(embedder_id: str, doc: dict) -> None:
    """Best-effort: send every chunk of `doc` into public.doc_chunks so the
    embeddings survive container restarts and can be queried at the DB layer
    when pgvector retrieval is enabled. No-op when Supabase isn't wired up
    or when the embedder isn't one we have a column for.

    Tracks success/failure counters in session_state (separate from the
    generic logging counters) so the cache tab can surface state. Never
    raises — local on-disk numpy arrays remain the source of truth."""
    client = _supabase_client()
    if client is None:
        return
    mapping = _EMBEDDER_TABLE_MAP.get(embedder_id)
    if mapping is None:
        return
    short_name, vec_col, _dim = mapping

    user_id = st.session_state.get('user_id', '_local')
    chunks = doc.get('chunks') or []
    embs = doc.get('embeddings')
    pages = doc.get('chunk_pages') or [[] for _ in chunks]
    if embs is None or len(embs) != len(chunks):
        return

    rows = []
    for i, ch in enumerate(chunks):
        v = embs[i]
        # supabase-py serializes lists as pgvector literals automatically;
        # numpy.ndarray needs .tolist() first.
        if hasattr(v, 'tolist'):
            v = v.tolist()
        rows.append({
            'user_id': user_id,
            'doc_id': doc['id'],
            'doc_name': doc.get('name', ''),
            'chunk_idx': i,
            'text': ch,
            'pages': pages[i] if i < len(pages) else [],
            'embedder': short_name,
            vec_col: v,
        })
    if not rows:
        return

    # Batch the upsert to keep individual requests under a few hundred KB
    # — large PDFs can have hundreds of chunks.
    BATCH = 100
    for start in range(0, len(rows), BATCH):
        batch = rows[start:start + BATCH]
        st.session_state['_pgv_attempts'] = (
            st.session_state.get('_pgv_attempts', 0) + len(batch)
        )
        try:
            (client.table('doc_chunks')
                .upsert(
                    _scrub_for_postgres(batch),
                    on_conflict='user_id,doc_id,chunk_idx,embedder',
                )
                .execute())
            st.session_state['_pgv_successes'] = (
                st.session_state.get('_pgv_successes', 0) + len(batch)
            )
        except Exception as e:
            st.session_state['_pgv_failures'] = (
                st.session_state.get('_pgv_failures', 0) + len(batch)
            )
            st.session_state['_pgv_last_err'] = (
                f'doc_chunks batch: {type(e).__name__}: {str(e)[:600]}'
            )
            # If the very first batch fails (e.g. table missing), don't
            # keep hammering — let the rest abort silently.
            return


def _pgvector_delete_doc(doc_id: str) -> None:
    """Best-effort cleanup: when a user deletes a document locally, also
    remove its chunks from pgvector so the two stores stay aligned."""
    client = _supabase_client()
    if client is None:
        return
    user_id = st.session_state.get('user_id', '_local')
    try:
        (client.table('doc_chunks')
            .delete()
            .eq('user_id', user_id)
            .eq('doc_id', doc_id)
            .execute())
    except Exception:
        pass


def load_doc(embedder_id: str, doc_id: str):
    d = _embedder_dir(embedder_id) / doc_id
    mp = d / 'meta.json'
    ep = d / 'embeddings.npy'
    if not (mp.exists() and ep.exists()):
        return None
    try:
        meta = json.loads(mp.read_text())
        embs = np.load(ep)
    except Exception:
        return None
    return {
        'id': doc_id,
        'name': meta['name'],
        'raw_text': meta.get('raw_text', ''),
        'chunks': meta['chunks'],
        'chunk_pages': meta.get('chunk_pages', []),
        'page_count': meta.get('page_count', 0),
        'has_page_images': meta.get('has_page_images', False),
        'is_pdf': meta.get('is_pdf', False),
        'chunk_size': meta.get('chunk_size'),
        'chunk_overlap': meta.get('chunk_overlap'),
        'embeddings': embs,
    }


def _pages_dir(embedder_id: str, doc_id: str) -> Path:
    return _embedder_dir(embedder_id) / doc_id / 'pages'


# ---------- Conversation sessions ----------

def _new_session_id() -> str:
    return datetime.datetime.now().strftime('%Y%m%d-%H%M%S-') + uuid.uuid4().hex[:6]


def _session_path(sid: str) -> Path:
    return _user_sessions_dir() / f'{sid}.json'


def save_current_session():
    """Persist current chat (if there is one). No-op if no user turns yet."""
    sid = st.session_state.get('current_session_id')
    if not sid:
        return
    if not st.session_state['user_inputs']:
        return
    data = {
        'id': sid,
        'title': st.session_state.get('current_session_title', '') or '',
        'created_at': st.session_state.get('current_session_created_at')
            or datetime.datetime.now().isoformat(),
        'updated_at': datetime.datetime.now().isoformat(),
        'model': st.session_state['model'],
        'provider': st.session_state.get('provider', ''),
        'user_inputs': st.session_state['user_inputs'],
        'generated_responses': st.session_state['generated_responses'],
        'thinking_traces': st.session_state['thinking_traces'],
        # retrieved chunks / variants intentionally not persisted (heavy + chunk ids may shift).
    }
    try:
        _session_path(sid).write_text(json.dumps(data, ensure_ascii=False))
    except Exception as e:
        st.warning(f'대화 저장 실패: {e}')


def _supabase_list_sessions(limit: int = 50) -> list:
    """Aggregate the current user's sessions from Supabase chat_turns.
    Returns [] when Supabase isn't configured, user is anonymous, or the
    RPC isn't available yet."""
    client = _supabase_client()
    if client is None:
        return []
    uid = st.session_state.get('user_id', '')
    if not uid or uid.startswith('_anon_'):
        return []
    try:
        resp = client.rpc(
            'list_user_sessions',
            {'p_user_id': uid, 'p_limit': int(limit)},
        ).execute()
        rows = resp.data or []
    except Exception:
        return []
    out = []
    for r in rows:
        sid = r.get('session_id')
        if not sid:
            continue
        out.append({
            'id': sid,
            'title': r.get('title') or '(제목 없음)',
            'updated_at': r.get('updated_at', '') or '',
            'model': r.get('model', '') or '',
            'n_turns': int(r.get('n_turns') or 0),
            '_source': 'supabase',
        })
    return out


def _supabase_load_session(sid: str) -> bool:
    """Pull every chat_turn row for (user, session) and populate session_state
    so the conversation re-renders. Returns False when no rows / errors."""
    client = _supabase_client()
    if client is None:
        return False
    uid = st.session_state.get('user_id', '')
    if not uid or uid.startswith('_anon_'):
        return False
    try:
        resp = (client.table('chat_turns')
                .select('turn_index, "timestamp", session_title, model, '
                        'user_message, assistant_message, reasoning, '
                        'query_variants')
                .eq('user_id', uid)
                .eq('session_id', sid)
                .order('turn_index')
                .execute())
        rows = resp.data or []
    except Exception:
        return False
    if not rows:
        return False

    title = rows[0].get('session_title') or '(제목 없음)'
    created = rows[0].get('timestamp', '')
    st.session_state['current_session_id'] = sid
    st.session_state['current_session_title'] = title
    st.session_state['current_session_created_at'] = created
    st.session_state['user_inputs'] = [r.get('user_message', '') or '' for r in rows]
    st.session_state['generated_responses'] = [
        r.get('assistant_message', '') or '' for r in rows
    ]
    st.session_state['thinking_traces'] = [
        r.get('reasoning', '') or '' for r in rows
    ]
    n = len(rows)
    st.session_state['retrieved_per_turn'] = [[] for _ in range(n)]
    st.session_state['query_variants_per_turn'] = [
        (r.get('query_variants') or []) for r in rows
    ]
    # Save back to disk as a warm cache for next access.
    try:
        save_current_session()
    except Exception:
        pass
    return True


def load_session(sid: str) -> bool:
    """Try local disk first (full retrieved cache); fall back to rebuilding
    the conversation from Supabase chat_turns when the disk file isn't on
    this container (fresh Cloud start, different device, etc.)."""
    p = _session_path(sid)
    if p.exists():
        try:
            data = json.loads(p.read_text())
        except Exception:
            data = None
        if data:
            st.session_state['current_session_id'] = data.get('id', sid)
            st.session_state['current_session_title'] = data.get('title', '')
            st.session_state['current_session_created_at'] = data.get('created_at', '')
            st.session_state['user_inputs'] = data.get('user_inputs', [])
            st.session_state['generated_responses'] = data.get('generated_responses', [])
            st.session_state['thinking_traces'] = data.get('thinking_traces', [])
            n = len(st.session_state['user_inputs'])
            st.session_state['retrieved_per_turn'] = [[] for _ in range(n)]
            st.session_state['query_variants_per_turn'] = [[] for _ in range(n)]
            return True
    # Disk miss → Supabase
    return _supabase_load_session(sid)


def list_sessions(limit: int = 30):
    """Sessions sorted by updated_at desc. Merges local disk with Supabase
    (logged-in users only); Supabase contributes sessions that aren't on the
    current container's disk. Disk entries win on duplicate id since they
    carry the original session metadata (created_at, etc.)."""
    sd = _user_sessions_dir()
    out = []
    seen_ids = set()
    if sd.exists():
        for p in sd.glob('*.json'):
            try:
                data = json.loads(p.read_text())
            except Exception:
                continue
            sid = data.get('id', p.stem)
            seen_ids.add(sid)
            out.append({
                'id': sid,
                'title': data.get('title') or '(제목 없음)',
                'updated_at': data.get('updated_at', ''),
                'model': data.get('model', ''),
            })
    for s in _supabase_list_sessions(limit=max(limit * 2, 100)):
        if s['id'] in seen_ids:
            continue
        out.append({
            'id': s['id'],
            'title': s.get('title') or '(제목 없음)',
            'updated_at': s.get('updated_at', ''),
            'model': s.get('model', ''),
        })
    out.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
    return out[:limit]


def rename_session(sid: str, new_title: str) -> bool:
    """Rename a saved conversation in both disk metadata and Supabase.

    Updates:
      - .data/{user}/sessions/{sid}.json — title field
      - public.chat_turns.session_title — every row of this session, so the
        sidebar list rebuilt from Supabase also reflects the new name
      - st.session_state['current_session_title'] if this is the active one
    Returns True when at least one storage layer updated successfully.
    """
    new_title = (new_title or '').strip()[:60]
    if not new_title:
        return False
    ok_any = False

    # Disk metadata
    p = _session_path(sid)
    if p.exists():
        try:
            data = json.loads(p.read_text())
            data['title'] = new_title
            data['updated_at'] = datetime.datetime.now().isoformat()
            p.write_text(json.dumps(data, ensure_ascii=False))
            ok_any = True
        except Exception:
            pass

    # Supabase chat_turns — update every row of this session for the user
    client = _supabase_client()
    uid = st.session_state.get('user_id', '')
    if client is not None and uid and not uid.startswith('_anon_'):
        try:
            (client.table('chat_turns')
                .update({'session_title': new_title})
                .eq('user_id', uid)
                .eq('session_id', sid)
                .execute())
            ok_any = True
        except Exception:
            pass

    # Active session — mirror the change in session_state too.
    if st.session_state.get('current_session_id') == sid:
        st.session_state['current_session_title'] = new_title

    if ok_any:
        _log_event('session_rename', {'session_id': sid, 'new_title': new_title})
    return ok_any


def delete_session(sid: str):
    """Remove a conversation entirely — disk file + every chat_turn row in
    Supabase for that (user, session). Audit logged."""
    p = _session_path(sid)
    existed = p.exists()
    if existed:
        try:
            p.unlink()
        except Exception:
            pass
    # Also wipe the corresponding chat_turns rows so the session does not
    # reappear after the next Cloud restart / device switch.
    client = _supabase_client()
    uid = st.session_state.get('user_id', '')
    if client is not None and uid and not uid.startswith('_anon_'):
        try:
            (client.table('chat_turns')
                .delete()
                .eq('user_id', uid)
                .eq('session_id', sid)
                .execute())
        except Exception:
            pass
    _log_event('session_delete', {'session_id': sid, 'existed': existed})


def start_new_session():
    """Reset state for a new conversation. ID is assigned lazily on first message."""
    st.session_state['current_session_id'] = None
    st.session_state['current_session_title'] = ''
    st.session_state['current_session_created_at'] = ''
    st.session_state['user_inputs'] = []
    st.session_state['generated_responses'] = []
    st.session_state['thinking_traces'] = []
    st.session_state['retrieved_per_turn'] = []
    st.session_state['query_variants_per_turn'] = []


def _session_jsonl_path(sid: str) -> Path:
    return _user_logs_dir() / f'{sid}.jsonl'


def log_turn_structured(user_input: str, response_text: str, reasoning: str,
                        retrieved: list, model: str, elapsed_sec,
                        query_variants: list = None):
    """Append one JSON record per turn to ./logs/{session_id}.jsonl.
    Postgres-ready schema. One line per turn; analysis-friendly.

    Schema fields: session_id, session_title, turn_index, timestamp,
      user_message, assistant_message, reasoning,
      model, provider, base_url, elapsed_seconds,
      retrieved[] (rank/source/score/text/doc_id/doc_name/chunk_index/pages/url/title),
      n_retrieved, citation_numbers_used[], n_cited,
      query_variants[], settings_snapshot{} (all retrieval/sampling/feature flags).
    """
    sid = st.session_state.get('current_session_id')
    if not sid:
        return
    path = _session_jsonl_path(sid)

    # Normalize retrieved entries. We store only a 200-char preview plus a
    # reference (doc_id + chunk_index) so chat_turns rows stay light — the
    # full chunk body is in doc_chunks and can be joined back when needed.
    retrieved_records = []
    for j, r in enumerate(retrieved or [], start=1):
        rec = {
            'rank': j,
            'source': r.get('source'),
            'score': r.get('score'),
            'text_preview': (r.get('text') or '')[:200],
        }
        if r.get('source') == 'web':
            rec['title'] = r.get('doc')
            rec['url'] = r.get('url')
        else:
            rec['doc_id'] = r.get('doc_id')
            rec['doc_name'] = r.get('doc')
            rec['chunk_index'] = r.get('chunk_idx')
            rec['pages'] = r.get('pages') or []
        retrieved_records.append(rec)

    n_chunks = len(retrieved) if retrieved else 0
    cited_nums = sorted(parse_citations(response_text, n_chunks)) if n_chunks else []

    settings = {
        'embedder_model': st.session_state.get('embedder_model'),
        'retrieval_mode': st.session_state.get('retrieval_mode'),
        'retrieve_top_n': st.session_state.get('retrieve_top_n'),
        'final_top_k': st.session_state.get('final_top_k'),
        'use_reranker': st.session_state.get('use_reranker'),
        'use_pgvector_search': st.session_state.get('use_pgvector_search'),
        'use_agentic_search': st.session_state.get('use_agentic_search'),
        'general_chat_mode': st.session_state.get('general_chat_mode'),
        'agentic_max_iters': st.session_state.get('agentic_max_iters'),
        'use_hyde': st.session_state.get('use_hyde'),
        'use_multi_query': st.session_state.get('use_multi_query'),
        'n_paraphrases': st.session_state.get('n_paraphrases'),
        'use_contextual_rewrite': st.session_state.get('use_contextual_rewrite'),
        'per_doc_balance': st.session_state.get('per_doc_balance'),
        'per_doc_reserve': st.session_state.get('per_doc_reserve'),
        'comparison_autodetect': st.session_state.get('comparison_autodetect'),
        'chat_doc_filter': st.session_state.get('chat_doc_filter'),
        'web_enabled': st.session_state.get('web_enabled'),
        'web_provider': st.session_state.get('web_provider'),
        'web_top_n': st.session_state.get('web_top_n'),
        'include_page_images': st.session_state.get('include_page_images'),
        'max_page_images': st.session_state.get('max_page_images'),
        'temperature': st.session_state.get('temperature'),
        'top_p': st.session_state.get('top_p'),
        'sampling_top_k': st.session_state.get('sampling_top_k'),
        'presence_penalty': st.session_state.get('presence_penalty'),
        'max_tokens': st.session_state.get('max_tokens'),
        'stream': st.session_state.get('stream'),
        'enable_thinking': st.session_state.get('enable_thinking'),
        'chunk_size': st.session_state.get('chunk_size'),
        'chunk_overlap': st.session_state.get('chunk_overlap'),
    }

    record = {
        'session_id': sid,
        'session_title': st.session_state.get('current_session_title', ''),
        'turn_index': len(st.session_state['user_inputs']),
        'timestamp': datetime.datetime.now().isoformat(timespec='microseconds'),
        'user_id': st.session_state.get('user_id', '_local'),
        'user_message': user_input,
        'assistant_message': response_text or '',
        'reasoning': reasoning or '',
        'model': model,
        'provider': st.session_state.get('provider', ''),
        'base_url': st.session_state.get('base_url', ''),
        'elapsed_seconds': float(elapsed_sec) if elapsed_sec is not None else None,
        'retrieved': retrieved_records,
        'n_retrieved': len(retrieved_records),
        'citation_numbers_used': cited_nums,
        'n_cited': len(cited_nums),
        'query_variants': list(query_variants or []),
        'settings_snapshot': settings,
    }

    try:
        with path.open('a', encoding='utf-8') as f:
            f.write(json.dumps(record, ensure_ascii=False, default=str) + '\n')
    except Exception as e:
        st.warning(f'구조화 로그 저장 실패: {e}')
    _supabase_insert('chat_turns', record)


def auto_title_session() -> str:
    """LLM-generate a short Korean title from the first turn.

    Falls back to a trimmed version of the user's first question when the LLM
    call fails or returns empty content (common with reasoning-style models
    like gpt-5 / o-series that can consume the whole max_tokens budget on
    internal thinking before emitting any visible output).
    """
    if not st.session_state['user_inputs'] or not st.session_state['generated_responses']:
        return ''
    user_msg = st.session_state['user_inputs'][0][:300]
    asst_msg = st.session_state['generated_responses'][0][:300]

    # Always-available fallback derived from the user's question itself.
    fallback = ' '.join((user_msg or '').split())[:16] or '새 대화'

    prompt = (
        '다음 대화 첫 턴을 보고 8글자 내외의 간결한 한국어 제목을 만드세요. '
        '큰따옴표 / 마침표 / 이모지 없이 제목만 한 줄로 출력하세요.\n\n'
        f'사용자: {user_msg}\n어시스턴트: {asst_msg}\n\n제목:'
    )
    try:
        client = get_openai_client()
        tparams = _build_completion_params(
            model=st.session_state['model'],
            messages=[{'role': 'user', 'content': prompt}],
            # Generous budget — reasoning models can spend hundreds of tokens
            # thinking before producing the 8-character title.
            max_tokens=400,
            temperature=0.3,
            extra_body=_thinking_off_extra_body(),
        )
        resp = client.chat.completions.create(**tparams)
        title = (resp.choices[0].message.content or '').strip()
        # If the model returned reasoning then the title on a final line,
        # keep only the last non-empty line.
        if '\n' in title:
            lines = [ln.strip() for ln in title.splitlines() if ln.strip()]
            if lines:
                title = lines[-1]
        title = title.strip('"').strip("'").strip('「').strip('」').strip()
        # Strip any "제목:" prefix the model might echo.
        for prefix in ('제목:', '제목 :', 'Title:'):
            if title.startswith(prefix):
                title = title[len(prefix):].strip()
        if len(title) > 30:
            title = title[:30]
        return title or fallback
    except Exception:
        return fallback


@st.cache_data(show_spinner=False)
def load_page_image_b64(embedder_id: str, doc_id: str, page_no: int):
    """Read and base64-encode a rendered PDF page. Cached so multimodal turns
    don't re-read the same PNG every time."""
    p = _pages_dir(embedder_id, doc_id) / f'{page_no}.png'
    if not p.exists():
        return None
    try:
        return base64.b64encode(p.read_bytes()).decode('ascii')
    except Exception:
        return None


def delete_saved_doc(embedder_id: str, doc_id: str) -> None:
    d = _embedder_dir(embedder_id) / doc_id
    if d.exists():
        shutil.rmtree(d, ignore_errors=True)


def list_saved_doc_ids(embedder_id: str):
    d = _embedder_dir(embedder_id)
    if not d.exists():
        return []
    return [p.name for p in d.iterdir() if p.is_dir() and (p / 'meta.json').exists()]


def _restore_docs_from_pgvector(embedder_id: str) -> list:
    """Pull every chunk this user has stored in Supabase doc_chunks and
    rebuild the in-memory docs[] list. Used on first login from a fresh
    Cloud container — the local .data/ tree is ephemeral, but the
    embeddings live in pgvector and survive restarts.

    Returns the rebuilt list (also persists each doc back to local disk so
    subsequent operations don't have to re-query). Empty list on any failure
    or when the user has no chunks yet."""
    client = _supabase_client()
    if client is None:
        return []
    mapping = _EMBEDDER_TABLE_MAP.get(embedder_id)
    if mapping is None:
        return []
    short_name, vec_col, _dim = mapping

    user_id = st.session_state.get('user_id', '_local')
    try:
        # Pull only what we need to reconstruct. Embeddings come back as
        # the vector type — supabase-py decodes pgvector to a Python list.
        # `.range(0, 9999)` is required because PostgREST default cap is 1000.
        resp = (client.table('doc_chunks')
                .select(f'doc_id, doc_name, chunk_idx, "text", pages, {vec_col}')
                .eq('user_id', user_id)
                .eq('embedder', short_name)
                .order('doc_id')
                .order('chunk_idx')
                .range(0, 9999)
                .execute())
        rows = resp.data or []
    except Exception:
        return []
    if not rows:
        return []

    # Group rows by doc_id, preserving chunk_idx order.
    docs_by_id = {}
    for r in rows:
        did = r.get('doc_id')
        if not did:
            continue
        d = docs_by_id.setdefault(did, {
            'id': did,
            'name': r.get('doc_name') or '',
            'chunks': [],
            'chunk_pages': [],
            'embeddings_list': [],
        })
        d['chunks'].append(r.get('text') or '')
        d['chunk_pages'].append(r.get('pages') or [])
        emb = r.get(vec_col)
        # PostgREST serializes pgvector values as text — e.g.
        # '[0.012, -0.005, ...]'. Decode to a Python list before we try to
        # stack them into a numpy matrix.
        if isinstance(emb, str):
            try:
                emb = json.loads(emb)
            except Exception:
                emb = None
        if not isinstance(emb, list) or len(emb) != mapping[2]:
            # Bad row (dim mismatch / null / decode failed) — fill with
            # zeros so the chunk index still aligns, retrieval will just
            # never match it.
            emb = [0.0] * mapping[2]
        d['embeddings_list'].append(emb)

    rebuilt = []
    for did, d in docs_by_id.items():
        if not d['chunks']:
            continue
        try:
            embs = np.asarray(d['embeddings_list'], dtype=np.float32)
        except Exception:
            # Defensive: if rows somehow have inconsistent dims after the
            # per-row check above, skip this doc entirely rather than
            # crashing the whole restore.
            continue
        # Re-normalize defensively — cosine math elsewhere assumes unit
        # length. supabase usually returns the original normalized values
        # but converting through json + back can drift floating-point.
        norms = np.linalg.norm(embs, axis=1, keepdims=True)
        norms = np.where(norms > 1e-9, norms, 1.0)
        embs = embs / norms

        doc = {
            'id': did,
            'name': d['name'] or did,
            'raw_text': '\n\n'.join(d['chunks']),
            'chunks': d['chunks'],
            'chunk_pages': d['chunk_pages'],
            'page_count': max(
                (max(p) for p in d['chunk_pages'] if p), default=0,
            ),
            'has_page_images': False,
            'is_pdf': False,  # cannot tell from pgvector alone
            'embeddings': embs,
            'chunk_size': 0,
            'chunk_overlap': 0,
        }
        try:
            save_doc(embedder_id, doc)
        except Exception:
            # Disk save is best-effort; the in-memory doc is still usable.
            pass
        rebuilt.append(doc)

    return rebuilt


def load_all_for_current_embedder():
    """Populate st.session_state['docs'] for the current embedder. Tries
    local disk first (fast); on a fresh Cloud container the disk is empty,
    so falls back to rebuilding from Supabase pgvector if configured.
    Skips the work if we already loaded the same embedder this session."""
    eid = st.session_state['embedder_model']
    if st.session_state.get('_loaded_for_embedder') == eid:
        return
    docs = []
    for did in list_saved_doc_ids(eid):
        d = load_doc(eid, did)
        if d is not None:
            docs.append(d)
    # Nothing on disk → try pgvector restore for logged-in users on Cloud.
    if not docs and _supabase_client() is not None:
        docs = _restore_docs_from_pgvector(eid)
        if docs:
            _log_event('docs_restored_from_pgvector', {
                'embedder': eid,
                'n_docs': len(docs),
                'n_chunks_total': sum(len(d['chunks']) for d in docs),
            })
    st.session_state['docs'] = docs
    st.session_state['_loaded_for_embedder'] = eid


# =============================================================================
# Models (cached)
# =============================================================================

@st.cache_resource(show_spinner=False)
def _make_openai_client(base_url: str, api_key: str):
    """Build (and cache) an OpenAI client keyed by base_url+api_key.
    Changing either creates a fresh client; identical configs share one."""
    return OpenAI(base_url=base_url, api_key=api_key)


def _active_api_key() -> str:
    """Return the API key for the currently selected provider."""
    p = st.session_state.get('provider', 'Hugging Face Router')
    if p == 'Hugging Face Router':
        return st.session_state.get('hf_api_key', '') or ''
    if p == 'OpenAI':
        return st.session_state.get('openai_api_key', '') or ''
    if p == 'Anthropic (Claude)':
        return st.session_state.get('anthropic_api_key', '') or ''
    if p == 'Fireworks AI':
        return st.session_state.get('fireworks_api_key', '') or ''
    if p == 'DashScope (Qwen)':
        return st.session_state.get('dashscope_api_key', '') or ''
    # vLLM / local / Custom
    return st.session_state.get('custom_api_key', '') or ''


def get_openai_client():
    """Resolve the active OpenAI-compatible client from session state."""
    return _make_openai_client(
        st.session_state.get('base_url') or '',
        _active_api_key(),
    )


@st.cache_resource(show_spinner=False)
def load_embedder(model_id: str):
    try:
        from transformers.utils import logging as _hf_logging
        _hf_logging.set_verbosity_error()
        _hf_logging.disable_progress_bar()
    except Exception:
        pass
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer(model_id)


@st.cache_resource(show_spinner=False)
def load_reranker(model_id: str):
    from sentence_transformers import CrossEncoder
    return CrossEncoder(model_id)


# =============================================================================
# Parsing & chunking
# =============================================================================

def _read_bytes(file) -> bytes:
    """Read full bytes from a Streamlit UploadedFile, resetting position."""
    pos = file.tell() if hasattr(file, 'tell') else 0
    data = file.read()
    try:
        file.seek(pos)
    except Exception:
        pass
    return data


def parse_pdf_docling(pdf_bytes: bytes) -> dict:
    """Parse PDF with Docling. Returns {elements, page_count, ok}.

    elements is a flat list of {text, page} ordered as on the page.
    Tables are emitted as markdown table strings; lists/headings preserved.
    """
    import tempfile
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return {'elements': [], 'page_count': 0, 'ok': False}
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        d = result.document
    except Exception as e:
        st.warning(f'Docling 파싱 실패, pypdf로 폴백합니다: {e}')
        return {'elements': [], 'page_count': 0, 'ok': False}
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    elements = []
    pages_seen = set()
    for item, _level in d.iterate_items():
        text = ''
        try:
            if hasattr(item, 'export_to_markdown'):
                # Tables, lists, headings: structured markdown.
                text = item.export_to_markdown(d)
        except Exception:
            text = ''
        if not text and hasattr(item, 'text'):
            text = item.text or ''
        text = (text or '').strip()
        if not text:
            continue
        page = None
        prov = getattr(item, 'prov', None)
        if prov:
            try:
                page = prov[0].page_no
            except Exception:
                page = None
        if page is not None:
            pages_seen.add(page)
        elements.append({'text': text, 'page': page})

    page_count = max(pages_seen) if pages_seen else 0
    return {'elements': elements, 'page_count': page_count, 'ok': True}


def parse_pdf_pypdf(pdf_bytes: bytes) -> dict:
    """Fallback PDF parser. Page boundaries known but no structure."""
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(pdf_bytes))
    elements = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or '').strip()
        if text:
            elements.append({'text': text, 'page': i + 1})
    return {'elements': elements, 'page_count': len(reader.pages), 'ok': True}


def parse_pdf_ocr(pdf_bytes: bytes) -> dict:
    """Last-resort PDF parser for scanned / image-only PDFs.

    Renders each page to PNG via PyMuPDF and runs Tesseract OCR on it.
    Returns {} when the OCR stack isn't installed so the caller can fall
    through gracefully. Korean + English are both enabled.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
        import io
    except Exception:
        return {'elements': [], 'page_count': 0, 'ok': False,
                'reason': 'ocr_dependencies_missing'}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    except Exception:
        return {'elements': [], 'page_count': 0, 'ok': False,
                'reason': 'pdf_open_failed'}
    elements = []
    page_count = len(doc)
    for i, page in enumerate(doc):
        try:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes('png')))
            # Korean + English. Falls back to English-only if kor pack missing.
            try:
                text = pytesseract.image_to_string(img, lang='kor+eng')
            except Exception:
                text = pytesseract.image_to_string(img)
            text = (text or '').strip()
            if text:
                elements.append({'text': text, 'page': i + 1})
        except Exception:
            continue
    doc.close()
    return {'elements': elements, 'page_count': page_count, 'ok': True}


def parse_docx_bytes(data: bytes) -> dict:
    """Parse DOCX (Word) into the same elements/page format we use for PDFs.

    Tries Docling first (better structure / table support), falls back to
    python-docx (simple paragraph extraction). DOCX doesn't have stable
    page numbers in the XML, so we leave page=None — chunking still works,
    citations just won't show page numbers."""
    # Docling path
    try:
        from docling.document_converter import DocumentConverter
        import io
        conv = DocumentConverter()
        # Docling accepts a path or a stream-like with name attribute.
        class _NamedBytesIO(io.BytesIO):
            name = 'in.docx'
        buf = _NamedBytesIO(data)
        result = conv.convert(buf)
        md = result.document.export_to_markdown()
        if md and md.strip():
            paragraphs = [p.strip() for p in md.split('\n\n') if p.strip()]
            return {
                'elements': [{'text': p, 'page': None} for p in paragraphs],
                'page_count': 0,
                'ok': True,
            }
    except Exception:
        pass

    # python-docx fallback
    try:
        from docx import Document
        import io
        d = Document(io.BytesIO(data))
        elements = []
        for p in d.paragraphs:
            t = (p.text or '').strip()
            if t:
                elements.append({'text': t, 'page': None})
        # Tables: each cell becomes its own element so a chunk can hold
        # one row's worth of facts.
        for table in d.tables:
            for row in table.rows:
                cells = [(c.text or '').strip() for c in row.cells]
                joined = ' | '.join(c for c in cells if c)
                if joined:
                    elements.append({'text': joined, 'page': None})
        if elements:
            return {'elements': elements, 'page_count': 0, 'ok': True}
    except Exception:
        pass

    return {'elements': [], 'page_count': 0, 'ok': False}


def parse_csv_bytes(data: bytes, file_name: str = '') -> dict:
    """Parse CSV. Each row becomes one element formatted as
    'col1: val1 | col2: val2 | ...' so the LLM sees field-name context
    around each value, not just numbers in isolation. Also includes a
    header-summary element so semantic search can match queries like
    "what columns does this file have"."""
    import csv as _csv
    import io as _io

    # Best-effort decoding: try UTF-8 (with BOM), fall back to CP949
    # which is what Excel saves CSV as on Korean Windows.
    for enc in ('utf-8-sig', 'utf-8', 'cp949', 'euc-kr', 'latin-1'):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        return {'elements': [], 'page_count': 0, 'ok': False}

    # Sniff delimiter — fall back to comma if sniffing fails.
    sample = text[:4096]
    try:
        dialect = _csv.Sniffer().sniff(sample, delimiters=',;\t|')
    except _csv.Error:
        class _Default(_csv.Dialect):
            delimiter = ','
            quotechar = '"'
            doublequote = True
            skipinitialspace = True
            lineterminator = '\n'
            quoting = _csv.QUOTE_MINIMAL
        dialect = _Default

    reader = _csv.reader(_io.StringIO(text), dialect)
    rows = list(reader)
    if not rows:
        return {'elements': [], 'page_count': 0, 'ok': True}

    header = [(c or '').strip() for c in rows[0]]
    data_rows = rows[1:]
    elements = []

    # 1) Summary element — column names + row count.
    summary = (
        f"[CSV 요약] 파일: {file_name or 'csv'}, "
        f"열 {len(header)}개, 행 {len(data_rows)}개. "
        f"열 이름: {', '.join(c for c in header if c)}."
    )
    elements.append({'text': summary, 'page': None})

    # 2) Per-row elements — formatted as "header: value | header: value".
    for i, row in enumerate(data_rows, start=1):
        cells = []
        for h, v in zip(header, row):
            v = (v or '').strip()
            if not v:
                continue
            if h:
                cells.append(f'{h}: {v}')
            else:
                cells.append(v)
        if cells:
            elements.append({
                'text': f'행 {i}: ' + ' | '.join(cells),
                'page': None,
            })

    return {'elements': elements, 'page_count': 0, 'ok': True}


def parse_hwpx_bytes(data: bytes) -> dict:
    """Parse HWPX (Hancom Office XML). HWPX is a ZIP archive containing
    section XML files under Contents/. We pull <hp:t> (text run) elements
    from each section in order and treat each section as a 'page' so the
    citation system can still surface a section number."""
    import io as _io
    import zipfile as _zip
    from xml.etree import ElementTree as _ET

    try:
        zf = _zip.ZipFile(_io.BytesIO(data))
    except _zip.BadZipFile:
        return {'elements': [], 'page_count': 0, 'ok': False}

    section_names = sorted(
        (n for n in zf.namelist()
         if n.startswith('Contents/section') and n.endswith('.xml')),
        key=lambda n: n,
    )
    if not section_names:
        zf.close()
        return {'elements': [], 'page_count': 0, 'ok': False}

    elements = []
    # HWPX uses the 'hp' namespace; strip namespaces so we don't have to
    # bind them — just match local tag names.
    def _local(tag: str) -> str:
        return tag.split('}', 1)[1] if '}' in tag else tag

    for section_idx, name in enumerate(section_names, start=1):
        try:
            raw = zf.read(name)
            root = _ET.fromstring(raw)
        except Exception:
            continue
        # Each <p> (paragraph) becomes one element. Within a paragraph we
        # concatenate the text of all <t> (text run) descendants.
        for p in root.iter():
            if _local(p.tag) != 'p':
                continue
            buf = []
            for t in p.iter():
                if _local(t.tag) == 't' and t.text:
                    buf.append(t.text)
            joined = ''.join(buf).strip()
            if joined:
                elements.append({'text': joined, 'page': section_idx})

    zf.close()
    return {'elements': elements, 'page_count': len(section_names), 'ok': True}


def parse_file(file) -> dict:
    """Returns a dict: {raw_text, elements, page_count, is_pdf, pdf_bytes}.

    elements: list of {text, page}.  page is int (1-indexed) for PDFs,
    None for plain text files.  pdf_bytes is retained so the caller can
    render page images later (only set when is_pdf=True).
    """
    name = file.name.lower()
    if name.endswith('.pdf'):
        data = _read_bytes(file)
        # Stage 1: Docling (structure-aware)
        parsed = parse_pdf_docling(data)
        # Stage 2: pypdf if Docling missed
        if not parsed['ok'] or not parsed['elements']:
            try:
                parsed = parse_pdf_pypdf(data)
            except Exception as e:
                parsed = {'elements': [], 'page_count': 0, 'ok': False,
                          'error': str(e)}
        # Stage 3: OCR if both text extractors got essentially nothing —
        # this catches scanned / image-only PDFs that don't have a text
        # layer to extract.
        total_chars = sum(len(e['text']) for e in parsed.get('elements', []))
        if total_chars < 50:
            try:
                with st.spinner(f'{file.name}: 텍스트 레이어 없음 — OCR 시도 중...'):
                    ocr = parse_pdf_ocr(data)
                if ocr['ok'] and ocr['elements']:
                    parsed = ocr
                    st.info(
                        f'{file.name}: 스캔 PDF 로 판단해 OCR (Tesseract) 로 텍스트 추출했습니다. '
                        f'페이지 {ocr["page_count"]}, 문자 {sum(len(e["text"]) for e in ocr["elements"])}.'
                    )
                elif ocr.get('reason') == 'ocr_dependencies_missing':
                    st.warning(
                        f'{file.name}: 텍스트 추출이 안 되고 OCR 라이브러리도 없어 인덱스 품질이 낮을 수 있습니다. '
                        f'`pip install pytesseract pymupdf` + 시스템 tesseract 설치 시 스캔 PDF 도 처리됩니다.'
                    )
            except Exception:
                pass
        if not parsed.get('elements'):
            st.error(f'PDF 파싱 실패 ({file.name}): 텍스트를 추출할 수 없습니다.')
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': True, 'pdf_bytes': data}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': parsed['page_count'],
            'is_pdf': True,
            'pdf_bytes': data,
        }
    if name.endswith('.docx'):
        data = _read_bytes(file)
        parsed = parse_docx_bytes(data)
        if not parsed['ok']:
            st.error(
                f'DOCX 파싱 실패 ({file.name}). '
                f'`pip install python-docx` 가 필요할 수 있습니다.'
            )
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': 0,
            'is_pdf': False,
            'pdf_bytes': None,
        }
    if name.endswith('.doc'):
        st.error(
            f'DOC (구버전 바이너리) 는 직접 지원하지 않습니다 ({file.name}). '
            f'Word 또는 LibreOffice 에서 DOCX 또는 PDF 로 다시 저장한 뒤 업로드해 주세요.'
        )
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    if name.endswith('.csv'):
        data = _read_bytes(file)
        parsed = parse_csv_bytes(data, file_name=file.name)
        if not parsed['ok']:
            st.error(f'CSV 디코딩 실패 ({file.name}): 인코딩이 UTF-8/CP949 가 아닌 듯합니다.')
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': 0,
            'is_pdf': False,
            'pdf_bytes': None,
        }
    if name.endswith('.hwpx'):
        data = _read_bytes(file)
        parsed = parse_hwpx_bytes(data)
        if not parsed['ok']:
            st.error(
                f'HWPX 파싱 실패 ({file.name}): 손상된 파일이거나 보안 처리된 '
                f'HWPX 일 수 있습니다. 한컴오피스에서 다시 저장하거나 PDF 로 '
                f'변환해 보세요.'
            )
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': parsed['page_count'],  # section count, used like "page"
            'is_pdf': False,
            'pdf_bytes': None,
        }
    # Reject .hwp (binary) explicitly with a helpful message — different format.
    if name.endswith('.hwp'):
        st.error(
            f'HWP (구버전 바이너리) 는 직접 지원하지 않습니다 ({file.name}). '
            f'한컴오피스에서 HWPX 또는 PDF 로 다시 저장한 뒤 업로드해 주세요.'
        )
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    try:
        text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        st.error(f'파일 읽기 실패 ({file.name}): {e}')
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    return {
        'raw_text': text,
        'elements': [{'text': text, 'page': None}] if text.strip() else [],
        'page_count': 0,
        'is_pdf': False,
        'pdf_bytes': None,
    }


_SENT_SPLIT_RE = re.compile(
    r'(?<=[.!?。!?])\s+(?=\S)|(?<=[다요죠음음됨함임함])\s*\n+\s*(?=\S)'
)


def _split_long_text_to_sentences(text: str) -> list:
    """Split a long block of text into rough sentence-like spans.

    Uses Latin sentence terminators (.!?) and the most common Korean sentence
    endings (다/요/죠/음/됨/함) followed by whitespace/newline. Output is a
    plain list of non-empty strings preserving original order; callers can
    glue them back together up to the size budget.
    """
    if not text:
        return []
    parts = _SENT_SPLIT_RE.split(text)
    # Some splits leave empty strings or leading/trailing whitespace.
    return [p.strip() for p in parts if p and p.strip()]


def _expand_into_paragraphs(elements: list) -> list:
    """Pre-process: explode any element whose text contains paragraph breaks
    into multiple smaller elements sharing the same page metadata.

    pypdf often returns one element per page = whole-page text glued together.
    Docling normally pre-splits at the layout level, so this is mostly a
    no-op for Docling input. The downstream chunker then groups paragraphs
    instead of being forced into raw char-step splits.
    """
    out = []
    for el in elements:
        text = (el.get('text') or '').strip()
        if not text:
            continue
        page = el.get('page')
        # Treat any run of newlines (with optional whitespace) as a boundary.
        paragraphs = re.split(r'\n\s*\n+', text)
        for p in paragraphs:
            p = p.strip()
            if p:
                out.append({'text': p, 'page': page})
    return out


def chunk_elements(elements: list, size: int, overlap: int):
    """Group elements (each {text, page}) into chunks of ~size chars.

    Paragraph-aware: adjacent short paragraphs are merged up to size; long
    paragraphs are split on sentence boundaries (not raw character offsets)
    so chunks don't end mid-sentence. Page metadata is preserved.

    Returns (chunks: list[str], chunk_pages: list[list[int]]).
    """
    if not elements:
        return [], []
    size = max(50, int(size))
    overlap = max(0, min(int(overlap), size - 1))

    # Pre-explode whole-page elements into paragraphs (no-op for already-fine
    # input — keeps Docling-style elements untouched).
    elements = _expand_into_paragraphs(elements)

    chunks, chunk_pages = [], []
    buf_text, buf_pages = '', set()

    def flush():
        nonlocal buf_text, buf_pages
        if buf_text:
            chunks.append(buf_text)
            chunk_pages.append(sorted(buf_pages))
            buf_text, buf_pages = '', set()

    for el in elements:
        text = (el.get('text') or '').strip()
        if not text:
            continue
        page = el.get('page')
        candidate = (buf_text + '\n\n' + text).strip() if buf_text else text
        if len(candidate) <= size:
            buf_text = candidate
            if page is not None:
                buf_pages.add(page)
            continue
        # candidate exceeds size: emit current buffer first.
        flush()
        if len(text) <= size:
            buf_text = text
            buf_pages = {page} if page is not None else set()
            continue
        # A single paragraph is longer than size — split by sentence boundary.
        sentences = _split_long_text_to_sentences(text)
        if not sentences:
            # Last-resort fallback: char-step (matches old behavior).
            step = max(1, size - overlap)
            for i in range(0, len(text), step):
                chunks.append(text[i:i + size])
                chunk_pages.append([page] if page is not None else [])
            continue
        sbuf = ''
        for s in sentences:
            cand = (sbuf + ' ' + s).strip() if sbuf else s
            if len(cand) <= size:
                sbuf = cand
                continue
            if sbuf:
                chunks.append(sbuf)
                chunk_pages.append([page] if page is not None else [])
                # Sentence-level overlap: keep tail of previous chunk if small.
                if overlap and len(sbuf) > overlap:
                    sbuf = sbuf[-overlap:] + ' ' + s
                else:
                    sbuf = s
            else:
                # A single sentence longer than size — char-step it.
                step = max(1, size - overlap)
                for i in range(0, len(s), step):
                    chunks.append(s[i:i + size])
                    chunk_pages.append([page] if page is not None else [])
                sbuf = ''
        if sbuf:
            chunks.append(sbuf)
            chunk_pages.append([page] if page is not None else [])
    flush()

    # Merge tiny trailing chunk into its predecessor.
    if len(chunks) >= 2 and len(chunks[-1]) < max(50, size // 5):
        chunks[-2] = (chunks[-2] + '\n\n' + chunks[-1])[: size + overlap]
        merged_pages = sorted(set(chunk_pages[-2]) | set(chunk_pages[-1]))
        chunk_pages[-2] = merged_pages
        chunks.pop()
        chunk_pages.pop()

    return chunks, chunk_pages


def chunk_text(text: str, size: int, overlap: int) -> list:
    """Backward-compatible: plain-text chunking returning list of strings only."""
    chunks, _pages = chunk_elements(
        [{'text': text, 'page': None}], size, overlap
    )
    return chunks


def render_pdf_pages_to_dir(pdf_bytes: bytes, out_dir: Path, dpi: int = 144) -> int:
    """Render each PDF page to PNG into out_dir/{page}.png. Returns number of pages."""
    import fitz
    out_dir.mkdir(parents=True, exist_ok=True)
    n = 0
    with fitz.open(stream=pdf_bytes, filetype='pdf') as d:
        for i, page in enumerate(d):
            pix = page.get_pixmap(dpi=dpi)
            pix.save(str(out_dir / f'{i + 1}.png'))
            n = i + 1
    return n


# =============================================================================
# Retrieval
# =============================================================================

def _tokenize_for_bm25(text: str):
    # Mixed-script tokenization: words/CJK/numbers. Lowercase ASCII.
    text = (text or '').lower()
    return re.findall(r"[\w가-힯一-鿿]+", text)


def build_bm25_over_docs(docs: list):
    try:
        from rank_bm25 import BM25Okapi
    except ImportError:
        return None, None
    if not docs:
        return None, None
    all_chunks, meta = [], []
    for doc in docs:
        for i, ch in enumerate(doc['chunks']):
            all_chunks.append(ch)
            meta.append((doc['id'], doc['name'], i))
    tokenized = [_tokenize_for_bm25(c) for c in all_chunks]
    bm25 = BM25Okapi(tokenized)
    return bm25, meta


def _flatten_chunks(docs: list):
    all_chunks, all_embs, meta = [], [], []
    for doc in docs:
        embs = doc.get('embeddings')
        if embs is None or len(embs) != len(doc['chunks']):
            continue
        for i, ch in enumerate(doc['chunks']):
            all_chunks.append(ch)
            meta.append((doc['id'], doc['name'], i))
        all_embs.append(embs)
    if not all_embs:
        return [], np.zeros((0, 1), dtype=np.float32), []
    return all_chunks, np.vstack(all_embs), meta


def _dense_search_pgvector(query: str, top_n: int):
    """Top-k cosine search via Supabase pgvector RPC. Returns the same shape
    as the in-memory variant — (chunk_idx_within_doc, score, meta, text) —
    where meta is (doc_id, doc_name, chunk_idx). Falls back to None on any
    issue so the caller can revert to the in-memory path."""
    client = _supabase_client()
    if client is None:
        return None
    embedder_id = st.session_state['embedder_model']
    mapping = _EMBEDDER_TABLE_MAP.get(embedder_id)
    if mapping is None:
        return None
    _short_name, _vec_col, _dim = mapping
    rpc_name = ('match_chunks_minilm' if embedder_id.endswith('MiniLM-L12-v2')
                else 'match_chunks_bgem3')

    embedder = load_embedder(embedder_id)
    q = embedder.encode(
        [query], convert_to_numpy=True, normalize_embeddings=True,
        show_progress_bar=False,
    )[0]

    # Honor the chat-side document filter so pgvector doesn't return chunks
    # the user has hidden.
    doc_filter = st.session_state.get('chat_doc_filter') or []
    docs = st.session_state.get('docs') or []
    if doc_filter and len(doc_filter) < len(docs):
        p_doc_ids = list(doc_filter)
    else:
        p_doc_ids = None

    try:
        params = {
            'p_user_id': st.session_state.get('user_id', '_local'),
            'p_query_embedding': q.tolist(),
            'p_match_count': int(top_n),
            'p_doc_ids': p_doc_ids,
        }
        resp = client.rpc(rpc_name, params).execute()
        rows = resp.data or []
    except Exception as e:
        st.session_state['_pgv_search_last_err'] = (
            f'{rpc_name}: {type(e).__name__}: {str(e)[:600]}'
        )
        return None

    if not rows:
        # Empty result is a legitimate answer; surface it as such instead
        # of silently falling back, so users notice if their data isn't in
        # pgvector yet (e.g. ingested before the schema was applied).
        st.session_state['_pgv_search_last_n'] = 0
        return []

    out = []
    for r in rows:
        meta = (r.get('doc_id'), r.get('doc_name'), int(r.get('chunk_idx', 0)))
        out.append((
            int(r.get('id', 0)),
            float(r.get('score', 0.0)),
            meta,
            r.get('text') or '',
        ))
    st.session_state['_pgv_search_last_n'] = len(out)
    return out


def dense_search(query: str, top_n: int):
    # Route to pgvector when the user has opted in (Phase 2b). On any
    # failure or when the result is unusable, transparently fall back to
    # the in-memory numpy path.
    if st.session_state.get('use_pgvector_search'):
        pg = _dense_search_pgvector(query, top_n)
        if pg is not None:
            return pg

    docs = st.session_state['docs']
    chunks, embs, meta = _flatten_chunks(docs)
    if not chunks:
        return []
    embedder = load_embedder(st.session_state['embedder_model'])
    q = embedder.encode(
        [query], convert_to_numpy=True, normalize_embeddings=True,
        show_progress_bar=False,
    )[0]
    sims = embs @ q
    idx = np.argsort(-sims)[:top_n]
    return [(int(i), float(sims[i]), meta[i], chunks[i]) for i in idx]


def bm25_search(query: str, top_n: int):
    docs = st.session_state['docs']
    bm25, meta = build_bm25_over_docs(docs)
    if bm25 is None or not meta:
        return []
    chunks, _, _ = _flatten_chunks(docs)
    scores = bm25.get_scores(_tokenize_for_bm25(query))
    idx = np.argsort(-scores)[:top_n]
    return [(int(i), float(scores[i]), meta[i], chunks[i]) for i in idx]


def rrf_fuse(rankings: list, k: int = 60):
    """rankings: list of [(idx, score, meta, chunk), ...]; fuse via Reciprocal Rank Fusion."""
    rrf_scores = {}
    payload = {}
    for ranking in rankings:
        for rank, item in enumerate(ranking):
            idx = item[0]
            rrf_scores[idx] = rrf_scores.get(idx, 0.0) + 1.0 / (k + rank)
            payload[idx] = item
    fused = sorted(rrf_scores.items(), key=lambda x: -x[1])
    out = []
    for idx, s in fused:
        i, _, m, c = payload[idx]
        out.append((i, s, m, c))
    return out


def rerank(query: str, candidates: list, top_k: int):
    if not candidates:
        return []
    reranker = load_reranker(RERANKER_MODEL)
    pairs = [[query, c[3]] for c in candidates]
    scores = reranker.predict(pairs, show_progress_bar=False)
    scored = list(zip(candidates, scores))
    scored.sort(key=lambda x: -float(x[1]))
    out = []
    for (i, _, m, c), s in scored[:top_k]:
        out.append((i, float(s), m, c))
    return out


def _is_openai_endpoint() -> bool:
    """Detect OpenAI's official API (rejects non-standard params like top_k)."""
    return 'api.openai.com' in (st.session_state.get('base_url') or '').lower()


def _is_dashscope_endpoint() -> bool:
    return 'dashscope' in (st.session_state.get('base_url') or '').lower()


def _is_hf_router_endpoint() -> bool:
    return 'router.huggingface.co' in (st.session_state.get('base_url') or '').lower()


def _is_anthropic_endpoint() -> bool:
    """Claude via Anthropic's OpenAI-SDK compatibility layer. Has its own
    parameter restrictions (no top_p, no penalty fields, no top_k)."""
    return 'api.anthropic.com' in (st.session_state.get('base_url') or '').lower()


def _is_fireworks_endpoint() -> bool:
    return 'fireworks.ai' in (st.session_state.get('base_url') or '').lower()


def _provider_supports_top_k() -> bool:
    """top_k via extra_body is provider-specific. Live-tested:
      - OpenAI: rejects.
      - Anthropic OpenAI-compat: rejects.
      - HF Router: many sub-providers reject (e.g. Cerebras for Llama-3.1-8B,
        gpt-oss-120b). Safest to omit unless DashScope/vLLM/Custom.
      - Fireworks: rejects via extra_body.
      - DashScope (Qwen): supports.
      - vLLM/local & Custom: typically supports."""
    if _is_openai_endpoint():
        return False
    if _is_anthropic_endpoint():
        return False
    if _is_hf_router_endpoint():
        return False
    if _is_fireworks_endpoint():
        return False
    return True


def _is_openai_reasoning_model(model: str) -> bool:
    """OpenAI o-series — strictest constraints: no temperature/top_p/penalty."""
    m = (model or '').lower()
    return m.startswith('o1') or m.startswith('o3') or m.startswith('o4')


def _is_openai_gpt5_family(model: str) -> bool:
    """GPT-5/5.x chat models — require max_completion_tokens, reject non-default
    temperature/top_p/penalty."""
    return (model or '').lower().startswith('gpt-5')


def _uses_max_completion_tokens(model: str) -> bool:
    """Models that require max_completion_tokens instead of max_tokens (OpenAI new tier)."""
    return _is_openai_reasoning_model(model) or _is_openai_gpt5_family(model)


# Models we used to expose but pulled because they can't be called through
# their configured endpoint. The runtime guard below swaps them on the fly
# before the API request goes out — protects users whose saved prefs / chat
# state still reference them after we removed them from the dropdown.
_DEPRECATED_MODEL_SWAPS = {
    # Korean-native models: HF Inference Providers don't deploy them.
    'LGAI-EXAONE/EXAONE-4.5-33B': ('Qwen/Qwen3-Next-80B-A3B-Instruct', 'Hugging Face Router'),
    'naver-hyperclovax/HyperCLOVAX-SEED-Think-32B': ('Qwen/Qwen3-Next-80B-A3B-Instruct', 'Hugging Face Router'),
    # OpenAI Responses-API-only models.
    'gpt-5-pro':     ('gpt-5-mini', 'OpenAI'),
    'gpt-5.5-pro':   ('gpt-5-mini', 'OpenAI'),
    'gpt-5.4-pro':   ('gpt-5-mini', 'OpenAI'),
    'o1-pro':        ('o4-mini',     'OpenAI'),
}


def _resolve_deprecated_model(model: str) -> str:
    """Swap a deprecated model id for its replacement at API-call time.
    Also updates st.session_state so the picker UI reflects the swap. Shows
    a one-time-per-session notice so the user knows it happened."""
    if model not in _DEPRECATED_MODEL_SWAPS:
        return model
    replacement, _expected_provider = _DEPRECATED_MODEL_SWAPS[model]
    notified = st.session_state.setdefault('_deprecated_model_notified', set())
    if model not in notified:
        st.info(
            f'`{model}` 은 이 endpoint 에서 서빙되지 않아 '
            f'`{replacement}` 으로 자동 전환했습니다. '
            f'설정 탭에서 다른 모델로 변경할 수도 있습니다.'
        )
        notified.add(model)
    st.session_state['model'] = replacement
    return replacement


def _build_completion_params(
    model: str, messages: list,
    *,
    max_tokens=None, temperature=None, top_p=None, presence_penalty=None,
    extra_body: dict = None, stream: bool = False,
) -> dict:
    """Build chat.completions.create kwargs that respect per-model constraints.
    Drops sampling params for GPT-5/o-series; uses max_completion_tokens for those."""
    model = _resolve_deprecated_model(model)
    out = {'model': model, 'messages': messages}
    if stream:
        out['stream'] = True

    is_openai = _is_openai_endpoint()
    if max_tokens is not None:
        budget = int(max_tokens)
        if is_openai and _is_openai_reasoning_model(model):
            # Reasoning eats budget internally; ensure room for actual output.
            budget = max(budget, 1500)
        if is_openai and _uses_max_completion_tokens(model):
            out['max_completion_tokens'] = budget
        else:
            out['max_tokens'] = budget

    # Sampling param compatibility per provider.
    # - OpenAI gpt-5/o-series: rejects ANY non-default sampling param.
    # - Anthropic (Claude 4.x via OpenAI compat): rejects temperature,
    #   top_p, and presence_penalty on most current models. Claude moved
    #   to model-managed sampling and returns
    #     'temperature is deprecated for this model.'
    #   when any of these are sent. Safest to omit all three; the
    #   default sampling produces good output.
    is_anthropic = _is_anthropic_endpoint()
    openai_strict = is_openai and (
        _is_openai_reasoning_model(model) or _is_openai_gpt5_family(model)
    )
    if not openai_strict and not is_anthropic:
        if temperature is not None:
            out['temperature'] = float(temperature)
        if top_p is not None:
            out['top_p'] = float(top_p)
        if presence_penalty is not None:
            out['presence_penalty'] = float(presence_penalty)

    # extra_body — strip provider-incompatible fields as a safety net.
    if extra_body:
        if is_openai or is_anthropic:
            extra_body = {k: v for k, v in extra_body.items()
                          if k not in ('top_k', 'chat_template_kwargs', 'enable_thinking')}
        if extra_body:
            out['extra_body'] = extra_body
    return out


def _thinking_off_extra_body() -> dict:
    """Return extra_body that disables thinking, branched by provider.

    - OpenAI: no thinking concept; rejects unknown fields → {}.
    - HF Router: providers vary; Cerebras-served models (Llama-3.1-8B,
      gpt-oss, etc.) reject chat_template_kwargs with HTTP 400. Safer to
      omit; thinking-capable models on the router will still emit <think>
      tags which we parse out in split_thinking().
    - DashScope (Qwen): uses enable_thinking directly in extra_body.
    - vLLM / Custom: chat_template_kwargs is the standard way (SGLang/vLLM).
    """
    if _is_openai_endpoint():
        return {}
    if _is_anthropic_endpoint():
        # Claude exposes extended thinking via a 'thinking' parameter at the
        # native API level, but the OpenAI-compat layer ignores it. Return
        # nothing extra to avoid 400 from unknown fields.
        return {}
    if _is_fireworks_endpoint():
        return {}
    if _is_dashscope_endpoint():
        return {'enable_thinking': False}
    if _is_hf_router_endpoint():
        return {}
    return {'chat_template_kwargs': {'enable_thinking': False}}


def rewrite_with_context(query: str) -> str:
    """If conversation history exists, rewrite the current query into a
    standalone form (resolving pronouns / implied subjects). Cheap fallback
    to original on any failure."""
    history_u = st.session_state['user_inputs']
    history_a = st.session_state['generated_responses']
    if not history_u:
        return query
    try:
        client = get_openai_client()
    except Exception:
        return query
    # Use only the last 3 turns to keep this cheap.
    hist_pairs = list(zip(history_u[-3:], history_a[-3:]))
    history_str = '\n'.join(f"User: {u}\nAssistant: {a}" for u, a in hist_pairs)
    prompt = (
        "Given the conversation history and the user's latest message, rewrite "
        "the latest message as a fully self-contained question. Resolve pronouns "
        "(그것/그게/it/this), fill in implied subjects, and make the question "
        "stand on its own. If it is already self-contained, return it unchanged. "
        "Output only the rewritten sentence — no preface, no quotes.\n\n"
        f"History:\n{history_str}\n\n"
        f"Latest: {query}\n\nRewritten:"
    )
    try:
        params = _build_completion_params(
            model=st.session_state['model'],
            messages=[{'role': 'user', 'content': prompt}],
            max_tokens=200, temperature=0.0,
            extra_body=_thinking_off_extra_body(),
        )
        resp = client.chat.completions.create(**params)
        rewritten = (resp.choices[0].message.content or '').strip()
        rewritten = rewritten.strip('"').strip("'").strip()
        return rewritten or query
    except Exception as e:
        st.warning(f'쿼리 재작성 실패 ({e}). 원본 질문으로 검색합니다.')
        return query


def expand_queries(query: str) -> list:
    """Return [effective_query, paraphrase1, ..., hyde]. The first variant is
    the contextually-rewritten query if rewriting is enabled and history exists;
    otherwise it is the original query."""
    base = query
    if st.session_state.get('use_contextual_rewrite') and st.session_state['user_inputs']:
        rewritten = rewrite_with_context(query)
        if rewritten and rewritten != query:
            base = rewritten
    variants = [base]
    # Always also keep the literal user query so retrieval is robust if the
    # rewrite drifted semantically.
    if base != query:
        variants.append(query)

    if not (st.session_state.get('use_multi_query') or st.session_state.get('use_hyde')):
        return variants
    try:
        client = get_openai_client()
    except Exception:
        return variants
    model = st.session_state['model']
    eb = _thinking_off_extra_body()

    if st.session_state.get('use_multi_query'):
        n = int(st.session_state.get('n_paraphrases', 3))
        prompt = (
            f"Rewrite the following question in {n} different ways while preserving "
            f"the meaning. Output one paraphrase per line. No numbering, no quotes, "
            f"no explanation.\n\nQuestion: {base}"
        )
        try:
            mq_params = _build_completion_params(
                model=model,
                messages=[{'role': 'user', 'content': prompt}],
                max_tokens=300, temperature=0.7,
                extra_body=eb,
            )
            resp = client.chat.completions.create(**mq_params)
            text = (resp.choices[0].message.content or '').strip()
            added = 0
            for line in text.split('\n'):
                clean = line.strip().lstrip('-•*').strip()
                # Strip leading numbering like "1." or "1)"
                m = re.match(r'^\d+\s*[.)]\s*(.*)$', clean)
                if m:
                    clean = m.group(1)
                if clean and clean not in variants:
                    variants.append(clean)
                    added += 1
                    if added >= n:
                        break
        except Exception as e:
            st.warning(f'Multi-query 생성 실패: {e}')

    if st.session_state.get('use_hyde'):
        prompt = (
            "Write a concise factual paragraph that would hypothetically answer "
            "the following question, as if extracted from an authoritative document. "
            "No preface, no commentary — just the paragraph.\n\n"
            f"Question: {base}"
        )
        try:
            hy_params = _build_completion_params(
                model=model,
                messages=[{'role': 'user', 'content': prompt}],
                max_tokens=250, temperature=0.3,
                extra_body=eb,
            )
            resp = client.chat.completions.create(**hy_params)
            hyde = (resp.choices[0].message.content or '').strip()
            if hyde:
                variants.append(hyde)
        except Exception as e:
            st.warning(f'HyDE 생성 실패: {e}')

    return variants


_COMPARISON_KEYWORDS = (
    '비교', '대조', '차이', '공통점', '공통', '유사점', '다른점', '대비', '구분',
    'compare', 'comparison', 'differ', 'difference', 'similar', 'versus', ' vs ',
)


def _is_comparison_query(query: str) -> bool:
    """Heuristic — does this query ask to compare/contrast across sources?"""
    q = (query or '').lower()
    return any(kw.lower() in q for kw in _COMPARISON_KEYWORDS)


def _select_with_per_doc_min(ranked: list, top_k: int, per_doc_min: int) -> list:
    """Pick top_k from a ranked list while reserving at least `per_doc_min` slots
    per document. Web results are treated as a single 'web' bucket.

    Phase 1: walk ranked list, assign each item to its doc's quota until quota fills.
    Phase 2: fill remaining slots with the highest-scoring not-yet-selected items."""
    if per_doc_min <= 0 or top_k <= 0:
        return ranked[:top_k]
    selected = []
    used = set()
    by_doc = {}

    def _bucket(item):
        # ranked item = (idx, score, meta, chunk). meta = (doc_id, doc_name, chunk_idx) for doc.
        # For web items we pass through a different shape; treat as 'web' bucket.
        try:
            meta = item[2]
            doc_id = meta[0] if isinstance(meta, tuple) else None
            return doc_id or 'web'
        except Exception:
            return 'web'

    for item in ranked:
        if len(selected) >= top_k:
            break
        b = _bucket(item)
        if by_doc.get(b, 0) < per_doc_min:
            selected.append(item)
            used.add(id(item))
            by_doc[b] = by_doc.get(b, 0) + 1

    for item in ranked:
        if len(selected) >= top_k:
            break
        if id(item) not in used:
            selected.append(item)
            used.add(id(item))
    return selected


def _effective_per_doc_min(query: str) -> int:
    """Apply user toggles + comparison autodetect. Returns the per-doc reserve to
    use for this query, considering how many documents the user has."""
    n_docs = len(st.session_state.get('docs', []))
    if n_docs < 2:
        return 0
    reserve = 0
    if st.session_state.get('per_doc_balance'):
        reserve = max(reserve, int(st.session_state.get('per_doc_reserve', 1)))
    if (st.session_state.get('comparison_autodetect')
            and _is_comparison_query(query)):
        reserve = max(reserve, 2)
    # Don't reserve more than would fit
    final_top_k = int(st.session_state.get('final_top_k', 5))
    if reserve * n_docs > final_top_k:
        # Allow at most floor(top_k / n_docs); minimum 1 if any reserve requested
        reserve = max(1, final_top_k // n_docs)
    return reserve


def _single_query_local_search(q: str, mode: str, top_n: int):
    """Run dense+/or BM25 for one query, return RRF-fused ranking (or single)."""
    rankings = []
    if mode in ('dense', 'hybrid'):
        rankings.append(dense_search(q, top_n))
    if mode in ('bm25', 'hybrid'):
        rankings.append(bm25_search(q, top_n))
    rankings = [r for r in rankings if r]
    if not rankings:
        return []
    if len(rankings) == 1:
        return rankings[0]
    return rrf_fuse(rankings)


def retrieve_local(query: str) -> list:
    """Hybrid retrieval with optional HyDE/multi-query expansion + reranker."""
    if not st.session_state['docs']:
        st.session_state['_last_variants'] = [query]
        return []
    mode = st.session_state['retrieval_mode']
    top_n = int(st.session_state['retrieve_top_n'])
    top_k = int(st.session_state['final_top_k'])

    variants = expand_queries(query)
    st.session_state['_last_variants'] = variants
    # variants[0] is the contextually-rewritten query (or the original if no
    # rewriting happened); used as the rerank target since it best captures intent.
    base_query = variants[0] if variants else query

    per_query_rankings = []
    for q in variants:
        ranking = _single_query_local_search(q, mode, top_n)
        if ranking:
            per_query_rankings.append(ranking)

    if not per_query_rankings:
        return []
    if len(per_query_rankings) == 1:
        fused = per_query_rankings[0][:top_n]
    else:
        fused = rrf_fuse(per_query_rankings)[:top_n]

    # Rerank operates on the full top_n candidates so we have a high-quality
    # ranked list to apply per-doc balancing on top of.
    if st.session_state['use_reranker']:
        try:
            fused = rerank(base_query, fused, top_k=len(fused))
        except Exception as e:
            st.warning(f'Reranker 실패 ({e}). RRF 결과를 그대로 사용합니다.')

    per_doc_min = _effective_per_doc_min(query)
    if per_doc_min > 0:
        fused = _select_with_per_doc_min(fused, top_k, per_doc_min)
    else:
        fused = fused[:top_k]

    docs_by_id = {d['id']: d for d in st.session_state['docs']}
    out = []
    for idx, score, meta, chunk in fused:
        doc_id, doc_name, chunk_idx = meta
        doc = docs_by_id.get(doc_id, {})
        pages_list = doc.get('chunk_pages') or []
        pages = pages_list[chunk_idx] if chunk_idx < len(pages_list) else []
        out.append({
            'source': 'doc',
            'doc_id': doc_id,
            'doc': doc_name,
            'chunk_idx': chunk_idx,
            'pages': pages,
            'text': chunk,
            'score': score,
        })
    return out


def web_search(query: str) -> list:
    """Return list of normalized web search results."""
    provider = st.session_state['web_provider']
    top_n = int(st.session_state['web_top_n'])
    try:
        if provider == 'duckduckgo':
            try:
                from ddgs import DDGS
            except ImportError:
                from duckduckgo_search import DDGS
            results = list(DDGS().text(query, max_results=top_n))
            out = []
            for r in results:
                out.append({
                    'source': 'web',
                    'doc': r.get('title') or r.get('href', ''),
                    'url': r.get('href') or r.get('link', ''),
                    'chunk_idx': 0,
                    'text': r.get('body') or r.get('snippet', ''),
                    'score': 0.0,
                })
            return out

        if provider == 'tavily':
            key = st.session_state['tavily_key']
            if not key:
                st.warning(
                    'Tavily API 키가 비어 있습니다. '
                    '.env의 TAVILY_API_KEY에 추가하거나 사이드바에서 입력하세요. '
                    '키 발급: https://app.tavily.com'
                )
                return []
            try:
                from tavily import TavilyClient
            except ImportError:
                st.error(
                    'tavily-python 패키지가 설치되지 않았습니다. '
                    '터미널에서 `pip install tavily-python`를 실행해 주세요.'
                )
                return []
            client = TavilyClient(api_key=key)
            resp = client.search(query=query, max_results=top_n, search_depth='basic')
            out = []
            for r in resp.get('results', []):
                out.append({
                    'source': 'web',
                    'doc': r.get('title', ''),
                    'url': r.get('url', ''),
                    'chunk_idx': 0,
                    'text': r.get('content', ''),
                    'score': float(r.get('score', 0.0)),
                })
            return out

        if provider == 'brave':
            key = st.session_state['brave_key']
            if not key:
                st.warning('Brave API 키가 비어 있습니다 (.env BRAVE_API_KEY).')
                return []
            import requests
            r = requests.get(
                'https://api.search.brave.com/res/v1/web/search',
                params={'q': query, 'count': top_n},
                headers={
                    'X-Subscription-Token': key,
                    'Accept': 'application/json',
                },
                timeout=15,
            )
            r.raise_for_status()
            data = r.json()
            out = []
            for item in data.get('web', {}).get('results', []):
                out.append({
                    'source': 'web',
                    'doc': item.get('title', ''),
                    'url': item.get('url', ''),
                    'chunk_idx': 0,
                    'text': item.get('description', ''),
                    'score': 0.0,
                })
            return out

    except Exception as e:
        st.warning(f'웹 검색 실패 ({provider}): {e}')
        return []
    return []


def retrieve(query: str) -> list:
    """Combine local document retrieval and web search results."""
    results = retrieve_local(query)
    if st.session_state['web_enabled']:
        results = results + web_search(query)
    return results


# =============================================================================
# Doc ingestion
# =============================================================================

def _embed_with_progress(chunks: list, embedder, status, label_prefix: str,
                          batch_size: int = 32):
    """Embed chunks in batches and surface per-batch progress to a Streamlit
    st.status container. Returns the stacked numpy embedding matrix."""
    total = len(chunks)
    if total == 0:
        import numpy as _np
        return _np.zeros((0, 1), dtype=_np.float32)
    if total <= batch_size:
        status.update(label=f'{label_prefix} 임베딩 ({total}청크)')
        return embedder.encode(
            chunks, convert_to_numpy=True, normalize_embeddings=True,
            show_progress_bar=False,
        )
    import numpy as _np
    parts = []
    done = 0
    bar = st.progress(0.0, text=f'{label_prefix} 임베딩 {done}/{total} 청크')
    for i in range(0, total, batch_size):
        batch = chunks[i:i + batch_size]
        parts.append(embedder.encode(
            batch, convert_to_numpy=True, normalize_embeddings=True,
            show_progress_bar=False,
        ))
        done = min(i + batch_size, total)
        bar.progress(done / total, text=f'{label_prefix} 임베딩 {done}/{total} 청크')
    bar.empty()
    return _np.vstack(parts)


def ingest_files(files):
    """For each new file: parse (Docling for PDF) → chunk with page metadata
    → embed (batched, progress-reported) → optionally render page images
    (only if multimodal is enabled, otherwise skipped to save time)
    → save → register."""
    eid = st.session_state['embedder_model']
    size = st.session_state['chunk_size']
    overlap = st.session_state['chunk_overlap']
    render_pages = bool(st.session_state.get('include_page_images'))
    existing_ids = {d['id'] for d in st.session_state['docs']}
    existing_names = {d['name'] for d in st.session_state['docs']}
    new_count = 0
    import time as _ing_time
    batch_t0 = _ing_time.time()

    with st.status(f'{len(files)}개 파일 인덱싱 시작', expanded=True) as status:
        for f in files:
            if f.name in existing_names:
                status.update(label=f'{f.name}: 이미 등록됨, 건너뜀')
                _log_event('doc_skip', {'name': f.name, 'reason': 'duplicate_name'})
                continue

            file_t0 = _ing_time.time()
            status.update(label=f'{f.name}: 파싱 중...')
            parsed = parse_file(f)
            raw = parsed['raw_text']
            if not raw.strip():
                st.warning(f'{f.name}: 추출된 텍스트가 없어 인덱스에서 제외합니다.')
                _log_event('doc_skip', {'name': f.name, 'reason': 'no_text'})
                continue

            status.update(label=f'{f.name}: 청크 분할 중...')
            chunks, chunk_pages = chunk_elements(parsed['elements'], size, overlap)
            if not chunks:
                st.warning(f'{f.name}: 청크가 생성되지 않았습니다.')
                _log_event('doc_skip', {'name': f.name, 'reason': 'no_chunks'})
                continue

            did = compute_doc_id(f.name, raw, size, overlap)
            if did in existing_ids:
                _log_event('doc_skip', {'name': f.name, 'reason': 'duplicate_content', 'doc_id': did})
                continue

            cached = load_doc(eid, did)
            if cached is not None:
                status.update(label=f'{f.name}: 캐시에서 즉시 복원 ({len(cached["chunks"])} 청크)')
                doc = cached
                # Cached restore: pgvector copy may not exist on this Cloud
                # container yet (different deploy / different region). Upsert
                # is idempotent so retrying is cheap.
                _pgvector_upsert_doc(eid, doc)
            else:
                # Load embedder lazily — the very first call also downloads the
                # model weights (~470 MB for MiniLM, ~2.2 GB for BGE-M3).
                status.update(label=f'{f.name}: 임베더 준비 중 (모델 로드)...')
                embedder = load_embedder(eid)
                embs = _embed_with_progress(
                    chunks, embedder, status, label_prefix=f.name,
                )

                has_imgs = False
                if render_pages and parsed['is_pdf'] and parsed['pdf_bytes']:
                    status.update(label=f'{f.name}: 페이지 이미지 렌더 중 (멀티모달용)...')
                    try:
                        n_pages = render_pdf_pages_to_dir(
                            parsed['pdf_bytes'], _pages_dir(eid, did)
                        )
                        has_imgs = n_pages > 0
                    except Exception as e:
                        st.warning(f'{f.name}: 페이지 이미지 렌더 실패 ({e}).')

                status.update(label=f'{f.name}: 디스크 저장 중...')
                doc = {
                    'id': did, 'name': f.name, 'raw_text': raw,
                    'chunks': chunks, 'chunk_pages': chunk_pages,
                    'page_count': parsed['page_count'],
                    'has_page_images': has_imgs,
                    'is_pdf': parsed['is_pdf'],
                    'embeddings': embs,
                    'chunk_size': size, 'chunk_overlap': overlap,
                }
                save_doc(eid, doc)
                # Also push chunk embeddings to pgvector if Supabase is
                # configured — no-op otherwise, never blocks ingestion.
                _pgvector_upsert_doc(eid, doc)

            st.session_state['docs'].append(doc)
            existing_ids.add(did)
            existing_names.add(doc['name'])
            new_count += 1
            status.update(label=f'{f.name}: 완료 ({len(doc["chunks"])} 청크)')
            _log_event('doc_ingest', {
                'name': doc['name'],
                'doc_id': did,
                'n_chunks': len(doc['chunks']),
                'page_count': doc.get('page_count', 0),
                'has_page_images': doc.get('has_page_images', False),
                'is_pdf': doc.get('is_pdf', False),
                'embedder': eid,
                'chunk_size': size,
                'chunk_overlap': overlap,
                'elapsed_seconds': round(_ing_time.time() - file_t0, 3),
                'from_cache': cached is not None,
            })

        if new_count > 0:
            status.update(
                label=f'{new_count}개 새 문서 인덱싱 완료',
                state='complete', expanded=False,
            )
        else:
            status.update(
                label='새로 인덱싱된 문서 없음', state='complete', expanded=False,
            )
    _log_event('doc_ingest_batch', {
        'files_offered': len(files),
        'new_count': new_count,
        'elapsed_seconds': round(_ing_time.time() - batch_t0, 3),
    })
    return new_count


def remove_doc(doc_id: str):
    eid = st.session_state['embedder_model']
    doc_name = next(
        (d['name'] for d in st.session_state['docs'] if d['id'] == doc_id),
        '',
    )
    delete_saved_doc(eid, doc_id)
    _pgvector_delete_doc(doc_id)
    st.session_state['docs'] = [d for d in st.session_state['docs'] if d['id'] != doc_id]
    _log_event('doc_delete', {
        'doc_id': doc_id, 'name': doc_name, 'embedder': eid,
    })


def reindex_all():
    """Re-chunk + re-embed all current docs with current settings.

    Note: re-chunking uses the stored raw_text, which loses Docling page
    boundaries. Page metadata is preserved only for original ingestion.
    """
    eid = st.session_state['embedder_model']
    size = st.session_state['chunk_size']
    overlap = st.session_state['chunk_overlap']
    new_docs = []
    embedder = load_embedder(eid)
    for d in st.session_state['docs']:
        raw = d.get('raw_text', '')
        if not raw:
            new_docs.append(d)
            continue
        chunks = chunk_text(raw, size, overlap)
        if not chunks:
            continue
        new_id = compute_doc_id(d['name'], raw, size, overlap)
        cached = load_doc(eid, new_id)
        if cached is not None:
            new_docs.append(cached)
            continue
        embs = embedder.encode(
            chunks, convert_to_numpy=True, normalize_embeddings=True,
            show_progress_bar=False,
        )
        nd = {
            'id': new_id, 'name': d['name'], 'raw_text': raw,
            'chunks': chunks,
            'chunk_pages': [[] for _ in chunks],
            'page_count': d.get('page_count', 0),
            'has_page_images': d.get('has_page_images', False),
            'is_pdf': d.get('is_pdf', False),
            'embeddings': embs,
            'chunk_size': size, 'chunk_overlap': overlap,
        }
        save_doc(eid, nd)
        new_docs.append(nd)
    st.session_state['docs'] = new_docs


# =============================================================================
# Chat helpers
# =============================================================================

RAG_SYSTEM_PROMPT = (
    "You are a helpful assistant. Ground every factual claim in the Context "
    "provided in the user's latest message. Also use the prior conversation "
    "to understand the user — resolve pronouns, follow-up references, and "
    "implied subjects from earlier turns. Each Context entry is marked with "
    "[N] and tagged as either a local document or a web result. Cite entry "
    "numbers like [1], [2] for every fact you take from the Context. If a "
    "factual answer is not supported by the Context or recent conversation, "
    "reply exactly: \"제공된 자료에서는 답을 찾을 수 없습니다.\""
)


def _format_pages(pages):
    if not pages:
        return ''
    if len(pages) == 1:
        return f' p.{pages[0]}'
    return f' pp.{pages[0]}-{pages[-1]}'


def _context_label(r: dict) -> str:
    if r.get('source') == 'web':
        url = r.get('url', '')
        return f'웹: {r.get("doc", "")} | {url}'
    pages = r.get('pages') or []
    page_part = _format_pages(pages)
    return f'로컬: {r.get("doc", "")} (chunk {r.get("chunk_idx", 0)}{page_part})'


def _collect_page_image_parts(retrieved: list, max_images: int) -> list:
    """Return a list of {type: image_url, image_url: {url: data:...}} parts
    for unique (doc_id, page) pairs that have rendered images on disk."""
    parts = []
    seen = set()
    embedder_id = st.session_state['embedder_model']
    for r in retrieved:
        if r.get('source') != 'doc':
            continue
        doc_id = r.get('doc_id')
        for p in (r.get('pages') or []):
            key = (doc_id, p)
            if key in seen:
                continue
            seen.add(key)
            b64 = load_page_image_b64(embedder_id, doc_id, p)
            if not b64:
                continue
            parts.append({
                'type': 'image_url',
                'image_url': {'url': f'data:image/png;base64,{b64}'},
            })
            if len(parts) >= max_images:
                return parts
    return parts


def build_messages(user_input: str, retrieved: list) -> list:
    msgs = []
    if retrieved:
        msgs.append({'role': 'system', 'content': RAG_SYSTEM_PROMPT})
    for u, a in zip(
        st.session_state['user_inputs'], st.session_state['generated_responses']
    ):
        msgs.append({'role': 'user', 'content': u})
        msgs.append({'role': 'assistant', 'content': a})

    if retrieved:
        ctx = '\n\n'.join(
            f'[{i + 1}] ({_context_label(r)}) {r["text"]}'
            for i, r in enumerate(retrieved)
        )
        text_content = f'Context:\n{ctx}\n\nQuestion: {user_input}'
    else:
        text_content = user_input

    # Multimodal: attach unique page images for current turn only.
    image_parts = []
    if retrieved and st.session_state.get('include_page_images'):
        max_imgs = int(st.session_state.get('max_page_images', 3))
        image_parts = _collect_page_image_parts(retrieved, max_imgs)

    if image_parts:
        content = [{'type': 'text', 'text': text_content}] + image_parts
        msgs.append({'role': 'user', 'content': content})
    else:
        msgs.append({'role': 'user', 'content': text_content})
    return msgs


def split_thinking(text: str):
    s = text or ''
    m = re.search(r'<think>(.*?)</think>\s*(.*)', s, flags=re.DOTALL)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    m = re.search(
        r'<\|channel\|?>thought\s*(.*?)<\|?channel\|>\s*(.*)',
        s, flags=re.DOTALL,
    )
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return '', s.strip()


_CITE_PATTERN = r'\[((?:\d+\s*,\s*)*\d+)\]'


def parse_citations(answer: str, n_chunks: int) -> set:
    nums = set()
    for group in re.findall(_CITE_PATTERN, answer or ''):
        for piece in group.split(','):
            try:
                n = int(piece.strip())
                if 1 <= n <= n_chunks:
                    nums.add(n)
            except ValueError:
                pass
    return nums


def format_answer_with_citations(answer: str, n_chunks: int) -> str:
    def repl(m):
        inner = m.group(1)
        for piece in inner.split(','):
            try:
                n = int(piece.strip())
                if 1 <= n <= n_chunks:
                    return f'**[{inner}]**'
            except ValueError:
                pass
        return m.group(0)
    return re.sub(_CITE_PATTERN, repl, answer or '')


# ---------- Agentic RAG (Phase 3) ----------
# The assistant gets one tool — search_documents(query) — and may invoke
# it up to N times to fetch additional context before producing the final
# answer. Each tool result is appended back into the message history so
# subsequent turns of the loop see all prior searches.

_SEARCH_TOOL_DEF = {
    'type': 'function',
    'function': {
        'name': 'search_documents',
        'description': (
            "사용자가 업로드한 문서에서 추가로 검색합니다. 초기 컨텍스트만으로 "
            "답하기 어려운 경우 더 구체적인 쿼리로 다시 검색해 근거를 보강하세요. "
            "다음 상황에 유용합니다: (1) 질문에 여러 하위 토픽이 있어 각각 따로 "
            "검색이 필요할 때, (2) 초기 컨텍스트에 핵심 정보가 빠져 보일 때, "
            "(3) 비교/대조 질문에서 한쪽 정보만 충분한 경우. 한 번에 한 가지에 "
            "집중된 쿼리를 작성하세요."
        ),
        'parameters': {
            'type': 'object',
            'properties': {
                'query': {
                    'type': 'string',
                    'description': '검색할 핵심 쿼리. 명확한 명사구 또는 self-contained 한 문장.',
                },
            },
            'required': ['query'],
        },
    },
}


def _format_tool_search_result(retrieved: list) -> str:
    """Compact representation of retrieved chunks for tool-message content."""
    if not retrieved:
        return '(검색 결과 없음)'
    lines = []
    for i, r in enumerate(retrieved, start=1):
        pages = r.get('pages') or []
        if len(pages) == 1:
            page_str = f' p.{pages[0]}'
        elif pages:
            page_str = f' pp.{pages[0]}-{pages[-1]}'
        else:
            page_str = ''
        src = r.get('doc', '') or r.get('url', '')
        lines.append(f"[검색 결과 {i}] ({src}{page_str}) {r.get('text', '')}")
    return '\n\n'.join(lines)


def agentic_chat_pass(client, params: dict, initial_retrieved: list,
                      max_iters: int = 3):
    """Run a chat completion loop where the assistant may invoke
    search_documents to refine context. Always non-streaming (tool-call
    streaming is brittle across providers). Returns
    (full_text, reasoning_text, augmented_retrieved).

    Falls back to a plain non-stream call without `tools` if the provider
    rejects tool calling (e.g. some HF Router-served small models)."""
    augmented = list(initial_retrieved)
    seen_keys = {(r.get('doc_id'), r.get('chunk_idx'))
                 for r in augmented if r.get('source') != 'web'}

    tools_params = {k: v for k, v in params.items() if k != 'stream'}
    tools_params['tools'] = [_SEARCH_TOOL_DEF]
    messages = list(tools_params['messages'])

    reasoning_acc = ''
    n_tool_calls_made = 0

    for iteration in range(max_iters):
        try:
            tools_params['messages'] = messages
            spinner_msg = (
                '생각 중…' if iteration == 0
                else f'추가 검색 결과 분석 중… (라운드 {iteration})'
            )
            with st.spinner(spinner_msg):
                resp = client.chat.completions.create(**tools_params)
        except Exception:
            # Provider/model rejects tools — fall back to plain non-stream.
            fallback = {k: v for k, v in params.items() if k not in ('stream',)}
            fallback['messages'] = messages
            with st.spinner('도구 호출 미지원 — 일반 모드로 재시도...'):
                resp = client.chat.completions.create(**fallback)
            msg = resp.choices[0].message
            full = msg.content or ''
            r = getattr(msg, 'reasoning_content', '') or ''
            return full, (reasoning_acc + r).strip(), augmented

        msg = resp.choices[0].message
        tool_calls = getattr(msg, 'tool_calls', None) or []
        r_part = getattr(msg, 'reasoning_content', '') or ''
        if r_part:
            reasoning_acc = (reasoning_acc + '\n\n' + r_part).strip()

        if not tool_calls:
            # Final answer — no more searches requested.
            return msg.content or '', reasoning_acc, augmented

        # Surface the search activity in the UI so the user sees the loop.
        st.info(f'추가 검색 라운드 {iteration + 1}/{max_iters}')

        # Append assistant turn with tool_calls intact (required by API).
        messages.append({
            'role': 'assistant',
            'content': msg.content or '',
            'tool_calls': [
                {
                    'id': tc.id,
                    'type': 'function',
                    'function': {
                        'name': tc.function.name,
                        'arguments': tc.function.arguments,
                    },
                }
                for tc in tool_calls
            ],
        })

        for tc in tool_calls:
            if tc.function.name != 'search_documents':
                tool_text = f'(알 수 없는 도구 호출: {tc.function.name})'
            else:
                try:
                    args = json.loads(tc.function.arguments or '{}')
                except Exception:
                    args = {}
                q = (args.get('query') or '').strip()
                if not q:
                    tool_text = '(빈 쿼리)'
                else:
                    st.caption(f'→ "{q}"')
                    n_tool_calls_made += 1
                    new_chunks = retrieve_local(q)
                    tool_text = _format_tool_search_result(new_chunks)
                    for nc in new_chunks:
                        key = (nc.get('doc_id'), nc.get('chunk_idx'))
                        if key not in seen_keys:
                            augmented.append(nc)
                            seen_keys.add(key)

            messages.append({
                'role': 'tool',
                'tool_call_id': tc.id,
                'content': tool_text,
            })

    # Max iterations reached — make one final tool-less call to extract the
    # answer so we never return mid-loop.
    final_params = {k: v for k, v in params.items() if k != 'stream'}
    final_params['messages'] = messages
    with st.spinner('최종 답변 생성 중...'):
        resp = client.chat.completions.create(**final_params)
    msg = resp.choices[0].message
    r_part = getattr(msg, 'reasoning_content', '') or ''
    if r_part:
        reasoning_acc = (reasoning_acc + '\n\n' + r_part).strip()
    return msg.content or '', reasoning_acc, augmented


def _record_response_model(model_id: str) -> None:
    """Capture the model id the provider actually served the response from.
    Useful for verifying that requests went to the model the user picked —
    OpenAI chat completions often respond with a dated variant like
    'gpt-5-2025-08-07', so seeing that next to the configured 'gpt-5'
    confirms the routing without trusting the model's self-introduction."""
    if model_id:
        st.session_state['_last_response_model'] = str(model_id)


def stream_chat(client, params: dict):
    """Stream response. Returns (full_text, reasoning_text).

    Renders an immediate "생각 중…" placeholder so the user can tell the
    difference between "model is thinking" and "app is hung". The
    placeholder updates each second with an elapsed-time counter while
    we wait for the first visible token (reasoning models like
    HyperCLOVAX-Think can take 10s+ before they emit anything)."""
    import time as _time
    placeholder = st.empty()
    placeholder.markdown('_생각 중…_')
    t0 = _time.time()
    last_idle_tick = t0
    full_text = ''
    reasoning_text = ''
    try:
        stream = client.chat.completions.create(stream=True, **params)
        for chunk in stream:
            # Some providers stamp the model id on every chunk; capture
            # it once we see it.
            cm = getattr(chunk, 'model', None)
            if cm and not st.session_state.get('_stream_model_captured'):
                _record_response_model(cm)
                st.session_state['_stream_model_captured'] = True
            if not getattr(chunk, 'choices', None):
                continue
            delta = chunk.choices[0].delta
            rc = getattr(delta, 'reasoning_content', None) or ''
            c = getattr(delta, 'content', None) or ''
            if rc:
                reasoning_text += rc
            if c:
                full_text += c
            if full_text:
                placeholder.markdown(full_text)
            elif reasoning_text:
                elapsed = int(_time.time() - t0)
                suffix = f' ({elapsed}초 경과)' if elapsed >= 2 else ''
                placeholder.markdown(
                    f'_생각 중{suffix}_\n\n> {reasoning_text}'
                )
            else:
                # No visible output yet — refresh the elapsed counter so
                # users can tell the connection isn't frozen.
                now = _time.time()
                if now - last_idle_tick >= 1.0:
                    placeholder.markdown(
                        f'_생각 중… ({int(now - t0)}초 경과)_'
                    )
                    last_idle_tick = now
        placeholder.empty()
    except Exception:
        placeholder.empty()
        raise
    finally:
        st.session_state.pop('_stream_model_captured', None)
    return full_text, reasoning_text


def non_stream_chat(client, params: dict):
    resp = client.chat.completions.create(**params)
    _record_response_model(getattr(resp, 'model', None))
    choice = resp.choices[0]
    full_text = choice.message.content or ''
    reasoning = ''
    for attr in ('reasoning_content', 'reasoning'):
        val = getattr(choice.message, attr, None)
        if val:
            reasoning = val
            break
    return full_text, reasoning


def render_assistant(answer: str, reasoning: str, retrieved: list, turn_idx: int,
                     variants: list = None):
    if variants and len(variants) > 1:
        with st.expander(
            f'검색에 사용된 쿼리 변형 {len(variants)}개', expanded=False
        ):
            for i, v in enumerate(variants):
                label = '원본 질문' if i == 0 else f'변형 {i}'
                preview = v if len(v) <= 300 else v[:300] + '...'
                st.markdown(f'**{label}**: {preview}')
    if reasoning:
        with st.expander(f'추론 과정 (turn {turn_idx + 1})', expanded=False):
            st.markdown(reasoning)

    n = len(retrieved)
    pretty = format_answer_with_citations(answer, n) if n else (answer or '')
    st.markdown(pretty or '*(empty response)*')

    if not retrieved:
        return

    cited = parse_citations(answer, n)
    cited_items = [(i + 1, r) for i, r in enumerate(retrieved) if (i + 1) in cited]
    uncited_items = [(i + 1, r) for i, r in enumerate(retrieved) if (i + 1) not in cited]

    if cited_items:
        st.markdown(f"**출처** · 인용 {len(cited_items)} / 검색 {n}")
        for j, r in cited_items:
            with st.expander(_citation_summary(j, r), expanded=False):
                _citation_body(r)
        if uncited_items:
            with st.expander(
                f'인용되지 않은 검색 결과 {len(uncited_items)}개 보기', expanded=False
            ):
                for j, r in uncited_items:
                    st.markdown(f"**{_citation_summary(j, r)}**")
                    _citation_body(r)
                    st.divider()
    else:
        with st.expander(
            f'검색된 자료 {n}개 (모델이 [N] 인용 표기를 사용하지 않음)', expanded=False
        ):
            for i, r in enumerate(retrieved):
                st.markdown(f"**{_citation_summary(i + 1, r)}**")
                _citation_body(r)
                st.divider()


def _citation_summary(j: int, r: dict) -> str:
    from urllib.parse import urlparse
    if r.get('source') == 'web':
        host = urlparse(r.get('url', '')).netloc or '웹'
        title = (r.get('doc') or '').strip()
        return f"[{j}] 웹 · {host} — {title[:60]}"
    score = r.get('score', 0.0)
    pages = r.get('pages') or []
    page_part = _format_pages(pages)
    return (
        f"[{j}] {r.get('doc', '')}{page_part} · "
        f"chunk {r.get('chunk_idx', 0)} · score {score:.3f}"
    )


def _citation_body(r: dict):
    if r.get('source') == 'web':
        url = r.get('url', '')
        if url:
            st.markdown(f"[{url}]({url})")
    st.text(r.get('text', ''))


# =============================================================================
# Boot: load persisted docs for current embedder
# =============================================================================

load_all_for_current_embedder()


# =============================================================================
# UI helpers
# =============================================================================

def _chip(text: str, kind: str = 'default') -> str:
    cls = 'chip'
    if kind == 'active':
        cls += ' active'
    elif kind == 'muted':
        cls += ' muted'
    return f'<span class="{cls}">{text}</span>'


def _section(title: str, sub: str = ''):
    st.markdown(f'<div class="section-title">{title}</div>', unsafe_allow_html=True)
    if sub:
        st.markdown(f'<div class="section-sub">{sub}</div>', unsafe_allow_html=True)


def _empty(text: str):
    st.markdown(f'<div class="empty-state">{text}</div>', unsafe_allow_html=True)


def model_picker(label: str, key_prefix: str, instant: bool = False):
    """Per-provider model dropdown with '직접 입력' fallback.

    Two modes:
      - instant=False (default): the choice is *pending* until the user
        clicks the 적용 button. Used in the Settings tab where casual
        clicks on the dropdown shouldn't immediately change the model
        being sent to the API.
      - instant=True: every dropdown change writes straight to
        st.session_state['model']. Used in the chat top-bar quick switch
        where the user clearly wants to apply right away.
    """
    provider = st.session_state.get('provider', PROVIDER_NAMES[0])
    known = PROVIDER_MODELS.get(provider, [])
    current = st.session_state.get('model', '')

    def _commit(value: str):
        if not value or value == current:
            return False
        st.session_state['model'] = value
        try:
            _save_user_prefs()
        except Exception:
            pass
        return True

    if not known:
        if instant:
            new_val = st.text_input(
                label, value=current, key=f'{key_prefix}_model_text',
            )
            if _commit(new_val):
                st.rerun()
            return
        pending = st.text_input(
            label, value=current, key=f'{key_prefix}_model_text',
        )
    else:
        options = known + [_CUSTOM]
        if current in known:
            initial_idx = known.index(current)
        else:
            initial_idx = len(options) - 1   # 직접 입력
        choice = st.selectbox(
            label, options, index=initial_idx,
            format_func=lambda x: '직접 입력...' if x == _CUSTOM else x,
            key=f'{key_prefix}_model_select',
        )
        if instant and choice != _CUSTOM:
            # Quick-switch: any direct selection from the list applies now.
            if _commit(choice):
                st.rerun()
            return
        if choice == _CUSTOM:
            text_val = st.text_input(
                '모델 ID 직접 입력',
                value=current if current not in known else '',
                key=f'{key_prefix}_model_custom',
                placeholder='예: gpt-4o, my-org/my-finetune',
            )
            if instant:
                if _commit(text_val):
                    st.rerun()
                return
            pending = text_val
        else:
            pending = choice

    # ---- pending mode (settings tab) ----
    # Apply / confirm row. Only shown when the user has actually changed
    # the dropdown (or typed a different model id) from the active value.
    if pending and pending != current:
        cols = st.columns([3, 1])
        with cols[0]:
            st.caption(
                f'미적용 변경: `{current or "—"}` → `{pending}` '
                f'· 적용 버튼을 눌러야 채팅에 반영됩니다.'
            )
        with cols[1]:
            if st.button('적용', key=f'{key_prefix}_model_apply',
                         type='primary', use_container_width=True):
                _commit(pending)
                st.rerun()
    elif current:
        st.caption(f'현재 모델: `{current}`')


def _show_llm_error(e: Exception):
    """Render a Korean, actionable error message for common LLM call failures.

    Pattern-matches the most frequent ones (HF Inference Providers permission
    missing, gated model, invalid token, quota, model not deployed, network)
    and falls back to the raw text otherwise."""
    err_str = str(e) or ''
    el = err_str.lower()

    # Persist the error to events.jsonl so failures aren't only visible
    # in the live UI. Best-effort — never raises.
    try:
        _log_event('llm_error', {
            'provider': st.session_state.get('provider', ''),
            'model': st.session_state.get('model', ''),
            'base_url': st.session_state.get('base_url', ''),
            'exception_type': type(e).__name__,
            'error': err_str[:1500],
        })
    except Exception:
        pass

    def show(headline, body_md):
        st.error(headline)
        st.markdown(body_md)
        with st.expander('원본 오류 메시지'):
            st.code(err_str[:1500] or repr(e))

    # HF Router — Inference Providers permission missing (403)
    if ('inference providers' in el
            and ('insufficient permissions' in el or 'does not have' in el
                 or 'this authentication method' in el)):
        show(
            'Hugging Face 토큰에 "Inference Providers" 권한이 없습니다.',
            """
**해결 방법:**

1. https://huggingface.co/settings/tokens 접속
2. 사용 중인 토큰 이름 클릭 → **Edit permissions** (또는 + Create new token → Fine-grained)
3. 다음 권한 체크:
   - ✅ **Make calls to Inference Providers** ← 필수
   - ✅ Make calls to the serverless Inference API (권장)
   - ✅ Read access to public repositories (자동 포함)
   - Llama / Gemma 같은 gated 모델 쓰면 → Read access to selected gated repositories 추가
4. **Save** → 설정 탭에서 Hugging Face 토큰 칸 갱신 (또는 .env / Cloud secrets 의 HF_TOKEN 갱신)
5. (Cloud 배포면) Manage app → Reboot
            """,
        )
        return

    # HF — gated model (Llama, Gemma 등) 라이선스 미수락 / gated 권한 누락
    if ('gated' in el or
            ('access to model' in el and ('granted' in el or 'requires' in el)) or
            ('is restricted' in el and 'license' in el)):
        show(
            'Gated 모델 접근 권한 없음.',
            """
**해결 방법:**

1. 해당 모델 페이지 (예: https://huggingface.co/meta-llama/Llama-3.1-8B-Instruct) 에서 라이선스 약관 수락
2. https://huggingface.co/settings/tokens → 토큰 권한에 **Read access to selected gated repositories** 추가하고 해당 모델 체크
3. 토큰 저장 → 앱에서 다시 시도
            """,
        )
        return

    # Fireworks AI — 모델 ID 오타 또는 계정 권한 부족
    if (_is_fireworks_endpoint() and
            ('not found' in el or 'inaccessible' in el or 'not deployed' in el)):
        show(
            'Fireworks 에서 이 모델을 찾을 수 없거나 계정에 접근 권한이 없습니다.',
            """
**해결 방법:**

1. **모델 ID 오타 확인** — Fireworks 모델 ID 는 항상 `accounts/fireworks/models/<name>` 풀 경로. 짧은 이름만 적으면 안 됩니다.
2. **계정 등급 제한** — Llama 405B 같은 일부 대형 모델은 유료 등급 / 신청 필요. 무료 계정에서는 70B 이하 권장.
3. **권장 동작 모델:** 설정 → Fireworks AI 모델 드롭다운에서:
   - `accounts/fireworks/models/llama-v3p3-70b-instruct` (기본)
   - `accounts/fireworks/models/qwen2p5-72b-instruct` (한국어 우수)
   - `accounts/fireworks/models/deepseek-v3` (강력)
4. 현재 본인 계정에서 호출 가능한 모델 전체 목록: https://fireworks.ai/models
            """,
        )
        return

    # Provider 미지원 / 모델 deploy 안 됨
    if 'not supported by any provider' in el or 'model_not_supported' in el:
        show(
            '이 모델을 서빙하는 활성 Provider가 없습니다.',
            """
**해결 방법:**

1. 설정 탭에서 다른 모델로 변경 (한국어가 강한 추천: `Qwen/Qwen3-Next-80B-A3B-Instruct`, `meta-llama/Llama-3.3-70B-Instruct`, `deepseek-ai/DeepSeek-V4-Pro`)
2. 또는 모델 카드 우측 "Inference Providers" 박스에서 서빙 가능한 provider 확인 후 https://huggingface.co/settings/inference-providers 에서 활성화 (Together AI / Cerebras / Hyperbolic 등)
            """,
        )
        return

    # OpenAI — Responses API 전용 모델 (gpt-5-pro, o1-pro 등) → Chat Completions 거부
    if 'v1/responses' in el or 'only supported in v1/responses' in el:
        show(
            '이 모델은 OpenAI 의 Responses API 전용 — 우리 앱은 호환되지 않습니다.',
            """
**원인:** `gpt-5-pro`, `o1-pro` 같은 "Pro" 등급 일부 모델은 OpenAI 의 새 `/v1/responses` 엔드포인트로만 제공됩니다. 본 앱은 표준 `/v1/chat/completions` 를 사용해 호출 자체가 거부됩니다.

**해결: 설정 → OpenAI 모델 변경.** 다음은 Chat Completions 로 정상 동작합니다:

- `gpt-5` — Pro 의 약 7할 성능, 같은 추론
- `gpt-5-mini` — 빠르고 저렴, 일반 RAG 충분
- `gpt-4.1` — 안정적인 차세대
- `o3` — 추론 강화 (논문 분석 / 수학에 강함)
- `o4-mini` — 추론 + 빠름
            """,
        )
        return

    # 토큰 자체가 무효 / 만료
    if ('invalid' in el and 'token' in el) or 'bad credentials' in el or '401' in el:
        show(
            'API 키 / 토큰 인증 실패 (401).',
            """
**해결 방법:**

1. 설정 탭에서 현재 공급자의 API 키가 비어있지 않은지 확인
2. 토큰이 만료됐다면 발급처에서 새로 만들기:
   - Hugging Face: https://huggingface.co/settings/tokens
   - OpenAI: https://platform.openai.com/api-keys
3. 새 키로 갱신 → 다시 시도
            """,
        )
        return

    # 결제 / 쿼터 초과
    if ('quota' in el or 'rate limit' in el or 'insufficient_quota' in el
            or '429' in err_str or '402' in err_str):
        show(
            '사용량 / 쿼터 초과 또는 결제 필요.',
            """
**해결 방법:**

- OpenAI: https://platform.openai.com/account/billing 에서 결제 / 한도 확인
- HF Inference Providers: provider 별로 무료 크레딧 한도 다름. 다른 provider 활성화 시도.
- 잠시 후 재시도하거나 더 작은 모델로 변경.
            """,
        )
        return

    # 네트워크 / 타임아웃
    if 'timeout' in el or 'timed out' in el or 'connection' in el:
        show(
            '네트워크 / 타임아웃.',
            '잠시 후 다시 시도해 주세요. 문제가 지속되면 다른 공급자로 변경.',
        )
        return

    # Fallback
    st.error(f'요청 실패: {e}')


def handle_chat_turn(user_input: str):
    """Process one user turn: retrieve, call LLM, render, persist to history."""
    if (not st.session_state['model']
            or not _active_api_key()
            or not st.session_state['base_url']):
        provider = st.session_state.get('provider', 'Hugging Face Router')
        st.error(
            f'설정 탭에서 모델 · 엔드포인트 · {provider} 용 API 키를 모두 입력해 주세요.'
        )
        return

    general_only = bool(st.session_state.get('general_chat_mode'))
    local_active = (not general_only) and bool(st.session_state['docs'])
    web_active = (not general_only) and bool(st.session_state['web_enabled'])
    rag_active = local_active or web_active

    with st.chat_message('user'):
        st.markdown(user_input)

    with st.chat_message('assistant'):
        retrieved = []
        if rag_active:
            if local_active and web_active:
                msg = '문서 및 웹 검색 중...'
            elif web_active:
                msg = '웹 검색 중...'
            else:
                msg = '문서 검색 중...'
            # Apply chat-time document filter (None = use all)
            all_docs_full = st.session_state['docs']
            chat_filter = st.session_state.get('chat_doc_filter')
            need_filter = (
                chat_filter is not None
                and 0 < len(chat_filter) < len(all_docs_full)
            )
            if need_filter:
                filt_set = set(chat_filter)
                st.session_state['docs'] = [
                    d for d in all_docs_full if d['id'] in filt_set
                ]
            try:
                with st.spinner(msg):
                    retrieved = retrieve(user_input)
            finally:
                if need_filter:
                    st.session_state['docs'] = all_docs_full

        messages = build_messages(user_input, retrieved)

        is_openai = _is_openai_endpoint()
        is_dashscope = _is_dashscope_endpoint()
        extra_body = {}
        if _provider_supports_top_k():
            extra_body['top_k'] = int(st.session_state['sampling_top_k'])
        if not st.session_state['enable_thinking']:
            if is_dashscope:
                extra_body['enable_thinking'] = False
            elif not is_openai and not _is_hf_router_endpoint():
                # vLLM / SGLang / self-hosted only. HF Router providers vary
                # and several (e.g. Cerebras-served Llama / gpt-oss) reject
                # chat_template_kwargs with HTTP 400.
                extra_body['chat_template_kwargs'] = {'enable_thinking': False}

        params = _build_completion_params(
            model=st.session_state['model'],
            messages=messages,
            max_tokens=st.session_state['max_tokens'],
            temperature=st.session_state['temperature'],
            top_p=st.session_state['top_p'],
            presence_penalty=st.session_state['presence_penalty'],
            extra_body=extra_body,
        )

        full_text, reasoning_text = '', ''
        elapsed_sec = None
        try:
            import time as _time
            client = get_openai_client()
            _t0 = _time.time()
            # Agentic mode: only meaningful when we have local docs to search
            # against. Web-only retrieval would need a different tool.
            if st.session_state.get('use_agentic_search') and local_active:
                max_iters = int(st.session_state.get('agentic_max_iters', 3))
                full_text, reasoning_text, retrieved = agentic_chat_pass(
                    client, params, retrieved, max_iters=max_iters,
                )
            elif st.session_state['stream']:
                full_text, reasoning_text = stream_chat(client, params)
            else:
                with st.spinner('생각 중…'):
                    full_text, reasoning_text = non_stream_chat(client, params)
            elapsed_sec = _time.time() - _t0
        except Exception as e:
            _show_llm_error(e)
            full_text, reasoning_text = '', ''

        if full_text and not reasoning_text:
            rt, ct = split_thinking(full_text)
            if rt:
                reasoning_text = rt
                full_text = ct

        if full_text or reasoning_text:
            turn_variants = st.session_state.get('_last_variants') or [user_input]
            st.session_state['user_inputs'].append(user_input)
            st.session_state['generated_responses'].append(full_text)
            st.session_state['thinking_traces'].append(reasoning_text)
            st.session_state['retrieved_per_turn'].append(retrieved)
            st.session_state['query_variants_per_turn'].append(turn_variants)
            render_assistant(
                full_text, reasoning_text, retrieved,
                len(st.session_state['generated_responses']) - 1,
                variants=turn_variants,
            )
            # Persist this conversation. Create a session ID on first turn.
            if not st.session_state.get('current_session_id'):
                st.session_state['current_session_id'] = _new_session_id()
                st.session_state['current_session_created_at'] = (
                    datetime.datetime.now().isoformat()
                )
            save_current_session()
            # Auto-title once, after the very first turn.
            if (len(st.session_state['user_inputs']) == 1
                    and not st.session_state.get('current_session_title')):
                title = auto_title_session()
                if title:
                    st.session_state['current_session_title'] = title
                    save_current_session()
            # Append this turn to the per-session structured JSONL log.
            log_turn_structured(
                user_input=user_input,
                response_text=full_text,
                reasoning=reasoning_text,
                retrieved=retrieved,
                model=st.session_state['model'],
                elapsed_sec=elapsed_sec,
                query_variants=turn_variants,
            )


# =============================================================================
# Sidebar — brand, new chat, navigation, current status
# =============================================================================

NAV = [
    ('chat',     '대화'),
    ('docs',     '문서'),
    ('agents',   '업무 도구(에이전트)'),
    ('settings', '설정'),
    ('cache',    '캐시'),
    ('about',    '소개'),
]
NAV_KEYS = [k for k, _ in NAV]

with st.sidebar:
    if _LOGO_URI:
        st.markdown(
            f'<div style="text-align:center; padding:4px 0 4px 0;">'
            f'<img src="{_LOGO_URI}" '
            f'style="width:100%; max-width:220px; height:auto; display:block; margin:0 auto;" />'
            f'</div>'
            f'<div style="text-align:center; margin-bottom:8px;">'
            f'<div class="sb-brand" style="font-size:14px;">Personal RAG</div>'
            f'<div class="sb-tagline">내 문서 기반 질의응답</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="sb-brand">Personal RAG</div>', unsafe_allow_html=True)
        st.markdown('<div class="sb-tagline">내 문서 기반 질의응답</div>', unsafe_allow_html=True)
    st.write('')

    if st.button('새 대화', use_container_width=True, type='primary'):
        start_new_session()
        st.session_state['active_view'] = 'chat'
        st.rerun()

    # ----- Menu (pinned, always visible above the fold) -----
    st.markdown('<div class="sb-section">메뉴</div>', unsafe_allow_html=True)
    current_view = st.session_state.get('active_view', 'chat')
    if current_view not in NAV_KEYS:
        current_view = 'chat'
    for key, label in NAV:
        is_active = (key == current_view)
        if st.button(
            label,
            key=f'nav_{key}',
            use_container_width=True,
            type='primary' if is_active else 'secondary',
        ):
            if not is_active:
                st.session_state['active_view'] = key
                st.rerun()

    # ----- Saved conversations (scrolls internally; does not push menu off) -----
    sessions = list_sessions()
    current_sid = st.session_state.get('current_session_id')
    st.markdown('<div class="sb-section">대화</div>', unsafe_allow_html=True)
    if not sessions:
        st.caption('저장된 대화가 없습니다.')
    else:
        # Cap visible height; the list scrolls inside its own container so the
        # menu above and status below stay anchored.
        with st.container(height=300, border=False):
            for s in sessions:
                is_active = (s['id'] == current_sid)
                label = s['title'] or '(제목 없음)'
                if len(label) > 22:
                    label = label[:20] + '...'
                row = st.columns([5, 1, 1])
                with row[0]:
                    if st.button(
                        label,
                        key=f"sess_{s['id']}",
                        use_container_width=True,
                        type='primary' if is_active else 'secondary',
                    ):
                        if not is_active:
                            load_session(s['id'])
                            st.session_state['active_view'] = 'chat'
                            st.rerun()
                with row[1]:
                    if hasattr(st, 'popover'):
                        with st.popover(
                            '✎', use_container_width=True,
                            help='대화 이름 변경',
                        ):
                            with st.form(
                                f'_rename_form_{s["id"]}',
                                clear_on_submit=False,
                            ):
                                new_t = st.text_input(
                                    '대화 이름',
                                    value=s['title'] or '',
                                    max_chars=60,
                                    key=f'rename_input_{s["id"]}',
                                )
                                if st.form_submit_button(
                                    '저장',
                                    use_container_width=True,
                                    type='primary',
                                ):
                                    if rename_session(s['id'], new_t):
                                        st.rerun()
                with row[2]:
                    if st.button('×', key=f"del_sess_{s['id']}",
                                 use_container_width=True,
                                 help='대화 삭제'):
                        delete_session(s['id'])
                        if current_sid == s['id']:
                            start_new_session()
                        st.rerun()

    # ----- Status -----
    st.markdown('<div class="sb-section">현재 상태</div>', unsafe_allow_html=True)
    n_docs = len(st.session_state['docs'])
    n_chunks = sum(len(d['chunks']) for d in st.session_state['docs'])
    model_short = st.session_state['model'] or '없음'
    if len(model_short) > 32:
        model_short = model_short[:29] + '...'
    st.caption(f"모델: `{model_short}`")
    prov_short = st.session_state.get('provider', '')
    if prov_short:
        st.caption(f"공급자: `{prov_short}`")
    # Provider-stamped model id from the last API response — bulletproof
    # confirmation of which model actually served the previous turn. Shown
    # always (not only on mismatch) because LLMs misidentify themselves
    # ("I'm Claude 3" when running Claude 4.6) and this header settles it.
    actual_model = st.session_state.get('_last_response_model')
    if actual_model:
        am_short = actual_model if len(actual_model) <= 36 else actual_model[:33] + '...'
        configured = st.session_state.get('model', '')
        if actual_model == configured:
            st.caption(f"직전 응답 (실제): `{am_short}` ✓ 일치")
        else:
            st.caption(f"직전 응답 (실제): `{am_short}`")
    if st.session_state.get('general_chat_mode'):
        st.caption(
            '모드: 일반 대화 (RAG 끔)'
            + (f' · 문서 {n_docs}개 보유' if n_docs else '')
        )
    elif n_docs:
        st.caption(f"문서 {n_docs}개 · 청크 {n_chunks}개")
    else:
        st.caption('문서 없음 (일반 챗)')
    if st.session_state['web_enabled'] and not st.session_state.get('general_chat_mode'):
        st.caption(f"웹 검색: {st.session_state['web_provider']}")

    # ----- User / logout -----
    uid = st.session_state.get('user_id', '_local')
    # Logged-in state = anything other than the anonymous/local fallbacks.
    is_logged_in = not (uid == '_local' or uid.startswith('_anon_'))
    st.markdown('<div class="sb-section">사용자</div>', unsafe_allow_html=True)
    if is_logged_in:
        st.caption(f"로그인: `{uid}`")
        if st.button('로그아웃', use_container_width=True, key='logout_btn'):
            _log_event('logout', {'username': uid})
            base_clear = (
                'user_id', 'user_inputs', 'generated_responses',
                'thinking_traces', 'retrieved_per_turn',
                'query_variants_per_turn', 'current_session_id',
                'current_session_title', 'current_session_created_at',
                'docs', 'doc_embs', 'doc_meta',
                '_loaded_for_embedder', '_prefs_snapshot', '_prefs_loaded',
            )
            # Also clear all persistable preference keys so the next user does
            # not inherit the previous one's API keys / model / settings.
            for k in base_clear + _PERSIST_KEYS:
                if k in st.session_state:
                    del st.session_state[k]
            st.rerun()
    elif uid.startswith('_anon_'):
        st.caption(f"익명 세션: `{uid}`")
        st.caption('이 브라우저 탭에서만 유효. 새 탭/창은 별도 데이터.')
    else:
        st.caption(f"로컬 단일 사용자: `{uid}`")


# =============================================================================
# View renderers
# =============================================================================

def view_chat():
    # Auto-restore the last active conversation if session_state was wiped
    # by an idle reconnect. preferences.json carries current_session_id; the
    # actual messages live in .data/{user}/sessions/{id}.json on disk and
    # are pulled in here by load_session().
    _restore_sid = st.session_state.get('current_session_id')
    if _restore_sid and not st.session_state.get('user_inputs'):
        try:
            load_session(_restore_sid)
        except Exception:
            pass

    local_active = bool(st.session_state['docs'])
    web_active = bool(st.session_state['web_enabled'])

    # Top bar — title + model picker + (optional) doc filter.
    has_multi_docs = len(st.session_state['docs']) >= 2
    if has_multi_docs:
        top_left, top_mid, top_right = st.columns([3, 1, 2])
    else:
        top_left, top_right = st.columns([3, 2])
        top_mid = None

    with top_left:
        title = st.session_state.get('current_session_title') or '새 대화'
        current_sid_for_rename = st.session_state.get('current_session_id')
        # Title behaves as a popover trigger when there's an actual saved
        # session — gives the user an inline rename UI. Pre-session it's
        # just static text since there's nothing to persist yet.
        if current_sid_for_rename and hasattr(st, 'popover'):
            with st.popover(title, use_container_width=False):
                with st.form('_rename_form', clear_on_submit=False):
                    new_title = st.text_input(
                        '대화 이름',
                        value=(st.session_state.get('current_session_title') or ''),
                        max_chars=60,
                        placeholder='예: 회의록 요약 / 신입사원 매뉴얼 Q&A',
                    )
                    submit = st.form_submit_button(
                        '저장', use_container_width=True, type='primary',
                    )
                if submit:
                    if rename_session(current_sid_for_rename, new_title):
                        st.success('이름을 변경했습니다.')
                        st.rerun()
                    else:
                        st.error('이름 변경에 실패했습니다.')
        else:
            st.markdown(
                f"<div style='font-size:15px; font-weight:600; padding-top:4px;'>{title}</div>",
                unsafe_allow_html=True,
            )

    # Doc filter popover (only when multiple docs are loaded).
    if has_multi_docs and top_mid is not None and hasattr(st, 'popover'):
        all_docs = st.session_state['docs']
        doc_id_to_name = {d['id']: d['name'] for d in all_docs}
        all_ids = list(doc_id_to_name.keys())
        all_id_set = set(all_ids)
        chat_filter = st.session_state.get('chat_doc_filter')
        n_total = len(all_ids)
        n_active = (
            len(chat_filter) if chat_filter is not None else n_total
        )
        with top_mid:
            with st.popover(f"문서 {n_active}/{n_total}",
                            use_container_width=True):
                # Default = filter if set (and still valid), else all
                if chat_filter is not None:
                    default_value = [d for d in chat_filter if d in all_id_set]
                else:
                    default_value = all_ids
                sel = st.multiselect(
                    '답변에 사용할 문서',
                    options=all_ids,
                    default=default_value,
                    format_func=lambda did: doc_id_to_name[did],
                    help='기본은 전체 문서. 일부만 사용하려면 체크를 조정하세요. '
                    '대화 중에도 언제든 변경할 수 있고, 다음 메시지부터 적용됩니다.',
                )
                new_filter = None if set(sel) == all_id_set else sel
                if new_filter != chat_filter:
                    st.session_state['chat_doc_filter'] = new_filter
                    st.rerun()

    with top_right:
        model_label = st.session_state['model'] or '모델 미설정'
        if len(model_label) > 36:
            model_label = model_label[:33] + '...'
        if hasattr(st, 'popover'):
            with st.popover(f"모델: {model_label}", use_container_width=True):
                provider = st.session_state.get('provider', '?')
                st.caption(f"공급자: {provider}")
                model_picker('모델 선택', key_prefix='inline', instant=True)
                if st.button('상세 설정 열기', key='inline_open_settings',
                             use_container_width=True):
                    st.session_state['active_view'] = 'settings'
                    st.rerun()
        else:
            st.caption(f"모델: {model_label}")
    st.divider()

    # Render existing history first.
    for i in range(len(st.session_state['generated_responses'])):
        with st.chat_message('user'):
            st.markdown(st.session_state['user_inputs'][i])
        with st.chat_message('assistant'):
            reasoning = (
                st.session_state['thinking_traces'][i]
                if i < len(st.session_state['thinking_traces']) else ''
            )
            retrieved = (
                st.session_state['retrieved_per_turn'][i]
                if i < len(st.session_state['retrieved_per_turn']) else []
            )
            variants = (
                st.session_state['query_variants_per_turn'][i]
                if i < len(st.session_state['query_variants_per_turn']) else []
            )
            render_assistant(
                st.session_state['generated_responses'][i],
                reasoning, retrieved, i, variants=variants,
            )

    # Process pending input from a suggestion click (rerun after click sets it).
    pending = st.session_state.pop('_pending_input', None)
    if pending:
        handle_chat_turn(pending)
    elif not st.session_state['generated_responses']:
        # Hero with full brand lockup + suggestion chips.
        logo_html = (
            f'<img src="{_LOGO_URI}" '
            f'style="width:100%; max-width:340px; height:auto; '
            f'margin:0 auto 18px auto; display:block; opacity:0.95;" />'
            if _LOGO_URI else ''
        )
        st.markdown(
            f'<div class="empty-hero">'
            f'{logo_html}'
            f'<h2>오늘은 어떤 걸 도와드릴까요?</h2>'
            f'<p>문서를 올려 그 내용으로 답을 받거나, 일반 챗으로 바로 시작할 수 있습니다.</p>'
            f'</div>',
            unsafe_allow_html=True,
        )
        if st.session_state['docs']:
            suggestions = [
                '업로드한 문서들을 한 문단으로 요약해 줘',
                '문서에서 가장 중요한 개념 세 가지는?',
                '문서들 사이의 주요 차이점을 비교해 줘',
                '내가 어떤 질문을 해볼 수 있을까',
            ]
        else:
            suggestions = [
                'Personal RAG가 어떻게 동작하는지 설명해 줘',
                '문서를 업로드하려면 어떻게 하는가',
                '임베딩 모델은 무엇을 골라야 좋을까',
                '검색 품질을 높이는 팁은',
            ]
        cols = st.columns(2)
        for i, s in enumerate(suggestions):
            with cols[i % 2]:
                if st.button(s, key=f'suggest_{i}', use_container_width=True):
                    st.session_state['_pending_input'] = s
                    st.rerun()

    if st.session_state['generated_responses']:
        cols = st.columns([1, 5])
        with cols[0]:
            if st.button('대화 초기화', use_container_width=True):
                st.session_state['user_inputs'] = []
                st.session_state['generated_responses'] = []
                st.session_state['thinking_traces'] = []
                st.session_state['retrieved_per_turn'] = []
                st.session_state['query_variants_per_turn'] = []
                st.rerun()

    if local_active and web_active:
        _placeholder = '문서와 웹에서 검색하여 답합니다'
    elif local_active:
        _placeholder = '문서에 대해 질문하세요'
    elif web_active:
        _placeholder = '웹 검색 후 답합니다'
    else:
        _placeholder = '메시지를 입력하세요 (.txt / .md / .pdf 파일 첨부 가능)'

    # Streamlit 1.42+: chat_input supports inline file attachments.
    # Fall back to plain chat_input on older versions.
    try:
        submitted = st.chat_input(
            _placeholder,
            accept_file='multiple',
            file_type=['txt', 'md', 'pdf', 'docx', 'csv', 'hwpx'],
        )
    except TypeError:
        submitted = st.chat_input(_placeholder)

    if submitted:
        # Normalize: newer Streamlit returns an object with .text and .files;
        # older returns a bare string.
        text_part = ''
        file_part = []
        if hasattr(submitted, 'text') or hasattr(submitted, 'files'):
            text_part = (getattr(submitted, 'text', '') or '').strip()
            file_part = list(getattr(submitted, 'files', []) or [])
        else:
            text_part = (submitted or '').strip()

        # Ingest any attached files first so they're searchable on the same turn.
        # ingest_files() appends to st.session_state['docs'] and saves to disk,
        # so the files immediately show up in the Documents tab too.
        added = 0
        if file_part:
            # ingest_files() shows its own st.status with per-batch progress.
            added = ingest_files(file_part)
            if added > 0:
                try:
                    st.toast(
                        f'{added}개 문서 인덱싱 완료 — 문서 탭에 추가되었고 '
                        '이번 질문부터 검색 대상입니다.'
                    )
                except Exception:
                    pass

        if text_part:
            handle_chat_turn(text_part)

        # Rerun after any file ingest so the sidebar status counter and
        # the Documents tab list refresh immediately (otherwise they stay
        # stale until the next user interaction).
        if added > 0:
            st.rerun()


# =============================================================================
# Documents view
# =============================================================================

def view_docs():
    _section(
        '문서 업로드',
        '.txt / .md / .pdf 파일을 업로드하면 자동으로 청킹·임베딩되어 디스크에 영속 저장됩니다. '
        '같은 파일을 다시 올려도 내용 해시가 일치하면 즉시 캐시에서 복원됩니다.',
    )
    uploaded = st.file_uploader(
        ' ',
        type=['txt', 'md', 'pdf', 'docx', 'csv', 'hwpx'],
        accept_multiple_files=True,
        label_visibility='collapsed',
    )
    if uploaded:
        # ingest_files() shows its own st.status with per-batch progress.
        added = ingest_files(uploaded)
        if added > 0:
            st.success(f'{added}개 새 문서 인덱싱 완료')
            st.rerun()

    st.write('')
    _section('인덱싱된 문서')

    if not st.session_state['docs']:
        _empty('아직 업로드된 문서가 없습니다. 위 영역에 파일을 끌어다 놓으세요.')
    else:
        total_chunks = sum(len(d['chunks']) for d in st.session_state['docs'])
        total_chars = sum(len(d.get('raw_text', '')) for d in st.session_state['docs'])
        st.caption(
            f"문서 {len(st.session_state['docs'])}개 · "
            f"청크 {total_chunks}개 · {total_chars:,}자"
        )
        for d in list(st.session_state['docs']):
            with st.container(border=True):
                row = st.columns([4, 1, 1, 1, 1])
                row[0].markdown(f"**{d['name']}**")
                row[0].caption(
                    f"id: `{d['id']}` · "
                    f"chunk_size={d.get('chunk_size', '?')} · "
                    f"overlap={d.get('chunk_overlap', '?')}"
                )
                row[1].metric('chunks', len(d['chunks']))
                row[2].metric('pages', d.get('page_count', 0) or '—')
                row[3].metric('images', 'on' if d.get('has_page_images') else 'off')
                with row[4]:
                    st.write('')
                    if st.button('삭제', key=f"del_doc_{d['id']}", use_container_width=True):
                        remove_doc(d['id'])
                        st.rerun()
        if st.button('모든 문서 삭제 (디스크 포함)', type='secondary'):
            for d in list(st.session_state['docs']):
                remove_doc(d['id'])
            st.rerun()

    st.write('')
    with st.expander('고급 청킹 설정'):
        st.caption(
            '값을 바꾸면 현재 보유한 모든 문서가 자동으로 재청킹·재임베딩됩니다. '
            '이전 설정의 캐시는 보존되어 되돌릴 수 있습니다.'
        )
        cs_cols = st.columns(2)
        with cs_cols[0]:
            new_size = st.number_input(
                '청크 크기 (문자)', 100, 4000,
                int(st.session_state['chunk_size']),
                help='하나의 청크에 들어가는 최대 글자 수.',
            )
        with cs_cols[1]:
            new_overlap = st.number_input(
                '청크 겹침 (문자)', 0, 1000,
                int(st.session_state['chunk_overlap']),
                help='이웃한 청크 사이에 중복으로 포함시킬 글자 수. 문맥 유실 방지.',
            )
        if (new_size, new_overlap) != (
            st.session_state['chunk_size'], st.session_state['chunk_overlap']
        ):
            st.session_state['chunk_size'] = new_size
            st.session_state['chunk_overlap'] = new_overlap
            if st.session_state['docs']:
                with st.spinner('재청킹 및 재임베딩 중...'):
                    reindex_all()
                st.rerun()

    st.write('')
    _section(
        '검색 미리보기',
        'LLM 호출 없이 검색 결과만 확인합니다. 어떤 청크가 어떤 점수로 잡히는지 점검하는 용도.',
    )
    preview_cols = st.columns([5, 1])
    with preview_cols[0]:
        preview_query = st.text_input(
            '검색어', key='preview_query', label_visibility='collapsed',
            placeholder='예: 매출 성장률 정의',
        )
    with preview_cols[1]:
        run_preview = st.button(
            '검색 실행',
            disabled=(not preview_query.strip() or not st.session_state['docs']),
            use_container_width=True,
        )
    if run_preview and preview_query.strip():
        with st.spinner('검색 중...'):
            results = retrieve_local(preview_query.strip())
        if not results:
            st.warning('결과 없음.')
        else:
            st.caption(f"{len(results)}개 청크 반환")
            for j, r in enumerate(results, start=1):
                with st.expander(_citation_summary(j, r), expanded=False):
                    _citation_body(r)


# =============================================================================
# Settings view
# =============================================================================

def view_settings():
    _section(
        '설정',
        '모델·검색·웹·응답을 카테고리별로 정리. 변경은 즉시 저장됩니다.',
    )

    # ----- Top status snapshot -----
    active_p = st.session_state.get('provider', 'Hugging Face Router')
    model_short = (st.session_state.get('model') or '—')
    if len(model_short) > 30:
        model_short = model_short[:27] + '...'
    embedder_label = {
        'BAAI/bge-m3': 'BGE-M3',
        'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2': 'MiniLM',
    }.get(st.session_state.get('embedder_model', ''), '—')
    has_key = bool(_active_api_key())
    with st.container(border=True):
        m = st.columns(4)
        m[0].metric('공급자', active_p)
        m[1].metric('모델', model_short)
        m[2].metric('임베더', embedder_label)
        m[3].metric('API 키', '설정됨' if has_key else '없음')

    # ----- Tabbed sections -----
    tab_llm, tab_search, tab_web, tab_response, tab_advanced = st.tabs(
        ['LLM 공급자', 'RAG 검색', '웹 검색', '응답 / 멀티모달', '고급']
    )

    # ============== LLM ==============
    with tab_llm:
        with st.container(border=True):
            st.markdown('##### 공급자 선택')
            current_provider = st.session_state.get('provider', PROVIDER_NAMES[0])
            if current_provider not in PROVIDER_NAMES:
                current_provider = PROVIDER_NAMES[0]
            new_provider = st.selectbox(
                '공급자', PROVIDER_NAMES,
                index=PROVIDER_NAMES.index(current_provider),
                help='프리셋을 선택하면 엔드포인트 주소·기본 모델·환경변수에서 API 키를 자동으로 채웁니다.',
                label_visibility='collapsed',
            )
            if new_provider != current_provider:
                cfg = PROVIDERS[new_provider]
                if cfg['base_url']:
                    st.session_state['base_url'] = cfg['base_url']
                if cfg['default_model']:
                    st.session_state['model'] = cfg['default_model']
                st.session_state['provider'] = new_provider
                st.rerun()
            else:
                st.session_state['provider'] = new_provider
            model_picker('모델', key_prefix='settings')

        # API keys — staged via pending state. Typing into a field doesn't
        # touch session_state[<active_key>] until the user presses 저장.
        # That keeps every keystroke from triggering _save_user_prefs and
        # the disk/Supabase round-trips it carries, and gives the user a
        # clear "I'm done editing" moment.
        with st.container(border=True):
            st.markdown('##### API 키')
            st.caption(
                '공급자별로 따로 저장됩니다. "(사용 중)" 표시가 현재 공급자 키. '
                '입력 후 아래 **저장** 버튼을 눌러야 반영됩니다.'
            )

            def _key_label(name, owner):
                return f'{name} (사용 중)' if owner == active_p else name

            _key_specs = [
                ('hf_api_key',         'Hugging Face',         'Hugging Face Router',
                 'hf_...',
                 'HF Inference Router (Gemma / DeepSeek / Qwen 등). '
                 'fine-grained 토큰 권장: https://huggingface.co/settings/tokens'),
                ('openai_api_key',     'OpenAI',               'OpenAI',
                 'sk-...',
                 'gpt-4o / gpt-5 / o3 등. https://platform.openai.com/api-keys'),
                ('anthropic_api_key',  'Anthropic (Claude)',   'Anthropic (Claude)',
                 'sk-ant-...',
                 'Claude 4.x. https://console.anthropic.com/settings/keys'),
                ('fireworks_api_key',  'Fireworks AI',         'Fireworks AI',
                 'fw_...',
                 '오픈모델 빠른 추론. https://fireworks.ai/account/api-keys'),
                ('dashscope_api_key',  'DashScope (Qwen)',     'DashScope (Qwen)',
                 'sk-...',
                 'Qwen 공식 API.'),
                ('custom_api_key',     'Custom / vLLM',        'vLLM / local',
                 '(self-host endpoint key)',
                 'vLLM 등 셀프 호스팅 / Custom OpenAI-호환 endpoint.'),
            ]

            def _render_key_row(active_key, label_text, owner, ph, help_text):
                pending_key = f'_pending_{active_key}'
                # First-time render seeds the widget from the active value.
                # On subsequent reruns Streamlit retains the user's typing
                # via the widget key, so we don't re-seed.
                if pending_key not in st.session_state:
                    st.session_state[pending_key] = (
                        st.session_state.get(active_key, '') or ''
                    )
                st.text_input(
                    _key_label(label_text, owner),
                    type='password',
                    placeholder=ph,
                    help=help_text,
                    key=pending_key,
                )
                pending_val = st.session_state.get(pending_key, '') or ''
                active_val = st.session_state.get(active_key, '') or ''
                if pending_val != active_val:
                    bcols = st.columns([2, 1, 1])
                    bcols[0].caption('변경 대기 — 적용을 눌러야 호출에 반영됩니다.')
                    if bcols[1].button(
                        '취소', key=f'reset_{active_key}',
                        use_container_width=True,
                    ):
                        st.session_state[pending_key] = active_val
                        st.rerun()
                    if bcols[2].button(
                        '적용', key=f'apply_{active_key}',
                        type='primary', use_container_width=True,
                    ):
                        st.session_state[active_key] = pending_val
                        try:
                            _save_user_prefs()
                        except Exception:
                            pass
                        st.rerun()

            kc1, kc2 = st.columns(2)
            for i, spec in enumerate(_key_specs):
                col = kc1 if i % 2 == 0 else kc2
                with col:
                    _render_key_row(*spec)

            # Saved-count summary at the bottom.
            set_count = sum(
                1 for active_key, *_ in _key_specs
                if (st.session_state.get(active_key, '') or '').strip()
            )
            st.caption(f'현재 {set_count}/{len(_key_specs)} 개의 키가 저장돼 있습니다.')

    # ============== RAG 검색 ==============
    with tab_search:
        # Mode + embedder
        with st.container(border=True):
            st.markdown('##### 동작 모드')
            st.session_state['general_chat_mode'] = st.checkbox(
                '일반 대화 모드 (RAG / 웹 검색 끔)',
                value=st.session_state['general_chat_mode'],
                help='업로드 문서·웹 검색을 모두 건너뛰고 LLM 본연 지식으로 답합니다.',
            )
            prev_embedder = st.session_state['embedder_model']
            emb_labels = {
                'BAAI/bge-m3': 'BGE-M3 (한국어 강함, 2.2GB)',
                'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2':
                    'MiniLM 다국어 (가벼움, 470MB)',
            }
            emb_idx = (
                EMBEDDER_CHOICES.index(st.session_state['embedder_model'])
                if st.session_state['embedder_model'] in EMBEDDER_CHOICES else 0
            )
            st.session_state['embedder_model'] = st.selectbox(
                '임베딩 모델', EMBEDDER_CHOICES,
                index=emb_idx,
                format_func=lambda x: emb_labels.get(x, x),
                help='문서·질문을 벡터로 변환하는 모델.',
            )
            if st.session_state['embedder_model'] != prev_embedder:
                st.session_state['_loaded_for_embedder'] = None
                load_all_for_current_embedder()
                st.rerun()

        # Retrieval pipeline
        with st.container(border=True):
            st.markdown('##### 검색 파이프라인')
            mode_labels = {'hybrid': '하이브리드 (권장)',
                           'dense': '의미 기반만', 'bm25': '키워드만'}
            st.session_state['retrieval_mode'] = st.radio(
                '검색 방식',
                ['hybrid', 'dense', 'bm25'],
                index=['hybrid', 'dense', 'bm25'].index(
                    st.session_state['retrieval_mode']
                ),
                horizontal=True,
                format_func=lambda x: mode_labels[x],
            )
            cc1, cc2 = st.columns(2)
            with cc1:
                st.session_state['use_reranker'] = st.checkbox(
                    '재정렬 모델 사용 (Cross-encoder)',
                    value=st.session_state['use_reranker'],
                    help='정확도 향상, 응답 약간 느려짐.',
                )
                st.session_state['use_contextual_rewrite'] = st.checkbox(
                    '이어지는 질문 자동 보완',
                    value=st.session_state['use_contextual_rewrite'],
                    help='이전 대화 맥락으로 self-contained 질문 재작성.',
                )
                st.session_state['per_doc_balance'] = st.checkbox(
                    '여러 문서 균형 검색',
                    value=st.session_state['per_doc_balance'],
                    help='상위 결과가 한 문서에 쏠리지 않게 강제.',
                )
                st.session_state['comparison_autodetect'] = st.checkbox(
                    '비교 질문 자동 감지',
                    value=st.session_state['comparison_autodetect'],
                )
            with cc2:
                _pgv_available = _supabase_client() is not None
                st.session_state['use_pgvector_search'] = st.checkbox(
                    'pgvector 의미 검색 (Supabase)',
                    value=st.session_state['use_pgvector_search'],
                    help='Supabase pgvector 로 dense 검색. 미연결 시 자동 폴백.',
                    disabled=not _pgv_available,
                )
                if not _pgv_available:
                    st.caption('Supabase 미연동 — 비활성')
                st.session_state['use_agentic_search'] = st.checkbox(
                    '에이전트 검색 (LLM 추가 검색)',
                    value=st.session_state['use_agentic_search'],
                    help='LLM 이 도구 호출로 추가 검색 발행. function calling 지원 모델 필요.',
                )
                if st.session_state['use_agentic_search']:
                    st.session_state['agentic_max_iters'] = st.slider(
                        '최대 라운드', 1, 5,
                        int(st.session_state.get('agentic_max_iters', 3)),
                    )

    # ============== 웹 ==============
    with tab_web:
        with st.container(border=True):
            st.markdown('##### 실시간 웹 검색')
            st.caption(
                '질문 시 웹 검색 결과를 컨텍스트에 함께 포함. DuckDuckGo 는 API 키 불필요.'
            )
            st.session_state['web_enabled'] = st.checkbox(
                '웹 검색 사용',
                value=st.session_state['web_enabled'],
            )
            if st.session_state['web_enabled']:
                wp_labels = {
                    'duckduckgo': 'DuckDuckGo (키 불필요)',
                    'tavily': 'Tavily (LLM 최적화, 키 필요)',
                    'brave': 'Brave (키 필요)',
                }
                wc1, wc2 = st.columns(2)
                with wc1:
                    st.session_state['web_provider'] = st.selectbox(
                        '검색 제공자', ['duckduckgo', 'tavily', 'brave'],
                        index=['duckduckgo', 'tavily', 'brave'].index(
                            st.session_state['web_provider']
                        ),
                        format_func=lambda x: wp_labels[x],
                    )
                with wc2:
                    st.session_state['web_top_n'] = st.number_input(
                        '결과 수', 1, 20, int(st.session_state['web_top_n']),
                    )
                if st.session_state['web_provider'] == 'tavily':
                    st.session_state['tavily_key'] = st.text_input(
                        'Tavily API 키', st.session_state['tavily_key'],
                        type='password', placeholder='tvly-...',
                    )
                elif st.session_state['web_provider'] == 'brave':
                    st.session_state['brave_key'] = st.text_input(
                        'Brave API 키', st.session_state['brave_key'],
                        type='password',
                    )
            else:
                st.caption('웹 검색이 꺼져 있어 추가 설정이 보이지 않습니다.')

    # ============== 응답 / 멀티모달 ==============
    with tab_response:
        with st.container(border=True):
            st.markdown('##### 응답 동작')
            rc1, rc2 = st.columns(2)
            with rc1:
                st.session_state['stream'] = st.checkbox(
                    '스트리밍 응답', value=st.session_state['stream'],
                    help='응답을 토큰 단위로 실시간 표시.',
                )
            with rc2:
                st.session_state['enable_thinking'] = st.checkbox(
                    '추론 모드 사용', value=st.session_state['enable_thinking'],
                    help='지원 모델에서 reasoning 토큰을 분리 표시.',
                )

        with st.container(border=True):
            st.markdown('##### 멀티모달 (이미지)')
            st.caption(
                'PDF 페이지 이미지를 LLM 에 함께 전달해 표·차트·도식을 이해하게 합니다. '
                '비전 입력 지원 모델 필요.'
            )
            st.session_state['include_page_images'] = st.checkbox(
                'PDF 페이지 이미지 첨부',
                value=st.session_state['include_page_images'],
            )
            if st.session_state['include_page_images']:
                st.session_state['max_page_images'] = st.number_input(
                    '한 턴에 보낼 이미지 수', 1, 10,
                    int(st.session_state['max_page_images']),
                )

    # ============== 고급 ==============
    with tab_advanced:
        with st.container(border=True):
            st.markdown('##### 엔드포인트')
            st.session_state['base_url'] = st.text_input(
                '엔드포인트 주소', st.session_state['base_url'],
                help='OpenAI 호환 endpoint. {base_url}/chat/completions 가 호출됩니다.',
            )

        with st.container(border=True):
            st.markdown('##### 샘플링 / 응답 길이')
            ac1, ac2 = st.columns(2)
            with ac1:
                st.session_state['max_tokens'] = st.number_input(
                    '최대 응답 토큰', 16, 131072,
                    int(st.session_state['max_tokens']),
                )
                st.session_state['temperature'] = st.slider(
                    'temperature', 0.0, 2.0,
                    float(st.session_state['temperature']), 0.05,
                )
                st.session_state['top_p'] = st.slider(
                    'top_p', 0.0, 1.0, float(st.session_state['top_p']), 0.01,
                )
            with ac2:
                st.session_state['sampling_top_k'] = st.number_input(
                    'top_k', 1, 200, int(st.session_state['sampling_top_k']),
                )
                st.session_state['presence_penalty'] = st.slider(
                    'presence_penalty', 0.0, 2.0,
                    float(st.session_state['presence_penalty']), 0.1,
                )

        with st.container(border=True):
            st.markdown('##### 검색 정밀도')
            ad1, ad2 = st.columns(2)
            with ad1:
                st.session_state['retrieve_top_n'] = st.number_input(
                    '1차 후보 수', 1, 200,
                    int(st.session_state['retrieve_top_n']),
                    help='재정렬 전에 가져올 후보 청크 수.',
                )
                st.session_state['final_top_k'] = st.number_input(
                    '최종 청크 수', 1, 50,
                    int(st.session_state['final_top_k']),
                    help='LLM 에 컨텍스트로 전달할 최종 청크 수.',
                )
            with ad2:
                st.session_state['per_doc_reserve'] = st.number_input(
                    '문서당 최소 청크', 1, 5,
                    int(st.session_state['per_doc_reserve']),
                )
            st.session_state['use_multi_query'] = st.checkbox(
                '다중 쿼리 (paraphrase)',
                value=st.session_state['use_multi_query'],
                help='질문을 여러 표현으로 변형 후 합집합 검색.',
            )
            if st.session_state['use_multi_query']:
                st.session_state['n_paraphrases'] = st.number_input(
                    '변형 개수', 1, 8,
                    int(st.session_state['n_paraphrases']),
                )
            st.session_state['use_hyde'] = st.checkbox(
                'HyDE (가상 답안 검색)',
                value=st.session_state['use_hyde'],
            )


# =============================================================================
# Cache view
# =============================================================================

def view_cache():
    _section(
        '데이터 & 저장소',
        '클라우드 영속 저장, 로컬 디스크, 모델 캐시를 한 곳에서 관리.',
    )

    user_logs = _user_logs_dir()
    user_dd = _user_data_dir()

    # ----- Compute compact stats for the overview row -----
    sb_connected = _supabase_client() is not None
    sb_attempts = st.session_state.get('_sb_attempts', 0)
    sb_failures = st.session_state.get('_sb_failures', 0)
    pgv_successes = st.session_state.get('_pgv_successes', 0)
    pgv_failures = st.session_state.get('_pgv_failures', 0)

    # Local disk usage (sum of all files under user_dd).
    local_bytes = 0
    if user_dd.exists():
        for p in user_dd.rglob('*'):
            if p.is_file():
                try:
                    local_bytes += p.stat().st_size
                except Exception:
                    pass

    # HF model cache size.
    hf_cached = []
    try:
        from huggingface_hub import scan_cache_dir
        info = scan_cache_dir()
        hf_cached = [(r.repo_id, r.size_on_disk_str, str(r.repo_path), r.size_on_disk)
                     for r in info.repos]
    except Exception:
        hf_cached = []
    hf_total_bytes = sum(x[3] for x in hf_cached)

    def _fmt_size(b):
        if b < 1024:
            return f'{b} B'
        if b < 1024 * 1024:
            return f'{b / 1024:.1f} KB'
        if b < 1024 ** 3:
            return f'{b / 1024 / 1024:.1f} MB'
        return f'{b / 1024 ** 3:.2f} GB'

    # ----- Top overview: 4 metric cards -----
    with st.container(border=True):
        m = st.columns(4)
        if sb_connected:
            sb_label = '정상' if sb_failures == 0 else f'{sb_failures}건 실패'
            m[0].metric('영속 로깅', sb_label, delta=f'{sb_attempts} 시도' if sb_attempts else None)
        else:
            m[0].metric('영속 로깅', '미설정')
        if sb_connected and pgv_successes:
            pgv_label = f'{pgv_successes}건' if pgv_failures == 0 else f'{pgv_failures}건 실패'
            m[1].metric('pgvector', pgv_label)
        elif sb_connected:
            m[1].metric('pgvector', '대기')
        else:
            m[1].metric('pgvector', '미연결')
        m[2].metric('로컬 디스크', _fmt_size(local_bytes))
        m[3].metric('HF 모델 캐시', _fmt_size(hf_total_bytes))

    # ----- Tabbed body -----
    tab_supabase, tab_local, tab_hf = st.tabs(
        ['클라우드 영속 (Supabase)', '로컬 디스크', 'HF 모델 캐시']
    )

    # ============== Supabase ==============
    with tab_supabase:
        if not sb_connected:
            with st.container(border=True):
                if _is_streamlit_cloud():
                    st.warning(
                        '영속 로깅 미설정 — 아래 로컬 JSONL 은 컨테이너 재시작 시 사라집니다.'
                    )
                    st.markdown(
                        '**활성화 방법:** Manage app → Settings → Secrets 에 다음 추가:\n'
                        '```toml\n'
                        'SUPABASE_URL = "https://xxxx.supabase.co"\n'
                        'SUPABASE_KEY = "eyJ..."\n'
                        '```\n'
                        '그 후 SQL Editor 에서 `db_schema.sql` / `db_schema_pgvector.sql` / `db_schema_users.sql` 실행.'
                    )
                else:
                    st.info(
                        '로컬 개발 중에는 JSONL 만으로 충분합니다. '
                        'Cloud 배포 시 `SUPABASE_URL` / `SUPABASE_KEY` 설정 권장.'
                    )
        else:
            # Status cards
            with st.container(border=True):
                st.markdown('##### 로깅 (chat_turns / agent_runs / events)')
                successes = st.session_state.get('_sb_successes', 0)
                last_err = st.session_state.get('_sb_last_err')
                if sb_attempts == 0:
                    st.info('이 세션에서 INSERT 시도 없음. 아래 진단 버튼으로 즉시 테스트 가능.')
                elif sb_failures == 0:
                    st.success(f'INSERT {successes}/{sb_attempts} 성공. 컨테이너 재시작에도 보존.')
                else:
                    st.error(
                        f'INSERT {sb_failures}/{sb_attempts} 실패. '
                        'RLS / 스키마 누락 등 점검 필요.'
                    )
                    if last_err:
                        with st.expander('마지막 실패 메시지', expanded=True):
                            st.code(last_err)

                if st.button(
                    '연결 진단 (events 1행 INSERT + DELETE)',
                    key='sb_diagnose_btn',
                ):
                    import time as _t
                    client = _supabase_client()
                    probe = {
                        'event_type': 'diagnostic_probe',
                        'user_id': st.session_state.get('user_id', '_local'),
                        'payload': {'ts': _t.time(), 'note': '캐시 탭 진단 버튼'},
                    }
                    try:
                        ins = client.table('events').insert(probe).execute()
                        ins_id = (ins.data[0]['id']
                                  if getattr(ins, 'data', None) and ins.data else None)
                        st.success(f'INSERT 성공 (id={ins_id}). 정리 중...')
                        if ins_id is not None:
                            try:
                                client.table('events').delete().eq('id', ins_id).execute()
                                st.info(f'정리 완료 (id={ins_id} 삭제). DB 정상.')
                            except Exception as de:
                                st.warning(
                                    f'INSERT 됐는데 DELETE 실패 ({de}).'
                                )
                    except Exception as e:
                        st.error(f'INSERT 실패: {type(e).__name__}')
                        st.code(str(e)[:1500])
                        msg = str(e).lower()
                        if 'row-level security' in msg or 'rls' in msg or 'policy' in msg:
                            st.markdown(
                                '**원인: RLS 차단.** SQL Editor 에서 한 번 실행:\n\n'
                                '```sql\n'
                                'alter table public.events     disable row level security;\n'
                                'alter table public.chat_turns disable row level security;\n'
                                'alter table public.agent_runs disable row level security;\n'
                                '```'
                            )
                        elif 'does not exist' in msg or '42p01' in msg:
                            st.markdown('**원인: 테이블 없음.** `db_schema.sql` 실행 필요.')
                        elif '401' in msg or '403' in msg or 'unauthorized' in msg:
                            st.markdown('**원인: 키 권한.** anon public 키인지 확인.')

            # pgvector status
            with st.container(border=True):
                st.markdown('##### pgvector (청크 임베딩 영속화)')
                pgv_attempts = st.session_state.get('_pgv_attempts', 0)
                pgv_last_err = st.session_state.get('_pgv_last_err')
                if pgv_attempts == 0:
                    st.info(
                        '이 세션 인덱싱 없음. 문서 업로드 시 청크 임베딩이 '
                        '`doc_chunks` 테이블에 자동 저장됩니다. '
                        '(스키마: `db_schema_pgvector.sql`)'
                    )
                elif pgv_failures == 0:
                    st.success(
                        f'청크 임베딩 {pgv_successes}/{pgv_attempts} 영속화 성공.'
                    )
                else:
                    st.error(
                        f'영속화 {pgv_failures}/{pgv_attempts} 실패. '
                        '`db_schema_pgvector.sql` 적용 확인.'
                    )
                    if pgv_last_err:
                        with st.expander('마지막 pgvector 에러', expanded=False):
                            st.code(pgv_last_err)

                # Active retrieval source.
                if st.session_state.get('use_pgvector_search'):
                    n = st.session_state.get('_pgv_search_last_n')
                    err = st.session_state.get('_pgv_search_last_err')
                    if err:
                        st.warning('의미 검색 경로: pgvector (직전 호출 실패 → 로컬 폴백)')
                        with st.expander('검색 에러', expanded=False):
                            st.code(err)
                    elif n is None:
                        st.caption('의미 검색 경로: pgvector (아직 호출 없음)')
                    else:
                        st.caption(f'의미 검색 경로: pgvector — 직전 호출 {n}건')
                else:
                    st.caption('의미 검색 경로: 로컬 numpy (in-memory)')

    # ============== 로컬 디스크 ==============
    with tab_local:
        # Aggregate logs first (agents + events)
        with st.container(border=True):
            st.markdown('##### 통합 로그 파일')
            found_any = False
            for fname, label_text in (
                ('events.jsonl', '로그인 / 문서 / 세션 / LLM 에러 이벤트'),
                ('agents.jsonl', '에이전트 실행 기록'),
            ):
                fpath = user_logs / fname
                if not fpath.exists():
                    continue
                found_any = True
                try:
                    n_lines = sum(1 for _ in fpath.open('r', encoding='utf-8'))
                except Exception:
                    n_lines = '?'
                size_kb = fpath.stat().st_size / 1024
                unit = 'events' if 'events' in fname else 'runs'
                cols = st.columns([5, 1, 1])
                cols[0].markdown(f"**`{fname}`** · {label_text}")
                cols[1].caption(f'{n_lines} {unit} · {size_kb:.1f} KB')
                with cols[2]:
                    try:
                        st.download_button(
                            '다운로드', data=fpath.read_bytes(),
                            file_name=fname, mime='application/x-jsonlines',
                            key=f'dl_{fname}', use_container_width=True,
                        )
                    except Exception:
                        pass
            if not found_any:
                st.caption('아직 통합 로그 파일이 없습니다.')

        # Per-session chat logs
        _AGGREGATE_LOGS = {'agents.jsonl', 'events.jsonl'}
        jsonl_files = sorted(
            [p for p in user_logs.glob('*.jsonl') if p.name not in _AGGREGATE_LOGS],
            key=lambda p: p.stat().st_mtime, reverse=True,
        )
        with st.container(border=True):
            st.markdown('##### 세션별 대화 로그')
            st.caption(f'경로: `{user_logs}`')
            if not jsonl_files:
                st.caption('대화 로그는 아직 없습니다.')
            else:
                for p in jsonl_files[:30]:
                    try:
                        n_lines = sum(1 for _ in p.open('r', encoding='utf-8'))
                    except Exception:
                        n_lines = '?'
                    size_kb = p.stat().st_size / 1024
                    cols = st.columns([5, 1, 1])
                    cols[0].markdown(f"`{p.name}`")
                    cols[1].caption(f'{n_lines} turns · {size_kb:.1f} KB')
                    with cols[2]:
                        try:
                            st.download_button(
                                '다운로드', data=p.read_bytes(),
                                file_name=p.name,
                                mime='application/x-jsonlines',
                                key=f'dl_jsonl_{p.name}',
                                use_container_width=True,
                            )
                        except Exception:
                            pass
                if len(jsonl_files) > 30:
                    st.caption(f'표시 30개 / 전체 {len(jsonl_files)}개')

        # Local vector store
        with st.container(border=True):
            st.markdown('##### 로컬 벡터 스토어')
            st.caption(f'경로: `{user_dd}` — 임베더 모델마다 하위 폴더.')
            if not user_dd.exists():
                st.caption('아직 저장된 벡터 인덱스가 없습니다.')
            else:
                rows = []
                for sub in sorted(user_dd.iterdir()):
                    if not sub.is_dir() or sub.name == 'sessions':
                        continue
                    doc_dirs = [p for p in sub.iterdir() if p.is_dir()]
                    total = 0
                    for p in sub.rglob('*'):
                        if p.is_file():
                            try:
                                total += p.stat().st_size
                            except Exception:
                                pass
                    rows.append((sub.name, len(doc_dirs), total))
                if not rows:
                    st.caption('아직 저장된 벡터 인덱스가 없습니다.')
                else:
                    for name, n, size in rows:
                        cols = st.columns([4, 1, 1])
                        cols[0].markdown(f"`{name}`")
                        cols[1].caption(f'{n} docs')
                        cols[2].caption(_fmt_size(size))

    # ============== HF 모델 캐시 ==============
    with tab_hf:
        with st.container(border=True):
            st.markdown('##### Hugging Face 다운로드 모델')
            st.caption(
                '임베더 / reranker / 자가호스팅 모델 등 로컬 다운로드된 가중치. '
                '사용하지 않는 항목은 삭제해 디스크 회수.'
            )
            if not hf_cached:
                st.caption('캐시된 모델이 없습니다.')
            else:
                for repo_id, size_str, path, _bytes in hf_cached:
                    cols = st.columns([4, 1, 1])
                    cols[0].markdown(f"`{repo_id}`")
                    cols[1].caption(size_str)
                    if cols[2].button(
                        '삭제', key=f'cache_{repo_id}',
                        use_container_width=True,
                    ):
                        try:
                            shutil.rmtree(path)
                            st.success(f'{repo_id} 삭제 완료')
                            st.rerun()
                        except Exception as e:
                            st.error(f'삭제 실패: {e}')


# =============================================================================
# Agentic workflows
# =============================================================================

# NOTE: legacy global. Per-user log path is _agent_log_path(); kept here only
# to avoid accidental import-time errors. Use _agent_log_path() everywhere.


def _agent_format_context(retrieved: list) -> str:
    """Format retrieved chunks as numbered context lines for the LLM prompt."""
    if not retrieved:
        return '(검색된 근거 자료 없음)'
    lines = []
    for i, r in enumerate(retrieved, start=1):
        if r.get('source') == 'web':
            lines.append(f"[{i}] (웹: {r.get('doc', '')} {r.get('url', '')}) {r.get('text', '')}")
        else:
            pages = r.get('pages') or []
            page_str = ''
            if pages:
                page_str = (f' p.{pages[0]}' if len(pages) == 1
                            else f' pp.{pages[0]}-{pages[-1]}')
            lines.append(
                f"[{i}] (로컬: {r.get('doc', '')}{page_str}) {r.get('text', '')}"
            )
    return '\n\n'.join(lines)


def _agent_run_llm(messages: list, model: str):
    """Common runner: builds params, streams or non-streams, post-processes."""
    extra_body = {}
    if _provider_supports_top_k():
        extra_body['top_k'] = int(st.session_state['sampling_top_k'])
    if not st.session_state['enable_thinking']:
        if _is_dashscope_endpoint():
            extra_body['enable_thinking'] = False
        elif not _is_openai_endpoint() and not _is_hf_router_endpoint():
            extra_body['chat_template_kwargs'] = {'enable_thinking': False}

    params = _build_completion_params(
        model=model, messages=messages,
        max_tokens=st.session_state['max_tokens'],
        temperature=st.session_state['temperature'],
        top_p=st.session_state['top_p'],
        presence_penalty=st.session_state['presence_penalty'],
        extra_body=extra_body,
    )
    import time as _time
    client = get_openai_client()
    t0 = _time.time()
    if st.session_state['stream']:
        full_text, reasoning = stream_chat(client, params)
    else:
        with st.spinner('생성 중...'):
            full_text, reasoning = non_stream_chat(client, params)
    elapsed = _time.time() - t0

    if full_text and not reasoning:
        rt, ct = split_thinking(full_text)
        if rt:
            reasoning, full_text = rt, ct
    return full_text, reasoning, elapsed


def _agent_log(task_key: str, inputs: dict, output: str, retrieved: list,
               model: str, elapsed: float):
    """Append one agent run to ./logs/agents.jsonl. Same schema family as chat
    JSONL — analysis-friendly."""
    # Slim: 200-char preview + reference, matches log_turn_structured.
    retrieved_records = []
    for j, r in enumerate(retrieved or [], start=1):
        rec = {'rank': j, 'source': r.get('source'),
               'score': r.get('score'),
               'text_preview': (r.get('text') or '')[:200]}
        if r.get('source') == 'web':
            rec['title'] = r.get('doc'); rec['url'] = r.get('url')
        else:
            rec['doc_id'] = r.get('doc_id')
            rec['doc_name'] = r.get('doc')
            rec['chunk_index'] = r.get('chunk_idx')
            rec['pages'] = r.get('pages') or []
        retrieved_records.append(rec)

    record = {
        'kind': 'agent',
        'task': task_key,
        'timestamp': datetime.datetime.now().isoformat(timespec='microseconds'),
        'user_id': st.session_state.get('user_id', '_local'),
        'inputs': inputs,
        'output': output or '',
        'model': model,
        'provider': st.session_state.get('provider', ''),
        'elapsed_seconds': float(elapsed) if elapsed is not None else None,
        'retrieved': retrieved_records,
        'n_retrieved': len(retrieved_records),
    }
    try:
        with _agent_log_path().open('a', encoding='utf-8') as f:
            f.write(json.dumps(record, ensure_ascii=False, default=str) + '\n')
    except Exception as e:
        st.warning(f'에이전트 로그 저장 실패: {e}')
    _supabase_insert('agent_runs', record)


# ---- Task templates ----

def _email_build(inputs, retrieved):
    sys_msg = (
        '당신은 한국어 비즈니스 이메일 작성 도우미입니다. 사용자가 제공한 Context '
        '(사용자 문서에서 발췌)를 사실 근거로 사용해, 요청된 톤·길이·핵심 포인트에 '
        '맞춰 이메일 초안을 작성합니다. 형식은 "제목: ..." 한 줄 + 빈 줄 + 본문. '
        'Context에 사실이 부족하면 추측하지 말고 자연스럽게 일반론으로 처리하세요.'
    )
    user_msg = (
        f"작업: 이메일 초안 작성\n"
        f"수신자: {inputs.get('recipient') or '미지정'}\n"
        f"주제 힌트: {inputs.get('subject_hint') or '미지정'}\n"
        f"톤: {inputs.get('tone') or '공식'}\n"
        f"분량: {inputs.get('length') or '보통'}\n\n"
        f"꼭 포함할 핵심 요점:\n{inputs.get('key_points') or '(자유)'}\n\n"
        f"Context (참고 자료):\n{_agent_format_context(retrieved)}\n\n"
        f"이제 이메일 초안을 작성해 주세요."
    )
    return [{'role': 'system', 'content': sys_msg},
            {'role': 'user', 'content': user_msg}]


def _report_build(inputs, retrieved):
    sys_msg = (
        '당신은 보고서 작성 전문가입니다. 제공된 Context를 사실 근거로 사용해 마크다운 '
        '구조의 보고서를 작성합니다. 모든 사실 주장에는 출처 번호를 [1], [2] 형식으로 '
        '인용하세요. 섹션 제목은 ##, 소제목은 ###을 사용하세요. Context에 없는 사실은 '
        '추측하지 말고 "추가 자료 필요"로 표시하세요.'
    )
    user_msg = (
        f"작업: 보고서 작성\n"
        f"주제: {inputs.get('topic') or '미지정'}\n"
        f"대상 독자: {inputs.get('audience') or '일반'}\n"
        f"분량: {inputs.get('length') or '중간 (3~5 섹션)'}\n\n"
        f"포함해야 할 섹션·관점:\n{inputs.get('sections') or '(자동 구성)'}\n\n"
        f"Context (참고 자료):\n{_agent_format_context(retrieved)}\n\n"
        f"위 자료를 근거로 보고서를 작성해 주세요. 마크다운 구조."
    )
    return [{'role': 'system', 'content': sys_msg},
            {'role': 'user', 'content': user_msg}]


def _summary_build(inputs, retrieved):
    sys_msg = (
        '당신은 한국어 문서 요약 전문가입니다. 제공된 Context를 종합해 요청된 깊이·초점에 '
        '맞게 요약합니다. 출처는 [1], [2]로 인용하고, 핵심 포인트를 글머리 기호로 정리한 후 '
        '단락 형식의 본문 요약을 덧붙이세요.'
    )
    user_msg = (
        f"작업: 문서 요약\n"
        f"초점: {inputs.get('focus') or '전반'}\n"
        f"깊이: {inputs.get('depth') or '단락 요약'}\n\n"
        f"Context:\n{_agent_format_context(retrieved)}\n\n"
        f"요약해 주세요."
    )
    return [{'role': 'system', 'content': sys_msg},
            {'role': 'user', 'content': user_msg}]


def _analysis_build(inputs, retrieved):
    sys_msg = (
        '당신은 데이터/통찰 분석 도우미입니다. 제공된 Context에서 수치·표·핵심 사실을 '
        '추출하고, 사용자의 질문에 대해 마크다운 표 + 해석 + 통찰 순으로 답변합니다. '
        '모든 수치 인용에는 [1], [2] 형식의 출처 번호를 붙이세요. Context에 없는 추정은 '
        '"가정"으로 명시하세요.'
    )
    user_msg = (
        f"작업: 데이터/통찰 분석\n"
        f"분석 질문: {inputs.get('question') or '(미지정)'}\n"
        f"추가 맥락: {inputs.get('context_note') or '(없음)'}\n\n"
        f"Context:\n{_agent_format_context(retrieved)}\n\n"
        f"분석을 진행해 주세요. 구조: (1) 핵심 수치 표, (2) 해석, (3) 통찰/시사점."
    )
    return [{'role': 'system', 'content': sys_msg},
            {'role': 'user', 'content': user_msg}]


def _comparison_build(inputs, retrieved):
    sys_msg = (
        '당신은 비교 분석 도우미입니다. 두 대상에 대해 마크다운 비교 표를 만들고, '
        '각 항목 차이의 의미를 해석합니다. 모든 사실에 [N]으로 출처를 표기하세요. '
        'Context에 한쪽 정보가 부족하면 "정보 부족"으로 표시하세요.'
    )
    user_msg = (
        f"작업: 비교 분석\n"
        f"대상 A: {inputs.get('item_a') or '(미지정)'}\n"
        f"대상 B: {inputs.get('item_b') or '(미지정)'}\n"
        f"비교 기준 (필요시): {inputs.get('criteria') or '(자동 선택)'}\n\n"
        f"Context:\n{_agent_format_context(retrieved)}\n\n"
        f"비교 분석을 진행해 주세요. 구조: (1) 비교 표, (2) 주요 차이 해석, (3) 결론."
    )
    return [{'role': 'system', 'content': sys_msg},
            {'role': 'user', 'content': user_msg}]


AGENT_TASKS = {
    'email': {
        'label': '이메일 초안',
        'description': '문서 내용을 참고해 비즈니스 이메일 초안을 작성합니다.',
        'fields': [
            {'key': 'recipient', 'label': '수신자', 'type': 'text',
             'placeholder': '예: 박팀장님 / 거래처 ABC사 김부장님'},
            {'key': 'subject_hint', 'label': '주제 / 목적', 'type': 'text',
             'placeholder': '예: Q1 실적 공유 및 다음 분기 전략 미팅 일정 제안'},
            {'key': 'tone', 'label': '톤', 'type': 'select',
             'options': ['공식', '친근', '간결']},
            {'key': 'length', 'label': '분량', 'type': 'select',
             'options': ['짧게', '보통', '상세']},
            {'key': 'key_points', 'label': '꼭 포함할 핵심 요점', 'type': 'textarea',
             'placeholder': '한 줄에 한 가지씩 적어주세요.'},
        ],
        'requires_docs': False,
        'search_query': lambda i: f"{i.get('subject_hint', '')} {i.get('key_points', '')}".strip(),
        'build_messages': _email_build,
        'output_ext': 'md',
    },
    'report': {
        'label': '보고서 작성',
        'description': '문서를 근거로 구조화된 마크다운 보고서를 생성합니다.',
        'fields': [
            {'key': 'topic', 'label': '보고서 주제', 'type': 'text',
             'placeholder': '예: 2026 Q1 매출 분석 및 전략 제안'},
            {'key': 'audience', 'label': '대상 독자', 'type': 'text',
             'placeholder': '예: 임원진 / 실무팀 / 외부 파트너'},
            {'key': 'length', 'label': '분량', 'type': 'select',
             'options': ['짧게 (1~2 섹션)', '중간 (3~5 섹션)', '상세 (6+ 섹션)']},
            {'key': 'sections', 'label': '포함해야 할 섹션·관점', 'type': 'textarea',
             'placeholder': '비워두면 자동 구성됩니다.'},
        ],
        'requires_docs': True,
        'search_query': lambda i: f"{i.get('topic', '')} {i.get('sections', '')}".strip(),
        'build_messages': _report_build,
        'output_ext': 'md',
    },
    'summary': {
        'label': '문서 요약',
        'description': '업로드된 문서를 종합 요약합니다.',
        'fields': [
            {'key': 'focus', 'label': '요약 초점', 'type': 'text',
             'placeholder': '예: 매출과 마진 중심 / 리스크 위주 (비워두면 전반)'},
            {'key': 'depth', 'label': '깊이', 'type': 'select',
             'options': ['핵심만 (3줄)', '단락 요약', '상세 요약']},
        ],
        'requires_docs': True,
        'search_query': lambda i: i.get('focus') or '핵심 주제, 주요 논점, 결론',
        'build_messages': _summary_build,
        'output_ext': 'md',
    },
    'analysis': {
        'label': '데이터 분석',
        'description': '문서에서 수치를 뽑아 표 + 해석 + 통찰을 정리합니다.',
        'fields': [
            {'key': 'question', 'label': '분석 질문', 'type': 'text',
             'placeholder': '예: 분기별 매출 성장률과 마진 추이는?'},
            {'key': 'context_note', 'label': '추가 맥락 (선택)', 'type': 'textarea',
             'placeholder': '분석 시 고려할 추가 정보가 있으면 적어주세요.'},
        ],
        'requires_docs': True,
        'search_query': lambda i: f"{i.get('question', '')} 수치 통계 결과".strip(),
        'build_messages': _analysis_build,
        'output_ext': 'md',
    },
    'comparison': {
        'label': '비교 분석',
        'description': '두 대상을 문서 근거로 비교합니다 (여러 문서 균형 검색 자동 적용).',
        'fields': [
            {'key': 'item_a', 'label': '대상 A', 'type': 'text',
             'placeholder': '예: 자사 제품 A / 2025년 실적'},
            {'key': 'item_b', 'label': '대상 B', 'type': 'text',
             'placeholder': '예: 경쟁사 제품 B / 2026년 실적'},
            {'key': 'criteria', 'label': '비교 기준 (선택)', 'type': 'textarea',
             'placeholder': '비워두면 자동 선택.'},
        ],
        'requires_docs': True,
        'search_query': lambda i: f"{i.get('item_a', '')} {i.get('item_b', '')} 비교 차이",
        'build_messages': _comparison_build,
        'output_ext': 'md',
    },
}


def run_agent_task(task_key: str, inputs: dict, doc_ids_filter: list = None):
    """Run one agent task: retrieve → build messages → call LLM → log → return.

    If doc_ids_filter is given and shorter than the full doc list, retrieval is
    scoped to just those documents (state is swapped during retrieval and
    restored in finally).
    """
    task = AGENT_TASKS[task_key]

    query = task['search_query'](inputs)
    retrieved = []
    if query.strip():
        if st.session_state['docs']:
            all_docs = st.session_state['docs']
            # None        → use every loaded doc (selector wasn't shown)
            # empty list  → user explicitly unchecked all → retrieve nothing
            # subset      → filter to those doc_ids
            # full set    → equivalent to None, skip the swap
            need_filter = (
                doc_ids_filter is not None
                and len(doc_ids_filter) < len(all_docs)
            )
            if need_filter:
                filt_set = set(doc_ids_filter)
                st.session_state['docs'] = [d for d in all_docs if d['id'] in filt_set]
            try:
                if task_key == 'comparison':
                    _saved_bal = st.session_state.get('per_doc_balance')
                    st.session_state['per_doc_balance'] = True
                    try:
                        retrieved = retrieve_local(query.strip())
                    finally:
                        st.session_state['per_doc_balance'] = _saved_bal
                else:
                    retrieved = retrieve_local(query.strip())
            finally:
                if need_filter:
                    st.session_state['docs'] = all_docs
        if st.session_state['web_enabled']:
            retrieved = retrieved + web_search(query.strip())

    messages = task['build_messages'](inputs, retrieved)
    full_text, reasoning, elapsed = _agent_run_llm(messages, st.session_state['model'])

    # Augment inputs with doc selection for the audit log
    log_inputs = dict(inputs)
    if doc_ids_filter is not None:
        id_to_name = {d['id']: d['name'] for d in st.session_state['docs']}
        log_inputs['_selected_documents'] = [
            id_to_name.get(did, did) for did in doc_ids_filter
        ]

    _agent_log(task_key, log_inputs, full_text, retrieved,
               st.session_state['model'], elapsed)
    return full_text, reasoning, retrieved, elapsed


def view_agents():
    _section(
        '에이전트',
        '내 문서를 토대로 이메일/보고서/요약/데이터 분석/비교 분석을 자동 생성합니다. '
        '각 작업은 retrieval로 근거를 가져온 뒤 작업별 전용 프롬프트로 LLM에 요청합니다.',
    )

    # Task picker
    task_keys = list(AGENT_TASKS.keys())
    if 'agent_task_key' not in st.session_state:
        st.session_state['agent_task_key'] = task_keys[0]
    selected = st.radio(
        '작업 선택',
        task_keys,
        index=task_keys.index(st.session_state.get('agent_task_key', task_keys[0])),
        format_func=lambda k: AGENT_TASKS[k]['label'],
        horizontal=True,
    )
    st.session_state['agent_task_key'] = selected
    task = AGENT_TASKS[selected]
    st.caption(task['description'])

    if task.get('requires_docs') and not st.session_state['docs']:
        st.warning(
            '이 작업은 문서가 필요합니다. **문서** 탭에서 먼저 파일을 업로드해 주세요.'
        )

    # Input form
    inputs = {}
    selected_doc_ids = None
    with st.form(f'agent_form_{selected}'):
        # Per-task document subset selector. Shown whenever the user has
        # 2+ documents loaded — every agent task uses retrieved chunks, so
        # the choice "which docs do you want this draft to lean on" is
        # always meaningful even when requires_docs=False (e.g. email).
        if len(st.session_state['docs']) >= 2:
            doc_id_to_name = {d['id']: d['name'] for d in st.session_state['docs']}
            all_ids = list(doc_id_to_name.keys())
            selected_doc_ids = st.multiselect(
                '참고할 문서',
                options=all_ids,
                default=all_ids,
                format_func=lambda did: doc_id_to_name[did],
                help='이 작업이 검색·인용할 문서를 선택합니다. 기본은 전체 문서. '
                '특정 문서만 골라 좁히면 응답이 더 정확해집니다. '
                '문서가 1개라면 선택지가 없으니 자동 사용됩니다.',
            )

        for f in task['fields']:
            label = f['label']
            placeholder = f.get('placeholder', '')
            if f['type'] == 'text':
                inputs[f['key']] = st.text_input(label, placeholder=placeholder)
            elif f['type'] == 'textarea':
                inputs[f['key']] = st.text_area(label, height=120, placeholder=placeholder)
            elif f['type'] == 'select':
                inputs[f['key']] = st.selectbox(label, f['options'])
        submitted = st.form_submit_button('실행', use_container_width=True, type='primary')

    if submitted:
        if (not st.session_state['model']
                or not _active_api_key()
                or not st.session_state['base_url']):
            provider = st.session_state.get('provider', 'Hugging Face Router')
            st.error(
                f'설정 탭에서 공급자/모델/{provider} 용 API 키를 먼저 확인해 주세요.'
            )
            st.stop()
        if task.get('requires_docs') and not st.session_state['docs']:
            st.stop()
        if (task.get('requires_docs') and selected_doc_ids is not None
                and len(selected_doc_ids) == 0):
            st.error('최소 한 개 이상의 문서를 선택해 주세요.')
            st.stop()

        try:
            full_text, reasoning, retrieved, elapsed = run_agent_task(
                selected, inputs, doc_ids_filter=selected_doc_ids,
            )
        except Exception as e:
            _show_llm_error(e)
            st.stop()

        st.write('')
        _section('결과', f'생성 시간: {elapsed:.1f}s')
        if reasoning:
            with st.expander('추론 과정', expanded=False):
                st.markdown(reasoning)
        st.markdown(full_text or '*(빈 응답)*')

        if full_text:
            stamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            ext = task.get('output_ext', 'md')
            mime = 'text/markdown' if ext == 'md' else 'text/plain'
            st.download_button(
                '다운로드',
                data=full_text,
                file_name=f"{selected}_{stamp}.{ext}",
                mime=mime,
                use_container_width=True,
                key=f'dl_agent_{selected}_{stamp}',
            )

        if retrieved:
            with st.expander(f'사용된 근거 ({len(retrieved)}개)', expanded=False):
                for j, r in enumerate(retrieved, start=1):
                    if r.get('source') == 'web':
                        from urllib.parse import urlparse
                        host = urlparse(r.get('url', '')).netloc or '웹'
                        st.markdown(f"**[{j}] 웹 · {host}** — {r.get('doc', '')[:80]}")
                        if r.get('url'):
                            st.markdown(f"<{r.get('url')}>")
                    else:
                        pages = r.get('pages') or []
                        page_str = ''
                        if pages:
                            page_str = (f' p.{pages[0]}' if len(pages) == 1
                                        else f' pp.{pages[0]}-{pages[-1]}')
                        st.markdown(
                            f"**[{j}] {r.get('doc', '')}{page_str}** · "
                            f"chunk {r.get('chunk_idx', 0)} · "
                            f"score {r.get('score', 0):.3f}"
                        )
                    st.text((r.get('text', '') or '')[:600])
                    st.divider()


# =============================================================================
# About view
# =============================================================================

def view_about():
    if _LOGO_URI:
        st.markdown(
            f'<div style="text-align:center; margin-bottom:18px;">'
            f'<img src="{_LOGO_URI}" '
            f'style="width:100%; max-width:280px; height:auto; display:block; margin:0 auto 8px auto;" />'
            f'<div style="font-size:22px; font-weight:700; letter-spacing:-0.01em;">Personal RAG</div>'
            f'<div style="font-size:13px; color:rgba(128,128,128,0.95);">내 문서 기반 질의응답 시스템</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    st.markdown(
        """
        내 문서를 기반으로 답하는 개인용 RAG (Retrieval-Augmented Generation) 시스템.

        #### 검색 파이프라인
        1. 문서 파싱 — PDF는 **Docling**으로 표·헤딩·리스트 구조를 보존하며 마크다운으로 변환, 페이지 메타를 청크별로 기록.
        2. 청킹 — 문단 단위로 묶고 size 한도를 넘으면 슬라이딩 윈도우로 분할.
        3. 임베딩 — **BGE-M3** (다국어 SOTA) 또는 MiniLM multilingual.
        4. 하이브리드 검색 — Dense 코사인 + BM25 의 **Reciprocal Rank Fusion**.
        5. (옵션) **HyDE** / **Multi-query** / **Contextual rewrite** 로 쿼리 보강.
        6. (옵션) Cross-encoder **bge-reranker-v2-m3** 로 top_n → top_k 재정렬.
        7. (옵션) PDF 페이지 이미지를 **멀티모달 LLM** 에 함께 전달.
        8. LLM 응답을 토큰 단위로 스트리밍, reasoning_content 와 본문을 분리.
        9. 답변 안의 `[N]` 인용 마커를 파싱하여 청크별로 출처 표시.

        #### 저장
        - `./.data/{embedder}/{doc_hash}/` — 임베더별 격리.
        - `meta.json` (청크, 페이지 메타, 원본 텍스트) + `embeddings.npy` + `pages/{page}.png`.
        - 같은 파일·청크 설정으로 다시 업로드 시 캐시에서 즉시 복원.
        - 사용자별 환경설정 (API 키, 모델, 검색 설정 등) 영속 저장: `./.data/{user}/preferences.json`. Streamlit Cloud의 idle 재연결 시 메모리가 초기화되더라도 다음 접속에서 자동 복원됩니다. (Cloud 컨테이너가 재시작되면 사라지므로 영구 보존이 필요하면 Streamlit Secrets 사용 권장.)
        - 대화 세션 메타: `./.data/{user}/sessions/{id}.json` (사이드바 대화 목록의 원천).
        - **대화 로그**: `./logs/{user}/{session_id}.jsonl` — 세션별로 한 파일, 한 줄당 한 턴. 분석·DB 친화적 구조.
          필드: session_id, turn_index, timestamp, user_message, assistant_message, reasoning, model/provider, elapsed_seconds, retrieved (rank·source·score·page·url), citation_numbers_used, query_variants, settings_snapshot (rerank/HyDE/multi-query/per-doc 등 모든 설정 스냅샷).
        - **에이전트 실행 로그**: `./logs/{user}/agents.jsonl` — 에이전트 워크플로(이메일/보고서/요약/분석/비교)의 실행 기록. task, inputs, output, retrieved, model, elapsed_seconds.
        - **이벤트 로그**: `./logs/{user}/events.jsonl` — 로그인/로그아웃, 문서 업로드·삭제, 세션 삭제, LLM 호출 실패 등 비-턴 이벤트. 각 줄에 event_type, timestamp, user_id, payload.
          pandas: `pd.concat([pd.read_json(f, lines=True) for f in glob.glob('logs/*/*.jsonl')])`.
          Postgres: `\\copy turns FROM '...' WITH (FORMAT json)` 또는 batched insert.

        #### 지원 LLM 엔드포인트
        Hugging Face Inference Router · OpenAI · Alibaba DashScope · vLLM / 로컬 OpenAI-호환 서버.

        #### 사용 라이브러리
        Streamlit, OpenAI SDK, sentence-transformers, rank_bm25, Docling, PyMuPDF, pypdf, ddgs.
        """
    )


# =============================================================================
# View router — dispatch to selected view
# =============================================================================

_VIEWS = {
    'chat':     view_chat,
    'docs':     view_docs,
    'agents':   view_agents,
    'settings': view_settings,
    'cache':    view_cache,
    'about':    view_about,
}
_VIEWS.get(st.session_state.get('active_view', 'chat'), view_chat)()

# Persist user preferences after every rerun (cheap; only writes on change).
_save_user_prefs()
