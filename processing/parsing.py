"""Format-specific parsers + the dispatch entrypoint ``parse_file``.

Every parser returns the same dict shape:

    {'elements': [{'text': str, 'page': int | None}, ...],
     'page_count': int,
     'ok': bool}

so the rest of the pipeline can stay format-agnostic. ``parse_file`` is
the front door: it sniffs the upload's filename, runs the right parser
(with graceful fallbacks for PDFs — Docling → pypdf → Tesseract OCR),
and wraps the result with ``raw_text`` + ``pdf_bytes`` for downstream
chunking / page-image rendering.
"""

import os
from pathlib import Path

import streamlit as st


def _read_bytes(file) -> bytes:
    """Read full bytes from a Streamlit UploadedFile, resetting position."""
    pos = file.tell() if hasattr(file, 'tell') else 0
    data = file.read()
    try:
        file.seek(pos)
    except Exception:
        pass
    return data


def parse_pdf_docling(pdf_bytes: bytes) -> dict:
    """Parse PDF with Docling. Returns {elements, page_count, ok}.

    elements is a flat list of {text, page} ordered as on the page.
    Tables are emitted as markdown table strings; lists/headings preserved.
    """
    import tempfile
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return {'elements': [], 'page_count': 0, 'ok': False}
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        d = result.document
    except Exception as e:
        st.warning(f'Docling 파싱 실패, pypdf로 폴백합니다: {e}')
        return {'elements': [], 'page_count': 0, 'ok': False}
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    elements = []
    pages_seen = set()
    for item, _level in d.iterate_items():
        text = ''
        try:
            if hasattr(item, 'export_to_markdown'):
                # Tables, lists, headings: structured markdown.
                text = item.export_to_markdown(d)
        except Exception:
            text = ''
        if not text and hasattr(item, 'text'):
            text = item.text or ''
        text = (text or '').strip()
        if not text:
            continue
        page = None
        prov = getattr(item, 'prov', None)
        if prov:
            try:
                page = prov[0].page_no
            except Exception:
                page = None
        if page is not None:
            pages_seen.add(page)
        elements.append({'text': text, 'page': page})

    page_count = max(pages_seen) if pages_seen else 0
    return {'elements': elements, 'page_count': page_count, 'ok': True}


def parse_pdf_pypdf(pdf_bytes: bytes) -> dict:
    """Fallback PDF parser. Page boundaries known but no structure."""
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(pdf_bytes))
    elements = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or '').strip()
        if text:
            elements.append({'text': text, 'page': i + 1})
    return {'elements': elements, 'page_count': len(reader.pages), 'ok': True}


def parse_pdf_ocr(pdf_bytes: bytes) -> dict:
    """Last-resort PDF parser for scanned / image-only PDFs.

    Renders each page to PNG via PyMuPDF and runs Tesseract OCR on it.
    Returns {} when the OCR stack isn't installed so the caller can fall
    through gracefully. Korean + English are both enabled.
    """
    try:
        import fitz  # PyMuPDF
        import pytesseract
        from PIL import Image
        import io
    except Exception:
        return {'elements': [], 'page_count': 0, 'ok': False,
                'reason': 'ocr_dependencies_missing'}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype='pdf')
    except Exception:
        return {'elements': [], 'page_count': 0, 'ok': False,
                'reason': 'pdf_open_failed'}
    elements = []
    page_count = len(doc)
    for i, page in enumerate(doc):
        try:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes('png')))
            # Korean + English. Falls back to English-only if kor pack missing.
            try:
                text = pytesseract.image_to_string(img, lang='kor+eng')
            except Exception:
                text = pytesseract.image_to_string(img)
            text = (text or '').strip()
            if text:
                elements.append({'text': text, 'page': i + 1})
        except Exception:
            continue
    doc.close()
    return {'elements': elements, 'page_count': page_count, 'ok': True}


def parse_docx_bytes(data: bytes) -> dict:
    """Parse DOCX (Word) into the same elements/page format we use for PDFs.

    Tries Docling first (better structure / table support), falls back to
    python-docx (simple paragraph extraction). DOCX doesn't have stable
    page numbers in the XML, so we leave page=None — chunking still works,
    citations just won't show page numbers."""
    # Docling path
    try:
        from docling.document_converter import DocumentConverter
        import io
        conv = DocumentConverter()
        # Docling accepts a path or a stream-like with name attribute.
        class _NamedBytesIO(io.BytesIO):
            name = 'in.docx'
        buf = _NamedBytesIO(data)
        result = conv.convert(buf)
        md = result.document.export_to_markdown()
        if md and md.strip():
            paragraphs = [p.strip() for p in md.split('\n\n') if p.strip()]
            return {
                'elements': [{'text': p, 'page': None} for p in paragraphs],
                'page_count': 0,
                'ok': True,
            }
    except Exception:
        pass

    # python-docx fallback
    try:
        from docx import Document
        import io
        d = Document(io.BytesIO(data))
        elements = []
        for p in d.paragraphs:
            t = (p.text or '').strip()
            if t:
                elements.append({'text': t, 'page': None})
        # Tables: each cell becomes its own element so a chunk can hold
        # one row's worth of facts.
        for table in d.tables:
            for row in table.rows:
                cells = [(c.text or '').strip() for c in row.cells]
                joined = ' | '.join(c for c in cells if c)
                if joined:
                    elements.append({'text': joined, 'page': None})
        if elements:
            return {'elements': elements, 'page_count': 0, 'ok': True}
    except Exception:
        pass

    return {'elements': [], 'page_count': 0, 'ok': False}


def parse_csv_bytes(data: bytes, file_name: str = '') -> dict:
    """Parse CSV. Each row becomes one element formatted as
    'col1: val1 | col2: val2 | ...' so the LLM sees field-name context
    around each value, not just numbers in isolation. Also includes a
    header-summary element so semantic search can match queries like
    "what columns does this file have"."""
    import csv as _csv
    import io as _io

    # Best-effort decoding: try UTF-8 (with BOM), fall back to CP949
    # which is what Excel saves CSV as on Korean Windows.
    for enc in ('utf-8-sig', 'utf-8', 'cp949', 'euc-kr', 'latin-1'):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        return {'elements': [], 'page_count': 0, 'ok': False}

    # Sniff delimiter — fall back to comma if sniffing fails.
    sample = text[:4096]
    try:
        dialect = _csv.Sniffer().sniff(sample, delimiters=',;\t|')
    except _csv.Error:
        class _Default(_csv.Dialect):
            delimiter = ','
            quotechar = '"'
            doublequote = True
            skipinitialspace = True
            lineterminator = '\n'
            quoting = _csv.QUOTE_MINIMAL
        dialect = _Default

    reader = _csv.reader(_io.StringIO(text), dialect)
    rows = list(reader)
    if not rows:
        return {'elements': [], 'page_count': 0, 'ok': True}

    header = [(c or '').strip() for c in rows[0]]
    data_rows = rows[1:]
    elements = []

    # 1) Summary element — column names + row count.
    summary = (
        f"[CSV 요약] 파일: {file_name or 'csv'}, "
        f"열 {len(header)}개, 행 {len(data_rows)}개. "
        f"열 이름: {', '.join(c for c in header if c)}."
    )
    elements.append({'text': summary, 'page': None})

    # 2) Per-row elements — formatted as "header: value | header: value".
    for i, row in enumerate(data_rows, start=1):
        cells = []
        for h, v in zip(header, row):
            v = (v or '').strip()
            if not v:
                continue
            if h:
                cells.append(f'{h}: {v}')
            else:
                cells.append(v)
        if cells:
            elements.append({
                'text': f'행 {i}: ' + ' | '.join(cells),
                'page': None,
            })

    return {'elements': elements, 'page_count': 0, 'ok': True}


def parse_hwpx_bytes(data: bytes) -> dict:
    """Parse HWPX (Hancom Office XML). HWPX is a ZIP archive containing
    section XML files under Contents/. We pull <hp:t> (text run) elements
    from each section in order and treat each section as a 'page' so the
    citation system can still surface a section number."""
    import io as _io
    import zipfile as _zip
    from xml.etree import ElementTree as _ET

    try:
        zf = _zip.ZipFile(_io.BytesIO(data))
    except _zip.BadZipFile:
        return {'elements': [], 'page_count': 0, 'ok': False}

    section_names = sorted(
        (n for n in zf.namelist()
         if n.startswith('Contents/section') and n.endswith('.xml')),
        key=lambda n: n,
    )
    if not section_names:
        zf.close()
        return {'elements': [], 'page_count': 0, 'ok': False}

    elements = []
    # HWPX uses the 'hp' namespace; strip namespaces so we don't have to
    # bind them — just match local tag names.
    def _local(tag: str) -> str:
        return tag.split('}', 1)[1] if '}' in tag else tag

    for section_idx, name in enumerate(section_names, start=1):
        try:
            raw = zf.read(name)
            root = _ET.fromstring(raw)
        except Exception:
            continue
        # Each <p> (paragraph) becomes one element. Within a paragraph we
        # concatenate the text of all <t> (text run) descendants.
        for p in root.iter():
            if _local(p.tag) != 'p':
                continue
            buf = []
            for t in p.iter():
                if _local(t.tag) == 't' and t.text:
                    buf.append(t.text)
            joined = ''.join(buf).strip()
            if joined:
                elements.append({'text': joined, 'page': section_idx})

    zf.close()
    return {'elements': elements, 'page_count': len(section_names), 'ok': True}


def parse_file(file) -> dict:
    """Returns a dict: {raw_text, elements, page_count, is_pdf, pdf_bytes}.

    elements: list of {text, page}.  page is int (1-indexed) for PDFs,
    None for plain text files.  pdf_bytes is retained so the caller can
    render page images later (only set when is_pdf=True).
    """
    name = file.name.lower()
    if name.endswith('.pdf'):
        data = _read_bytes(file)
        # Stage 1: Docling (structure-aware)
        parsed = parse_pdf_docling(data)
        # Stage 2: pypdf if Docling missed
        if not parsed['ok'] or not parsed['elements']:
            try:
                parsed = parse_pdf_pypdf(data)
            except Exception as e:
                parsed = {'elements': [], 'page_count': 0, 'ok': False,
                          'error': str(e)}
        # Stage 3: OCR if both text extractors got essentially nothing —
        # this catches scanned / image-only PDFs that don't have a text
        # layer to extract.
        total_chars = sum(len(e['text']) for e in parsed.get('elements', []))
        if total_chars < 50:
            try:
                with st.spinner(f'{file.name}: 텍스트 레이어 없음 — OCR 시도 중...'):
                    ocr = parse_pdf_ocr(data)
                if ocr['ok'] and ocr['elements']:
                    parsed = ocr
                    st.info(
                        f'{file.name}: 스캔 PDF 로 판단해 OCR (Tesseract) 로 텍스트 추출했습니다. '
                        f'페이지 {ocr["page_count"]}, 문자 {sum(len(e["text"]) for e in ocr["elements"])}.'
                    )
                elif ocr.get('reason') == 'ocr_dependencies_missing':
                    st.warning(
                        f'{file.name}: 텍스트 추출이 안 되고 OCR 라이브러리도 없어 인덱스 품질이 낮을 수 있습니다. '
                        f'`pip install pytesseract pymupdf` + 시스템 tesseract 설치 시 스캔 PDF 도 처리됩니다.'
                    )
            except Exception:
                pass
        if not parsed.get('elements'):
            st.error(f'PDF 파싱 실패 ({file.name}): 텍스트를 추출할 수 없습니다.')
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': True, 'pdf_bytes': data}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': parsed['page_count'],
            'is_pdf': True,
            'pdf_bytes': data,
        }
    if name.endswith('.docx'):
        data = _read_bytes(file)
        parsed = parse_docx_bytes(data)
        if not parsed['ok']:
            st.error(
                f'DOCX 파싱 실패 ({file.name}). '
                f'`pip install python-docx` 가 필요할 수 있습니다.'
            )
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': 0,
            'is_pdf': False,
            'pdf_bytes': None,
        }
    if name.endswith('.doc'):
        st.error(
            f'DOC (구버전 바이너리) 는 직접 지원하지 않습니다 ({file.name}). '
            f'Word 또는 LibreOffice 에서 DOCX 또는 PDF 로 다시 저장한 뒤 업로드해 주세요.'
        )
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    if name.endswith('.csv'):
        data = _read_bytes(file)
        parsed = parse_csv_bytes(data, file_name=file.name)
        if not parsed['ok']:
            st.error(f'CSV 디코딩 실패 ({file.name}): 인코딩이 UTF-8/CP949 가 아닌 듯합니다.')
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': 0,
            'is_pdf': False,
            'pdf_bytes': None,
        }
    if name.endswith('.hwpx'):
        data = _read_bytes(file)
        parsed = parse_hwpx_bytes(data)
        if not parsed['ok']:
            st.error(
                f'HWPX 파싱 실패 ({file.name}): 손상된 파일이거나 보안 처리된 '
                f'HWPX 일 수 있습니다. 한컴오피스에서 다시 저장하거나 PDF 로 '
                f'변환해 보세요.'
            )
            return {'raw_text': '', 'elements': [], 'page_count': 0,
                    'is_pdf': False, 'pdf_bytes': None}
        raw = '\n\n'.join(e['text'] for e in parsed['elements'])
        return {
            'raw_text': raw,
            'elements': parsed['elements'],
            'page_count': parsed['page_count'],  # section count, used like "page"
            'is_pdf': False,
            'pdf_bytes': None,
        }
    # Reject .hwp (binary) explicitly with a helpful message — different format.
    if name.endswith('.hwp'):
        st.error(
            f'HWP (구버전 바이너리) 는 직접 지원하지 않습니다 ({file.name}). '
            f'한컴오피스에서 HWPX 또는 PDF 로 다시 저장한 뒤 업로드해 주세요.'
        )
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    try:
        text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        st.error(f'파일 읽기 실패 ({file.name}): {e}')
        return {'raw_text': '', 'elements': [], 'page_count': 0,
                'is_pdf': False, 'pdf_bytes': None}
    return {
        'raw_text': text,
        'elements': [{'text': text, 'page': None}] if text.strip() else [],
        'page_count': 0,
        'is_pdf': False,
        'pdf_bytes': None,
    }


def render_pdf_pages_to_dir(pdf_bytes: bytes, out_dir: Path, dpi: int = 144) -> int:
    """Render each PDF page to PNG into out_dir/{page}.png. Returns number of pages."""
    import fitz
    out_dir.mkdir(parents=True, exist_ok=True)
    n = 0
    with fitz.open(stream=pdf_bytes, filetype='pdf') as d:
        for i, page in enumerate(d):
            pix = page.get_pixmap(dpi=dpi)
            pix.save(str(out_dir / f'{i + 1}.png'))
            n = i + 1
    return n
